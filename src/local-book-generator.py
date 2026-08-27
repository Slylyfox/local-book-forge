import sys

# Force UTF-8 stdout/stderr before anything else runs. On Windows, the default
# console codepage (cp1252) can't encode emoji/unicode that CrewAI's console
# output uses, causing 'charmap' codec errors from CrewAI's internal event bus.
# This must happen before crewai is imported, since some of its output can
# fire on import.
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", line_buffering=True)
    sys.stderr.reconfigure(encoding="utf-8", line_buffering=True)

import argparse
import base64
import json
import os
import random
import re
import shutil
import subprocess
import time
from datetime import datetime
from typing import Type

import requests
from pydantic import BaseModel, Field
from PIL import Image, ImageDraw, ImageFont, ImageOps, ImageFilter, ImageChops
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from crewai import Agent, Crew, Process, Task, LLM
from crewai.tools import BaseTool

# =====================================================================
# 0. CONFIG
# =====================================================================

# Book length, retuned 2026-08-14 after the first clean run came in at 11,604
# words — 29% of the allowed budget. Nothing in the old constraints pushed the
# outline model toward the upper end of 5-15 chapters, so it took the cheap path
# every time. That was masked until now: before num_predict was actually
# enforced, runaway generation accidentally padded books to ~25k words. With the
# cap working, a book gets exactly what the outline asks for, so the outline has
# to ask for a publishable length.
MIN_CHAPTERS = 10          # what the outline is ASKED for
ABSOLUTE_MIN_CHAPTERS = 5  # below this, treat the response as malformed and retry
MAX_CHAPTERS = 15
MIN_TOTAL_WORDS = 30000
MAX_TOTAL_WORDS = 40000

# Per-chapter ceiling, and this number is the important one.
#
# Extra length must come from MORE CHAPTERS, not longer ones. Measured on the
# 2026-08-13 18:37 run, worst-window vocabulary richness against chapter length:
#
#     1,249 words -> 0.562        1,893 words -> 0.560
#     1,762 words -> 0.590        2,343 words -> 0.515
#     4,357 words -> 0.357   <-- the only chapter anywhere near the 0.36 floor
#
# Collapse risk climbs with chapter length, which stands to reason: the longer a
# single generation runs, the more of its own output is in context and the more
# likely it is to start echoing itself. So 3,000 is the cap rather than the old
# 5,000 — a 33,000-word book is 11 chapters of 3,000, not 7 of 4,700. Same total
# length, materially lower risk per chapter, and it keeps every chapter inside
# the band that measured healthiest.
MAX_CHAPTER_TARGET_WORDS = 3000
MIN_CHAPTER_TARGET_WORDS = 1800


# Curated genre pool. Python (not the LLM) picks the genre each run —
# local models tend to default to the same genre repeatedly if left
# to "choose freely," so we force variety here instead.
GENRE_POOL = [
    "sci-fi thriller",
    "cozy mystery",
    "epic fantasy",
    "romantic suspense",
    "psychological horror",
    "dystopian survival",
    "noir detective",
    "space opera",
]

# Reopen stdout/stderr line-buffered so output streams live to the dashboard
# even if something upstream (e.g. Windows console redirection) buffers by
# default. The "-u" flag the dashboard launches with already covers most
# cases, but this makes the script robust when run standalone too.
try:
    sys.stdout.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
except AttributeError:
    pass  # Python < 3.7 fallback, not expected here

parser = argparse.ArgumentParser(description="Local AI book generator.")
parser.add_argument(
    "--genre",
    choices=GENRE_POOL,
    default=None,
    help="Genre to write. Omit for a random pick from the genre pool.",
)
parser.add_argument(
    "--writer-model",
    default=None,
    help="Override WRITER_MODEL_NAME for this run only, e.g. to A/B test a different local "
         "model against the default (dashboard section 6.12). Must already exist on this "
         "Ollama instance. Passing this also skips title-history dedupe/recording for this "
         "run (see the dedupe_title() call site below) so testing the same premise across "
         "models doesn't trip the duplicate-title guard against itself, and the model actually "
         "used is logged to model_info.json in the book's output folder either way.",
)
parser.add_argument(
    "--min-chapters", type=int, default=None,
    help="Override the minimum chapter count asked of the outline model.",
)
parser.add_argument(
    "--max-chapters", type=int, default=None,
    help="Override the maximum chapter count.",
)
parser.add_argument(
    "--min-total-words", type=int, default=None,
    help="Override the whole-book word floor. An outline coming in under this is scaled up "
         "(and padded with extra chapters if scaling alone can't reach it).",
)
parser.add_argument(
    "--max-total-words", type=int, default=None,
    help="Override the whole-book word ceiling.",
)
parser.add_argument(
    "--min-chapter-words", type=int, default=None,
    help="Override the per-chapter target floor.",
)
parser.add_argument(
    "--max-chapter-words", type=int, default=None,
    help="Override the per-chapter target ceiling. Raising this above ~3000 measurably "
         "increases repetition-collapse risk — see MAX_CHAPTER_TARGET_WORDS for the numbers.",
)
args = parser.parse_args()

RUN_GENRE = args.genre if args.genre else random.choice(GENRE_POOL)
WRITER_MODEL_OVERRIDE = args.writer_model or None

# CLI overrides from the dashboard's Book Structure panel, applied here so every
# downstream reference picks them up. Validated rather than trusted: the ordering
# constraints matter (a floor above a ceiling would make parse_outline
# unsatisfiable and loop the outline retries), and the per-chapter ceiling has a
# real quality consequence, so exceeding the tested band warns loudly instead of
# silently accepting it.
_structure_overrides = {
    "MIN_CHAPTERS": args.min_chapters,
    "MAX_CHAPTERS": args.max_chapters,
    "MIN_TOTAL_WORDS": args.min_total_words,
    "MAX_TOTAL_WORDS": args.max_total_words,
    "MIN_CHAPTER_TARGET_WORDS": args.min_chapter_words,
    "MAX_CHAPTER_TARGET_WORDS": args.max_chapter_words,
}
for _name, _value in _structure_overrides.items():
    if _value is not None:
        globals()[_name] = _value
        print(f"[Config] {_name} overridden to {_value}")

if MIN_CHAPTERS > MAX_CHAPTERS:
    print(f"[FATAL] --min-chapters ({MIN_CHAPTERS}) is above --max-chapters ({MAX_CHAPTERS}).")
    sys.exit(1)
if MIN_CHAPTER_TARGET_WORDS > MAX_CHAPTER_TARGET_WORDS:
    print(f"[FATAL] --min-chapter-words ({MIN_CHAPTER_TARGET_WORDS}) is above "
          f"--max-chapter-words ({MAX_CHAPTER_TARGET_WORDS}).")
    sys.exit(1)
if MIN_TOTAL_WORDS > MAX_TOTAL_WORDS:
    print(f"[FATAL] --min-total-words ({MIN_TOTAL_WORDS}) is above --max-total-words "
          f"({MAX_TOTAL_WORDS}).")
    sys.exit(1)
# The word floor has to be reachable with the chapter budget available, or
# parse_outline can never satisfy it and every outline attempt gets rejected.
_max_reachable = MAX_CHAPTERS * MAX_CHAPTER_TARGET_WORDS
if MIN_TOTAL_WORDS > _max_reachable:
    print(f"[FATAL] --min-total-words ({MIN_TOTAL_WORDS}) is unreachable: at most "
          f"{MAX_CHAPTERS} chapters x {MAX_CHAPTER_TARGET_WORDS} words = {_max_reachable}. "
          f"Raise --max-chapters or --max-chapter-words, or lower the floor.")
    sys.exit(1)
if MAX_CHAPTER_TARGET_WORDS > 3500:
    print(f"[WARN] --max-chapter-words is {MAX_CHAPTER_TARGET_WORDS}. Measured on real runs, "
          f"chapters over ~3,000 words are where repetition collapse starts showing up (a "
          f"4,357-word chapter measured 0.357 against a 0.36 floor, while 1,200-2,400-word "
          f"chapters held 0.51-0.59). The collapse guard will catch it, but expect more rerolls "
          f"and longer runs. Prefer more chapters over longer ones.")

OUTPUT_ROOT = "output_books"

# --- Cover generation config ---
# Pen names are fixed per genre so branding/review history stays consistent for each
# genre over time. Every entry in GENRE_POOL has a permanent pen name assigned below;
# DEFAULT_AUTHOR_NAME only matters as a safety net if GENRE_POOL ever grows a new genre
# before this dict is updated to match.
DEFAULT_AUTHOR_NAME = "Your Pen Name"  # <-- only used if a genre is missing from GENRE_PEN_NAMES

GENRE_PEN_NAMES = {
    "sci-fi thriller": "Marcus Dax",
    "cozy mystery": "Penelope Croft",
    "epic fantasy": "Soren Ashcroft",
    "romantic suspense": "Nora Wilde",
    "psychological horror": "Arthur Blackwood",
    "dystopian survival": "Mason Cross",
    "noir detective": "Frankie Marlowe",
    "space opera": "Helena Voss",
}


def get_author_name(genre: str) -> str:
    """Pen name for this run's genre, falling back to DEFAULT_AUTHOR_NAME if this
    genre has no dedicated entry in GENRE_PEN_NAMES (or the entry is blank)."""
    return GENRE_PEN_NAMES.get(genre) or DEFAULT_AUTHOR_NAME


AUTHOR_NAME = get_author_name(RUN_GENRE)
print(f"[Config] Genre: {RUN_GENRE} | Pen name: {AUTHOR_NAME}")

# SDXL was trained on specific bucket resolutions; forcing arbitrary aspect ratios
# degrades quality. 832x1216 is an official SDXL bucket close to the standard ebook cover ratio.
BASE_GEN_WIDTH = 832
BASE_GEN_HEIGHT = 1216

# Standard ebook cover ratio is 1.6:1 (height:width). 1600x2560 comfortably
# clears the 1000px-shortest-side minimum most storefronts ask for, at good
# print/preview quality.
TARGET_COVER_SIZE = (1600, 2560)  # (width, height)

# Optional: point this at a bold display/serif .ttf you like better. If it doesn't
# exist, the tool falls back through a list of common system fonts, then to PIL's
# default bitmap font (ugly, but the run won't crash).
CUSTOM_TITLE_FONT_PATH = ""  # e.g. "/home/you/fonts/PlayfairDisplay-Bold.ttf"

# --- Manuscript (.docx) formatting config ---
BODY_FONT_NAME = "Garamond"   # common self-published fiction body font; swap if you prefer
BODY_FONT_SIZE_PT = 12
INCLUDE_TOC = True            # static chapter-list Contents page (always renders, incl. after ebook conversion)
COPYRIGHT_BOILERPLATE = (
    "This is a work of fiction. Names, characters, places, and incidents either are "
    "the product of the author's imagination or are used fictitiously. Any "
    "resemblance to actual events, locales, or persons, living or dead, is entirely "
    "coincidental."
)

GENRE_STYLE_HINTS = {
    "sci-fi thriller": "cinematic sci-fi, cool blue and neon accent lighting, sleek "
                        "technology, sense of paranoia and scale",
    "cozy mystery": "warm inviting illustrated style, soft muted palette, small-town "
                     "or village setting, gentle whimsy, no violence shown",
    "epic fantasy": "painterly fantasy illustration, dramatic sweeping landscape, "
                     "rich saturated colors, sense of scale and myth",
    "romantic suspense": "moody cinematic lighting, high contrast, one or two "
                          "silhouetted figures, tension between romance and danger",
    "psychological horror": "dark desaturated palette, unsettling atmosphere, "
                             "heavy shadow, restrained/suggestive rather than graphic",
    "dystopian survival": "gritty desaturated tones, harsh lighting, ruined or "
                           "stark environment, isolation, survival tension",
    "noir detective": "high-contrast black and white or muted sepia, rain-slicked "
                       "streets, dramatic shadow, 1940s noir cinematography",
    "space opera": "vast cosmic vistas, dramatic starfields, bold saturated color, "
                    "sense of scale and adventure",
}

# One small, high-contrast detail to nudge the art prompt toward — the kind of thing
# that makes a cover feel art-directed rather than generic. Local SD/SDXL won't
# reliably obey this level of specific instruction every time (it's a nudge, not a
# guarantee), but it costs nothing to ask for and sometimes lands.
GENRE_DETAIL_HINTS = {
    "sci-fi thriller": "one small glowing tech detail (an interface light, a lens "
                        "flare off metal) catching the eye against the darker scene",
    "cozy mystery": "one small charming prop detail (a teacup, a cat, a garden gate) "
                     "that rewards a closer look",
    "epic fantasy": "one small glinting detail (a rune, a blade's edge, jewelry) "
                     "catching the light against the grand backdrop",
    "romantic suspense": "one small sharp glinting detail on the figure (jewelry, a "
                          "key, a ring) catching the light, and if there's a "
                          "reflective surface, let the reflection feel subtly off",
    "psychological horror": "one small wrong detail that rewards a second look "
                             "(an object slightly out of place, an unnatural shadow)",
    "dystopian survival": "one small worn/damaged detail (scavenged gear, a cracked "
                           "surface) that grounds the scene in hardship",
    "noir detective": "one small glinting detail (a blade, a lighter, rain-lit chrome) "
                       "and if there's wet pavement, let the reflection feel subtly off",
    "space opera": "one small glowing detail (running lights, an energy trail) that "
                    "gives a sense of scale against the cosmic backdrop",
}

# =====================================================================
# 0b. TITLE HISTORY / DUPLICATE HANDLING
# =====================================================================
# The local outline model has a limited creative range and will occasionally
# regenerate a title it's already used in an earlier run. This file tracks every
# title ever produced so a repeat can be caught and turned into "Part 2" / "Part 3"
# instead of silently shipping two unrelated books under the identical title. It's
# a plain background JSON file for now — no dashboard UI yet — but "show history" /
# "clear history" buttons are planned, and they'll just read/write this same file.
TITLE_HISTORY_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "title_history.json")
MAX_TITLE_PARTS = 3  # Part 1 (original) + Part 2 + Part 3, then fall back to a disambiguated title


def _normalize_title(title: str) -> str:
    return re.sub(r"\s+", " ", title.strip().lower())


def _load_title_history() -> dict:
    if os.path.isfile(TITLE_HISTORY_PATH):
        try:
            with open(TITLE_HISTORY_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
        except (json.JSONDecodeError, OSError):
            print(f"[WARN] Could not read {TITLE_HISTORY_PATH} — starting a fresh title history.")
            return {}
        # Type check added 2026-08-14. The existing guard caught unparseable
        # files but not files that parse fine into the WRONG TYPE, and that gap
        # was real: project_cleanup.py reset this to "[]" instead of "{}", which
        # is valid JSON, so it sailed past the try/except and then died in
        # dedupe_title() with "AttributeError: 'list' object has no attribute
        # 'get'" — after the outline stage had already run. Title dedupe is a
        # convenience; it must never be able to abort a book.
        if not isinstance(data, dict):
            print(f"[WARN] {TITLE_HISTORY_PATH} contains a {type(data).__name__}, but the title "
                  f"history must be a JSON object keyed by title. Ignoring it and starting a "
                  f"fresh history — duplicate titles may not be caught on this run.")
            return {}
        return data
    return {}


def _save_title_history(history: dict) -> None:
    try:
        with open(TITLE_HISTORY_PATH, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=2)
    except OSError as e:
        print(f"[WARN] Could not save title history ({e}) — a repeat title may not be caught next run.")


def dedupe_title(title: str, genre: str) -> str:
    """Check `title` against every title used in past runs (case/whitespace-insensitive).
    First time seeing it: record and return unchanged. Seen once before: return
    '<title> Part 2'. Seen twice before: return '<title> Part 3'. Beyond that (rare),
    fall back to a date-stamped disambiguator rather than inventing a 'Part 4' and
    breaking the trilogy framing."""
    history = _load_title_history()
    key = _normalize_title(title)
    entry = history.get(key)

    if entry is None:
        history[key] = {"base_title": title, "genre": genre, "uses": [title], "count": 1}
        _save_title_history(history)
        return title

    count = entry["count"]
    if count < MAX_TITLE_PARTS:
        new_title = f"{entry['base_title']} Part {count + 1}"
        entry["count"] = count + 1
        entry["uses"].append(new_title)
        _save_title_history(history)
        print(f"[Config] Title '{title}' was already used — renamed to '{new_title}'.")
        return new_title

    disambiguated = f"{entry['base_title']} ({datetime.now().strftime('%Y-%m-%d')})"
    entry["uses"].append(disambiguated)
    _save_title_history(history)
    print(f"[WARN] Title '{title}' already has Parts 1-{MAX_TITLE_PARTS} — using '{disambiguated}' instead.")
    return disambiguated


# =====================================================================
# 1. LOCAL OLLAMA LLMs
# =====================================================================

# Two separate model configs, since the two jobs have very different needs:
#
# - OUTLINE_MODEL_NAME handles structured JSON generation (title/premise/chapter
#   list). Stock llama3.1 is reliable at this and rarely refuses — outline
#   requests don't read as "sensitive" the way chapter prose does. No reason to
#   change this one.
#
# - WRITER_MODEL_NAME writes the actual chapter prose. Stock llama3.1's instruct
#   tuning threw blunt, keyword-triggered refusals on completely mainstream
#   genre content (crime, noir, thriller, dark fantasy, rebellion-vs-empire
#   space opera) even with zero actual objectionable content in the request —
#   no amount of "this is fiction" prompt framing reliably fixes this, since
#   it's a base-model behavior, not a misunderstanding CrewAI can prompt its
#   way around. Switched (2026-08-11) to a storytelling-tuned uncensored model:
#       ollama pull Tohur/natsumura-storytelling-rp-llama-3.1
#   Fine-tuned specifically on storytelling/roleplay datasets, so it should
#   both refuse far less AND read more like genre-fiction prose than a general
#   chat model with the safety training stripped out. Requires pulling it once
#   on this machine before a run will find it — Ollama will error clearly if
#   it's missing. Second, separate one-time step (2026-08-12): the pipeline
#   doesn't target this base pull directly — see the WRITER_MODEL_NAME note
#   just below for why and what else needs to be run once.
OUTLINE_MODEL_NAME = "llama3.1"
# Points at a custom Ollama tag (not the base pull) with num_ctx raised to
# 16384 — see Modelfile.writer-16k and project plan section 5. Root cause:
# Ollama's 2048-token default context window was silently truncating long
# chapters mid-generation (confirmed 2026-08-12 against "Murder at Willowbrook
# Manor" — invented character names late in Chapter 1, prose dissolving into
# fragments at the Chapter 7 climax). Requires a one-time
# `ollama create natsumura-storytelling-rp-llama-3.1-16k -f Modelfile.writer-16k`
# on this machine before a run will find it — same as any other required
# model, the preflight check below will say so clearly if it's missing.
#
# --writer-model overrides this for a single run (see the CLI section above) —
# dashboard section 6.12's model-switch buttons pass it to A/B test an
# alternate writer model without editing this file.
WRITER_MODEL_NAME = WRITER_MODEL_OVERRIDE or "natsumura-storytelling-rp-llama-3.1-16k"

outline_llm = LLM(
    model=f"ollama/{OUTLINE_MODEL_NAME}",
    base_url="http://localhost:11434",
)

# =====================================================================
# WRITER OUTPUT LIMITS — added 2026-08-13 after "Rashomon's Rainy Nights"
# =====================================================================
# On that run every chapter collapsed into repetition in its second half.
# Chapter 2 was given a 2,800-word target and produced 8,303 words, the last
# ~5,500 of which were a loop where the single word "hell" made up 12.4% of
# all text. Measured type-token ratio (unique words / total words) fell to
# 0.055 in that stretch; healthy prose from earlier books measures 0.37-0.51.
#
# The mechanism is an overshoot that turns into a spiral: the model delivers
# the chapter beat, has nothing left to say, and keeps generating anyway
# because nothing tells it to stop. Two things were missing:
#
#   1. No num_predict — no ceiling on output length at all, so "approximately
#      2,800 words" in the prompt was a suggestion the model could ignore by
#      a factor of three.
#   2. No repeat_penalty — nothing discouraging it from cycling the same
#      phrases once it ran out of story.
#
# Worth recording the irony: raising num_ctx from 2048 to 16384 (project plan
# section 5) is what made this failure mode reachable. At 2048 tokens the
# model was forcibly cut off long before it could spiral, so the "coherence
# fix" removed the accidental brake. Section 5.5 was never confirmed on a real
# run; this was that confirmation, and it failed.
#
# CHAPTER_WORD_CAP_MULTIPLIER caps output at 1.3x the outline's target for
# that chapter, leaving room for natural variation while making an 8,000-word
# answer to a 2,800-word ask impossible. ~1.4 tokens per English word is the
# usual rule of thumb for Llama-family tokenizers.
CHAPTER_WORD_CAP_MULTIPLIER = 1.3
TOKENS_PER_WORD = 1.4
WRITER_REPEAT_PENALTY = 1.15  # >1 penalises reused tokens; 1.1-1.2 is the usual band

# Belt-and-suspenders, deliberately: the two parameters above are passed
# through CrewAI -> litellm -> Ollama, and that chain is worth not trusting
# blindly for a bug this expensive. detect_degeneration() below is a pure
# Python check on the returned text, so it catches a spiral even if the
# parameters are silently dropped somewhere in that stack.


writer_llm = LLM(
    model=f"ollama/{WRITER_MODEL_NAME}",
    base_url="http://localhost:11434",
)

# =====================================================================
# DIRECT OLLAMA CHAPTER WRITER — replaced the CrewAI path 2026-08-14
# =====================================================================
# The first attempt at capping output went through CrewAI's LLM(max_tokens=...),
# on the assumption litellm would map that to Ollama's num_predict. The
# 2026-08-13 17:11 run proved it does not: Chapter 2 was given a 2,730-word cap
# (~3,821 tokens) and came back with 8,759 words. The cap was silently dropped
# somewhere in the CrewAI -> litellm -> Ollama chain, and only the Python-side
# trim afterwards kept the chapter to size — after paying the full 779 seconds
# to generate text that was then thrown away.
#
# So chapter writing now talks to Ollama's /api/chat directly, where the
# options dict is unambiguous and verifiable. Three things this buys:
#
#   1. num_predict is actually enforced, so an over-long generation is
#      impossible rather than merely corrected after the fact.
#   2. Streaming lets the collapse detector run WHILE text arrives, so a
#      chapter that starts looping is abandoned within seconds instead of
#      after minutes. On the 17:11 run, Chapter 5 degenerated on three
#      consecutive full-length attempts — roughly 25 minutes spent generating
#      prose that was discarded, which is what looked like a hang.
#   3. It removes CrewAI's verbose panels from the writer stage, which were
#      ~78,000 of the 79,000 lines that killed the browser tab.
#
# Chapter writing is a single prompt in, one block of prose out — CrewAI's
# agent/task/crew machinery was adding overhead and opacity without adding
# anything here. The editor and scorer still use CrewAI; their JSON-schema
# retry behaviour genuinely benefits from it.
OLLAMA_URL = "http://localhost:11434"
WRITER_NUM_CTX = 16384          # matches Modelfile.writer-16k
WRITER_TEMPERATURE = 0.85
WRITER_STREAM_TIMEOUT = 1800    # seconds for the whole streamed response

# Rebuilt from the writer Agent's persona so switching off CrewAI doesn't
# change the prose voice — same role, goal, and backstory text, just sent as a
# system message instead of via an Agent object.
WRITER_SYSTEM_PROMPT = (
    "You are a Creative Fiction Writer. Your goal: draft vivid, atmospheric, "
    "continuity-consistent chapters based on an approved outline.\n\n"
    "You are a professional novelist writing mature, sophisticated genre "
    "fiction for adult readers, in the tradition of classic detective, "
    "thriller, and speculative fiction. You write immersive prose with "
    "strong sensory detail, natural dialogue, and forward-moving tension. "
    "You strictly follow the outline beat you're given for each chapter, "
    "stay consistent with what happened in prior chapters, and respect "
    "the target word count."
)


def generate_chapter_streaming(prompt: str, num_predict: int, ch_num: int) -> dict:
    """Stream one chapter from Ollama, aborting early if it starts to loop.

    Returns {"text": str, "aborted_early": bool, "reason": str}. Never raises
    for a degenerate response — collapse is an expected outcome the caller
    handles — but does raise requests exceptions for genuine transport
    failures, which the caller treats as a failed attempt.

    The early-abort check runs on a trailing slice rather than the whole text
    so far: collapse is a local property of where generation currently is, and
    scanning the full accumulated text every time would both get slower as the
    chapter grows and let a long healthy opening mask a spiralling ending."""
    payload = {
        "model": WRITER_MODEL_NAME,
        "messages": [
            {"role": "system", "content": WRITER_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        "stream": True,
        "options": {
            "num_predict": num_predict,
            "repeat_penalty": WRITER_REPEAT_PENALTY,
            "num_ctx": WRITER_NUM_CTX,
            "temperature": WRITER_TEMPERATURE,
        },
    }

    chunks = []
    words_seen = 0
    next_check_at = DEGEN_STREAM_FIRST_CHECK_WORDS
    next_report_at = 1000
    response = requests.post(f"{OLLAMA_URL}/api/chat", json=payload,
                             stream=True, timeout=WRITER_STREAM_TIMEOUT)
    response.raise_for_status()
    try:
        for raw_line in response.iter_lines():
            if not raw_line:
                continue
            try:
                event = json.loads(raw_line)
            except ValueError:
                continue
            piece = (event.get("message") or {}).get("content", "")
            if piece:
                chunks.append(piece)
            if event.get("done"):
                break

            text_so_far = "".join(chunks)
            words_seen = text_so_far.count(" ")  # cheap approximation, exact count is not needed here
            if words_seen >= next_report_at:
                print(f"[Writer]   ...chapter {ch_num}: ~{words_seen} words so far", flush=True)
                next_report_at += 1000
            if words_seen >= next_check_at:
                next_check_at += DEGEN_STREAM_CHECK_EVERY_WORDS
                tail = " ".join(text_so_far.split()[-DEGEN_STREAM_TAIL_WORDS:])
                verdict = detect_degeneration(tail)
                if verdict["degenerate"]:
                    # Stop paying for prose that is already being discarded.
                    response.close()
                    return {"text": text_so_far, "aborted_early": True,
                            "reason": verdict["detail"]}
    finally:
        response.close()

    return {"text": "".join(chunks), "aborted_early": False, "reason": ""}


# =====================================================================
# DEGENERATION DETECTION — added 2026-08-13
# =====================================================================
# Catches the repetition spiral described above by measuring the text itself
# rather than trusting generation parameters. Duplicated verbatim in
# scoring_agent.py (same "each script runs standalone" convention as the rest
# of this codebase) so the scorer can refuse to publish-rate a broken book
# even if it was written before this fix existed.
#
# Two signals, measured over a sliding 400-word window:
#
#   type-token ratio  — unique words / total words in the window. Collapses
#                       when the model starts cycling the same vocabulary.
#   runaway content word — the most frequent NON-function word's share of the
#                       window. Function words ("the", "a", "and") are
#                       legitimately 4-6% of any healthy English text, so
#                       they're excluded; a content word above 5% means the
#                       model is stuck on it.
#
# Thresholds were calibrated against 19 real chapters from this pipeline's own
# output_books/ — 5 known-bad (Rashomon) and 14 known-good (Murder at
# Willowbrook, Golden Heirloom, Aurora Initiative, Raining Ashes, Stardust
# Rebellion, Echoes of Eternity). Worst-window TTR measured 0.245-0.34 on the
# bad chapters and 0.365-0.505 on the good ones, so the 0.36 floor sits in a
# clean gap: it caught all 5 bad chapters and flagged not one window in any of
# the 14 good ones. Requiring two CONSECUTIVE bad windows keeps a single
# dialogue-heavy passage (naturally low TTR — short lines, repeated names)
# from tripping it.
DEGEN_WINDOW_WORDS = 400
DEGEN_STEP_WORDS = 200
DEGEN_TTR_FLOOR = 0.36
DEGEN_WORD_CEILING = 0.05
DEGEN_MIN_CONSECUTIVE = 2

# Streaming early-abort tuning (see generate_chapter_streaming). The tail slice
# is 1000 words so it spans enough windows for the 2-consecutive rule to be
# meaningful; checks start at 1200 words because a collapse before then is
# vanishingly rare and the first few hundred words of any chapter are naturally
# repetitive (establishing names, place, weather).
DEGEN_STREAM_TAIL_WORDS = 1000
DEGEN_STREAM_FIRST_CHECK_WORDS = 1200
DEGEN_STREAM_CHECK_EVERY_WORDS = 300

_DEGEN_WORD_RE = re.compile(r"[A-Za-z']+")
_DEGEN_STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "if", "of", "to", "in", "on", "at", "for", "with",
    "from", "by", "as", "is", "was", "were", "be", "been", "are", "am", "he", "she", "it",
    "they", "them", "his", "her", "him", "their", "its", "i", "you", "we", "me", "my", "that",
    "this", "these", "those", "not", "no", "had", "has", "have", "did", "do", "does", "would",
    "could", "should", "will", "can", "there", "then", "so", "up", "out", "down", "into",
    "over", "about", "just", "like", "what", "when", "who", "which", "all", "one", "more",
    "than", "too", "very", "said", "being", "own", "only", "other", "new",
}


def detect_degeneration(text: str) -> dict:
    """Return a verdict on whether `text` collapses into repetition.

    Keys: degenerate (bool), first_bad_word_index (int|None — where the
    collapse starts, for salvaging the healthy prefix), worst_ttr,
    runaway_word, total_words, longest_run, detail (human-readable)."""
    from collections import Counter as _Counter

    words = _DEGEN_WORD_RE.findall(text.lower())
    total = len(words)
    if total < DEGEN_WINDOW_WORDS:
        # Too short to measure meaningfully — don't guess. A chapter this short
        # has other problems the length check will catch anyway.
        return {"degenerate": False, "first_bad_word_index": None, "worst_ttr": None,
                "runaway_word": None, "total_words": total, "longest_run": 0,
                "detail": f"only {total} words — too short to assess"}

    # Scan on SEVERAL window phases and keep the worst verdict, rather than one
    # grid starting at word 0.
    #
    # Why (found 2026-08-14, "Rainy Night, Lonely Soul"): a single grid makes the
    # result phase-dependent, and on a chapter sitting right at the boundary that
    # decides the verdict. Chapter 9 of that book tested CLEAN as the writer saw
    # it (the chapter body, longest_run 1) and DEGENERATE as the scorer saw it
    # (the same body with a 6-word title line prepended, longest_run 2). Shifting
    # the start by 4 words flipped it. The chapter was in fact badly collapsed —
    # 20 repetitions of "murder" in one stretch — so the writer's miss let it
    # through to the editor, and only the scorer caught it.
    #
    # Four phases at quarter-step intervals means a bad stretch is measured
    # wherever it happens to begin, so the same text gets the same answer no
    # matter what precedes it.
    phases = [0, DEGEN_STEP_WORDS // 4, DEGEN_STEP_WORDS // 2, 3 * DEGEN_STEP_WORDS // 4]
    worst_ttr = 1.0
    runaway = None
    longest_run = 0
    first_bad = None
    for phase in phases:
        bad_offsets = []
        for i in range(phase, total - DEGEN_WINDOW_WORDS + 1, DEGEN_STEP_WORDS):
            window = words[i:i + DEGEN_WINDOW_WORDS]
            ttr = len(set(window)) / len(window)
            worst_ttr = min(worst_ttr, ttr)
            content = [w for w in window if w not in _DEGEN_STOPWORDS]
            share, top = 0.0, ""
            if content:
                top, count = _Counter(content).most_common(1)[0]
                share = count / len(window)
            if ttr < DEGEN_TTR_FLOOR or share > DEGEN_WORD_CEILING:
                bad_offsets.append(i)
                if share > DEGEN_WORD_CEILING and (runaway is None or share > runaway[1]):
                    runaway = (top, round(share, 4))
        run, prev = 0, None
        for i in bad_offsets:
            if prev is not None and i - prev == DEGEN_STEP_WORDS:
                run += 1
            else:
                run = 1
            longest_run = max(longest_run, run)
            prev = i
        if bad_offsets and (first_bad is None or bad_offsets[0] < first_bad):
            first_bad = bad_offsets[0]

    degenerate = longest_run >= DEGEN_MIN_CONSECUTIVE

    # first_bad_word_index is the FIRST bad window anywhere in the text, not the
    # start of the longest bad run — and that distinction matters for salvage.
    # Truncating at the longest run's start looked right but left earlier bad
    # patches in place: on Rashomon chapter 2 it produced a 2,596-word "salvaged"
    # prefix that was still degenerate, because that chapter was already
    # collapsing well before its worst stretch. Cutting at the first bad window
    # instead makes the surviving prefix clean by construction — every window
    # entirely inside it was measured and passed.
    detail = ""
    if degenerate:
        detail = (f"repetition collapse from around word {first_bad} of {total} "
                  f"(lowest vocabulary richness {round(worst_ttr, 3)}, floor {DEGEN_TTR_FLOOR})")
        if runaway:
            detail += f"; the word '{runaway[0]}' reaches {runaway[1] * 100:.1f}% of a 400-word stretch"
    return {"degenerate": degenerate, "first_bad_word_index": first_bad,
            "worst_ttr": round(worst_ttr, 3), "runaway_word": runaway, "total_words": total,
            "longest_run": longest_run, "detail": detail}


def truncate_at_word_index(text: str, word_index: int) -> str:
    """Cut `text` just before its `word_index`-th word, then back up to the last
    paragraph break so the salvaged prose ends on a clean break rather than
    mid-sentence. Used to rescue the healthy opening of a chapter whose ending
    spiralled, instead of throwing the whole chapter away."""
    matches = list(_DEGEN_WORD_RE.finditer(text))
    if word_index <= 0 or word_index >= len(matches):
        return text
    cut = matches[word_index].start()
    para_break = text.rfind("\n\n", 0, cut)
    if para_break > cut * 0.5:  # only if it doesn't cost us half the salvaged text
        cut = para_break
    return text[:cut].rstrip()


def check_ollama_models_available(
    required_models: list, ollama_url: str = "http://localhost:11434", fix_hints: dict = None
) -> None:
    """Preflight check, run once at startup before any CrewAI work begins.

    Without this, a missing/not-yet-pulled Ollama model surfaces as a wall of
    CrewAI retry-panel noise (3+ repeated 'LLM Call Failed' boxes) ending in a
    raw Python traceback — confirmed the hard way on 2026-08-12 when the run
    started Ollama manually mid-launch-sequence and it turned out not to have
    llama3.1 pulled. The 404 from Ollama itself is perfectly clear ("model
    'llama3.1' not found"); it's just buried three screens deep by the time
    CrewAI is done retrying. This checks Ollama's own /api/tags endpoint up
    front and fails fast with one unambiguous message and the exact fix
    command, before a single token of the (guaranteed-to-fail) crew run.

    fix_hints: optional {model_name: fix_command} overrides for models that
    aren't fixable with a plain `ollama pull` — added 2026-08-12 alongside the
    custom -16k context-window tags (see Modelfile.writer-16k /
    Modelfile.llama31-16k). Those are local-only tags built with `ollama
    create`, not registry models, so the default "ollama pull <name>" advice
    would send you chasing a model that doesn't exist on any registry.
    """
    import urllib.request as _urllib_request
    import json as _json

    fix_hints = fix_hints or {}

    try:
        with _urllib_request.urlopen(f"{ollama_url.rstrip('/')}/api/tags", timeout=5) as resp:
            data = _json.loads(resp.read())
    except Exception as e:
        print(f"[FATAL] Could not reach Ollama at {ollama_url} to verify required models: {e}")
        print("        Is Ollama running? Start it (e.g. 'ollama serve') and try again.")
        sys.exit(1)

    available = {m.get("name", "") for m in data.get("models", [])}
    # Ollama lists pulled models with an explicit tag ("llama3.1:latest"); a bare
    # required name like "llama3.1" should still match against that tag.
    available_bare = {name.split(":")[0] for name in available if name}

    missing = [
        m for m in required_models
        if m not in available and m.split(":")[0] not in available_bare
    ]
    if missing:
        print(f"[FATAL] Required Ollama model(s) not found on this machine: {', '.join(missing)}")
        for m in missing:
            print(f"        Fix: {fix_hints.get(m, f'ollama pull {m}')}")
        print(f"        Currently pulled models: {', '.join(sorted(available)) or '(none)'}")
        sys.exit(1)

    print(f"[Config] Ollama models verified present: {', '.join(required_models)}")


# Keyed by model name rather than built from whichever WRITER_MODEL_NAME happens
# to be active this run — --writer-model can swap in either custom -16k tag, and
# each needs its own correct 'ollama create ... -f Modelfile...' fix line, not
# whichever one happened to be the default when this dict was written.
WRITER_MODEL_FIX_HINTS = {
    "natsumura-storytelling-rp-llama-3.1-16k": (
        "one-time local setup, not a registry pull: "
        "ollama create natsumura-storytelling-rp-llama-3.1-16k -f Modelfile.writer-16k "
        "(requires 'ollama pull Tohur/natsumura-storytelling-rp-llama-3.1' first if you haven't already)"
    ),
    "llama3.1-16k": (
        "one-time local setup, not a registry pull: "
        "ollama create llama3.1-16k -f Modelfile.llama31-16k "
        "(requires 'ollama pull llama3.1' first if you haven't already)"
    ),
}

check_ollama_models_available(
    [OUTLINE_MODEL_NAME, WRITER_MODEL_NAME],
    fix_hints=WRITER_MODEL_FIX_HINTS,
)

# =====================================================================
# 2. CUSTOM AUTOMATIC1111 IMAGE GENERATION TOOL — SDXL resolution, hi-res
#    fix, cropped to ebook cover ratio, with Pillow text overlay
# =====================================================================


class CoverToolInput(BaseModel):
    """Input schema for the Local Cover Generator Tool."""
    prompt: str = Field(
        ...,
        description=(
            "Detailed, descriptive prompt for the cover ART ONLY. Do not include any "
            "title, lettering, text, or typography in this prompt — text is added "
            "separately after generation."
        ),
    )
    negative_prompt: str = Field(
        default="text, watermark, signature, letters, words, blurry, low quality, "
                "deformed, extra limbs, bad proportions",
        description="Things to avoid in the generated image."
    )


# =====================================================================
# 2a. COVER TYPOGRAPHY — genre-specific fonts, glow, "duality" shadow,
#     distressed grain, and a diagonal light-leak color grade.
# =====================================================================
# Fonts live in fonts/ next to this script (all open-license/OFL, safe for
# commercial genre covers). Each genre gets its own title/author typeface
# pairing plus a small set of "flare" parameters (accent glow color, light
# leak tint, shadow drama, grain, and whether one title letter gets subtly
# subverted into a mismatched style) instead of one generic plain-white look.
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")


def _font_path(filename: str) -> str:
    return os.path.join(FONT_DIR, filename)


GENRE_COVER_STYLE = {
    "sci-fi thriller": dict(
        title_font=_font_path("BebasNeue-Regular.ttf"),
        author_font=_font_path("CrimsonText-Italic.ttf"),
        glow_color=(80, 200, 255), light_leak_color=(60, 160, 255),
        shadow_style="dramatic", grain=False, letter_subvert=True,
    ),
    "cozy mystery": dict(
        title_font=_font_path("PlayfairDisplay-VF.ttf"),
        author_font=_font_path("CrimsonText-Italic.ttf"),
        glow_color=None, light_leak_color=(255, 200, 130),
        shadow_style="soft", grain=False, letter_subvert=False,
    ),
    "epic fantasy": dict(
        title_font=_font_path("CinzelDecorative-Black.ttf"),
        author_font=_font_path("Cinzel-Bold.ttf"),
        glow_color=(255, 205, 90), light_leak_color=(255, 215, 120),
        shadow_style="soft", grain=False, letter_subvert=False,
    ),
    "romantic suspense": dict(
        title_font=_font_path("AbrilFatface-Regular.ttf"),
        author_font=_font_path("CrimsonText-Italic.ttf"),
        glow_color=(255, 60, 140), light_leak_color=(255, 90, 190),
        shadow_style="dramatic", grain=False, letter_subvert=True,
    ),
    "psychological horror": dict(
        title_font=_font_path("Nosifer-Regular.ttf"),
        author_font=_font_path("CrimsonText-Regular.ttf"),
        glow_color=(180, 20, 20), light_leak_color=(120, 10, 40),
        shadow_style="dramatic", grain=True, letter_subvert=True,
    ),
    "dystopian survival": dict(
        title_font=_font_path("BlackOpsOne-Regular.ttf"),
        author_font=_font_path("Oswald-VF.ttf"),
        glow_color=(255, 150, 40), light_leak_color=(200, 110, 20),
        shadow_style="dramatic", grain=True, letter_subvert=False,
    ),
    "noir detective": dict(
        title_font=_font_path("Anton-Regular.ttf"),
        author_font=_font_path("CrimsonText-Italic.ttf"),
        glow_color=(200, 30, 30), light_leak_color=(180, 40, 40),
        shadow_style="dramatic", grain=True, letter_subvert=True,
    ),
    "space opera": dict(
        title_font=_font_path("Audiowide-Regular.ttf"),
        author_font=_font_path("Oswald-VF.ttf"),
        glow_color=(160, 90, 255), light_leak_color=(120, 80, 255),
        shadow_style="soft", grain=False, letter_subvert=False,
    ),
}


def _load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    """Load a specific font file at `size`. Falls back through CUSTOM_TITLE_FONT_PATH,
    then common system fonts, then PIL's default bitmap font if the requested file is
    missing (e.g. the fonts/ folder didn't make it onto this machine), so a cover
    still renders instead of crashing Stage C."""
    candidates = [
        path,
        CUSTOM_TITLE_FONT_PATH,
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",  # common on Linux
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",      # common on macOS
        "C:\\Windows\\Fonts\\arialbd.ttf",                        # common on Windows
    ]
    for candidate in candidates:
        if not candidate:
            continue
        if not os.path.isfile(candidate):
            if candidate == path:
                print(f"[WARN] Font file not found at '{candidate}' — trying fallback fonts.")
            continue
        try:
            return ImageFont.truetype(candidate, size)
        except Exception as e:
            print(f"[WARN] Found '{candidate}' but couldn't load it ({e}) — trying fallback fonts.")
            continue
    print(f"[WARN] Could not load '{path}' or any fallback font — using PIL's default bitmap font.")
    return ImageFont.load_default()


def _fit_text_to_width(draw: ImageDraw.ImageDraw, text: str, max_width: int,
                        start_size: int, min_size: int, font_path: str) -> tuple:
    """Shrink font size until the text (single line) fits max_width, down to min_size."""
    size = start_size
    while size > min_size:
        font = _load_font(font_path, size)
        bbox = draw.textbbox((0, 0), text, font=font)
        if (bbox[2] - bbox[0]) <= max_width:
            return font, size
        size -= 4
    return _load_font(font_path, min_size), min_size


def _wrap_at_size(draw: ImageDraw.ImageDraw, words: list, max_width: int,
                   font: ImageFont.FreeTypeFont) -> list:
    lines, current = [], ""
    for word in words:
        trial = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), trial, font=font)
        if (bbox[2] - bbox[0]) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _wrap_title(draw: ImageDraw.ImageDraw, title: str, max_width: int,
                 start_size: int, min_size: int, font_path: str) -> tuple:
    """Shrink font size until the title wraps to <=3 lines AND every line actually
    fits max_width. Line count alone isn't enough for wide/decorative display fonts
    (e.g. the horror title font) — a single long word placed alone on its own line
    passes the count check while still overflowing the canvas, since a lone word's
    width is never checked against max_width during packing. Both are verified here."""
    font, lines = None, None
    for size in range(start_size, min_size - 1, -4):
        font = _load_font(font_path, size)
        lines = _wrap_at_size(draw, title.split(), max_width, font)
        fits_width = all(
            (draw.textbbox((0, 0), ln, font=font)[2] - draw.textbbox((0, 0), ln, font=font)[0]) <= max_width
            for ln in lines
        )
        if len(lines) <= 3 and fits_width:
            return font, lines
    return font, lines


def _vertical_fade_band(width: int, band_height: int, top_opaque: bool, max_alpha: int) -> Image.Image:
    """A black RGBA band whose alpha fades from max_alpha at one edge to 0 at the
    other, so title/author legibility backing blends into the art instead of sitting
    on top of it as a hard-edged box (the "atmospheric blending" ask)."""
    grad = Image.linear_gradient("L").resize((width, band_height))  # 0 top -> 255 bottom
    if top_opaque:
        grad = grad.transpose(Image.FLIP_TOP_BOTTOM)  # 255 top -> 0 bottom
    alpha = grad.point(lambda p: int(p * max_alpha / 255))
    band = Image.new("RGBA", (width, band_height), (0, 0, 0, 0))
    band.putalpha(alpha)
    return band


def _light_leak_overlay(size: tuple, color: tuple, opacity: float = 0.30, angle: int = 25) -> Image.Image:
    """A soft diagonal colored streak blended over the whole cover, as if a light
    source in the scene is spilling across it — colored per genre so it reads as
    lit-by-the-environment rather than a generic overlay."""
    width, height = size
    diag = int((width ** 2 + height ** 2) ** 0.5)
    grad = Image.linear_gradient("L").resize((diag, diag))
    grad = grad.rotate(angle, resample=Image.BICUBIC)
    left, top = (grad.width - width) // 2, (grad.height - height) // 2
    grad = grad.crop((left, top, left + width, top + height))
    grad = grad.filter(ImageFilter.GaussianBlur(radius=max(width * 0.05, 20)))
    alpha = grad.point(lambda p: int(p * opacity))
    overlay = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    overlay.paste(Image.new("RGB", (width, height), color), (0, 0))
    overlay.putalpha(alpha)
    return overlay


def _draw_char(canvas: Image.Image, ch: str, xy: tuple, font: ImageFont.FreeTypeFont,
               fill: tuple, rotation: float = 0.0) -> None:
    """Draw one character onto `canvas`, optionally rotated in place (used for the
    single subverted title letter — "something about this word is off")."""
    if rotation == 0:
        ImageDraw.Draw(canvas).text(xy, ch, font=font, fill=fill)
        return
    bbox = font.getbbox(ch)
    w, h = max(bbox[2] - bbox[0], 1), max(bbox[3] - bbox[1], 1)
    pad = int(max(w, h) * 0.6) + 4
    tile = Image.new("RGBA", (w + pad * 2, h + pad * 2), (0, 0, 0, 0))
    ImageDraw.Draw(tile).text((pad - bbox[0], pad - bbox[1]), ch, font=font, fill=fill)
    tile = tile.rotate(rotation, resample=Image.BICUBIC, expand=True)
    canvas.alpha_composite(tile, (int(xy[0] - (tile.width - w) / 2), int(xy[1] - (tile.height - h) / 2)))


def _draw_title_line(canvas: Image.Image, line: str, x: int, y: int, font: ImageFont.FreeTypeFont,
                      fill: tuple, accent_index=None, accent_font=None, accent_fill=None,
                      accent_rotation: float = 7) -> None:
    """Draw one title line character-by-character so a single letter can be swapped
    to a mismatched font/color/rotation without disturbing the rest of the line."""
    cursor_x = x
    draw = ImageDraw.Draw(canvas)
    for i, ch in enumerate(line):
        char_font, char_fill, rotation, y_off = font, fill, 0.0, 0
        if accent_index is not None and i == accent_index and ch != " ":
            char_font, char_fill, rotation = accent_font or font, accent_fill or fill, accent_rotation
            y_off = int(font.size * 0.05)
        _draw_char(canvas, ch, (cursor_x, y + y_off), char_font, char_fill, rotation)
        cursor_x += draw.textlength(ch, font=font)


def _apply_grain_to_layer(layer: Image.Image, intensity: int = 45, strength: float = 0.4) -> Image.Image:
    """Roughen an RGBA text layer's alpha with noise so the fill reads as distressed/
    under-duress rather than a perfectly flat color, without fully punching holes
    through the letterforms."""
    noise = Image.effect_noise(layer.size, intensity)
    noise_mask = noise.point(lambda p: min(255, int(p * 1.6)))
    r, g, b, a = layer.split()
    eroded = ImageChops.multiply(a, noise_mask)
    layer.putalpha(Image.blend(a, eroded, strength))
    return layer


def compose_cover(art_image: Image.Image, title: str, author: str, genre: str, subtitle: str = "") -> Image.Image:
    """Fit the generated art to the ebook cover ratio and overlay genre-styled title/
    subtitle/author typography: a matched font pairing, a diagonal light-leak color
    grade, a gradient legibility band that blends into the art instead of a hard
    block, an offset "duality" shadow, an optional neon glow, optional distressed
    grain, and (for genres where it fits) one subtly subverted title letter.

    `subtitle` is an optional story-specific tagline rendered just beneath the title
    in a smaller weight, and `author` is rendered as a two-tier "Written by" / pen
    name block — both as default rules applied across every genre."""
    style = GENRE_COVER_STYLE.get(genre, GENRE_COVER_STYLE["sci-fi thriller"])
    cover = ImageOps.fit(art_image.convert("RGB"), TARGET_COVER_SIZE, method=Image.LANCZOS).convert("RGBA")
    width, height = cover.size
    margin = int(width * 0.08)
    max_text_width = width - (2 * margin)

    if style["light_leak_color"]:
        cover = Image.alpha_composite(cover, _light_leak_overlay((width, height), style["light_leak_color"]))

    overlay = Image.new("RGBA", cover.size, (0, 0, 0, 0))

    # Legibility bands — gradients now, not solid rectangles, so fog/lights show through.
    title_band_height = int(height * 0.36)
    overlay.alpha_composite(_vertical_fade_band(width, title_band_height, top_opaque=True, max_alpha=190), (0, 0))
    author_band_height = int(height * 0.16)
    author_band_top = height - author_band_height
    overlay.alpha_composite(
        _vertical_fade_band(width, author_band_height, top_opaque=False, max_alpha=190), (0, author_band_top)
    )

    # --- Title ---
    title_font, title_lines = _wrap_title(
        ImageDraw.Draw(overlay), title.upper(), max_text_width,
        start_size=int(width * 0.13), min_size=int(width * 0.05), font_path=style["title_font"],
    )
    line_height = title_font.size + int(title_font.size * 0.28)
    total_title_height = line_height * len(title_lines)

    # --- Subtitle sizing (computed before title placement so the title+subtitle
    # block can be centered together, instead of the subtitle getting whatever
    # leftover space happens to remain below wherever the title landed) ---
    subtitle = (subtitle or "").strip()
    subtitle_font, subtitle_lines, sub_line_height, total_subtitle_height = None, [], 0, 0
    if subtitle:
        subtitle_font, subtitle_lines = _wrap_title(
            ImageDraw.Draw(overlay), subtitle, max_text_width,
            start_size=int(width * 0.042), min_size=int(width * 0.024), font_path=style["author_font"],
        )
        subtitle_lines = subtitle_lines[:2]  # a tagline, not a second title
        sub_line_height = subtitle_font.size + int(subtitle_font.size * 0.35)
        total_subtitle_height = sub_line_height * len(subtitle_lines)
    subtitle_gap = int(title_font.size * 0.22) if subtitle_lines else 0

    block_height = total_title_height + subtitle_gap + total_subtitle_height
    y = max(int(height * 0.05), int(title_band_height * 0.62) - block_height // 2)

    accent_font = _load_font(style["author_font"], int(title_font.size * 0.95)) if style["letter_subvert"] else None
    accent_line_idx = random.randrange(len(title_lines)) if style["letter_subvert"] else None

    title_layer = Image.new("RGBA", cover.size, (0, 0, 0, 0))
    for li, line in enumerate(title_lines):
        bbox = ImageDraw.Draw(title_layer).textbbox((0, 0), line, font=title_font)
        x = (width - (bbox[2] - bbox[0])) // 2
        accent_idx = None
        if style["letter_subvert"] and li == accent_line_idx:
            letter_positions = [i for i, c in enumerate(line) if c != " "]
            if letter_positions:
                pool = letter_positions[1:-1] or letter_positions  # avoid first/last letter of the line
                accent_idx = random.choice(pool)
        _draw_title_line(
            title_layer, line, x, y, title_font, (255, 255, 255, 255),
            accent_index=accent_idx, accent_font=accent_font,
            accent_fill=style["glow_color"] or (255, 255, 255, 255),
        )
        y += line_height

    if style["grain"]:
        title_layer = _apply_grain_to_layer(title_layer)

    # "Duality" shadow: a darker, offset + blurred copy underneath the crisp text —
    # not just a drop shadow, but a distorted double suggesting a second identity.
    dx, dy = (10, 10) if style["shadow_style"] == "dramatic" else (5, 6)
    blur_radius = 3 if style["shadow_style"] == "dramatic" else 1.5
    shadow_layer = Image.new("RGBA", cover.size, (0, 0, 0, 0))
    shadow_layer.paste(Image.new("RGBA", cover.size, (0, 0, 0, 190)), (0, 0), title_layer.split()[3])
    shadow_layer = shadow_layer.filter(ImageFilter.GaussianBlur(blur_radius))
    shifted_shadow = Image.new("RGBA", cover.size, (0, 0, 0, 0))
    shifted_shadow.paste(shadow_layer, (dx, dy), shadow_layer)
    overlay.alpha_composite(shifted_shadow)

    # Neon glow behind the text, colored per genre (skipped for genres — e.g. cozy
    # mystery — where a lit-sign look would fight the tone).
    if style["glow_color"]:
        glow_layer = Image.new("RGBA", cover.size, (0, 0, 0, 0))
        glow_layer.paste(Image.new("RGBA", cover.size, style["glow_color"] + (255,)), (0, 0), title_layer.split()[3])
        glow_layer = glow_layer.filter(ImageFilter.GaussianBlur(title_font.size * 0.12))
        overlay.alpha_composite(glow_layer)

    overlay.alpha_composite(title_layer)

    # --- Subtitle: story-specific tagline, sits directly under the title in a
    # lighter weight (the genre's author font) with a soft shadow only — no glow,
    # no letter subversion, so it reads as clearly secondary to the title. ---
    if subtitle_lines:
        sub_y = y + subtitle_gap
        subtitle_layer = Image.new("RGBA", cover.size, (0, 0, 0, 0))
        draw_sub = ImageDraw.Draw(subtitle_layer)
        for line in subtitle_lines:
            bbox = draw_sub.textbbox((0, 0), line, font=subtitle_font)
            x = (width - (bbox[2] - bbox[0])) // 2
            draw_sub.text((x, sub_y), line, font=subtitle_font, fill=(235, 235, 235, 255))
            sub_y += sub_line_height

        sub_shadow = Image.new("RGBA", cover.size, (0, 0, 0, 0))
        sub_shadow.paste(Image.new("RGBA", cover.size, (0, 0, 0, 160)), (0, 0), subtitle_layer.split()[3])
        sub_shadow = sub_shadow.filter(ImageFilter.GaussianBlur(2))
        shifted_sub_shadow = Image.new("RGBA", cover.size, (0, 0, 0, 0))
        shifted_sub_shadow.paste(sub_shadow, (3, 4), sub_shadow)
        overlay.alpha_composite(shifted_sub_shadow)
        overlay.alpha_composite(subtitle_layer)

    # --- Author: "Written by" (smaller) stacked above the pen name (larger) — a
    # default two-tier rule applied the same way across every genre. Both lines use
    # the genre's author font so the pairing still feels typographically matched. ---
    pen_font, _ = _fit_text_to_width(
        ImageDraw.Draw(overlay), author, max_text_width,
        start_size=int(width * 0.05), min_size=int(width * 0.03), font_path=style["author_font"],
    )
    byline_font = _load_font(style["author_font"], max(int(pen_font.size * 0.62), 18))
    byline_text = "Written by"

    draw = ImageDraw.Draw(overlay)
    byline_bbox = draw.textbbox((0, 0), byline_text, font=byline_font)
    pen_bbox = draw.textbbox((0, 0), author, font=pen_font)
    byline_h = byline_bbox[3] - byline_bbox[1]
    pen_h = pen_bbox[3] - pen_bbox[1]
    author_gap = int(pen_font.size * 0.16)

    block_h = byline_h + author_gap + pen_h
    y = author_band_top + (height - author_band_top - block_h) // 2

    x = (width - (byline_bbox[2] - byline_bbox[0])) // 2
    draw.text((x, y), byline_text, font=byline_font, fill=(220, 220, 220, 230))

    y += byline_h + author_gap
    x = (width - (pen_bbox[2] - pen_bbox[0])) // 2
    draw.text((x, y), author, font=pen_font, fill=(255, 255, 255, 255))

    return Image.alpha_composite(cover, overlay).convert("RGB")


# =====================================================================
# 2b. VRAM UTILITIES
# =====================================================================
# The RTX 4060 8GB in this setup has to host Ollama (Stage A/B) and
# AUTOMATIC1111/SDXL (Stage C) *sequentially*, not at the same time — there
# isn't enough VRAM for both to be resident at once. Ollama keeps a model
# loaded for ~5 minutes after its last call by default (`keep_alive`), so
# without an explicit unload, Stage C starts fighting Ollama for the same
# VRAM and A1111's hi-res fix pass (the most memory-hungry step) fails with
# "Not enough memory" even though the base generation would have fit fine.
#
# NOTE: Windows' per-app "Graphics performance preference" (the setting used
# to try routing things to the Intel iGPU) does NOT affect this — it only
# applies to DirectX/OpenGL apps. PyTorch/CUDA (what both Ollama and A1111
# use) always targets the NVIDIA GPU directly, regardless of that setting.
# Intel UHD graphics can't run CUDA at all, so there's no way to actually
# move this workload off the 4060; the only real lever is not making it
# share VRAM with Ollama at the same moment.


def get_free_vram_mb() -> "int | None":
    """Query free VRAM in MB via nvidia-smi. Returns None (not an exception) if
    nvidia-smi isn't available, so callers can treat 'unknown' as 'proceed
    cautiously' instead of crashing on machines without it on PATH."""
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=memory.free", "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode == 0 and result.stdout.strip():
            return int(result.stdout.strip().splitlines()[0])
    except (subprocess.SubprocessError, FileNotFoundError, ValueError, OSError):
        pass
    return None


def detect_sdxl_checkpoint() -> "bool | None":
    """Ask AUTOMATIC1111 which checkpoint is currently loaded so Stage C can pick a
    resolution/hi-res strategy that actually matches it, instead of hardcoding one
    model family. Returns None (unknown) if the request fails, so callers can fall
    back to conservative settings instead of guessing wrong and wasting VRAM."""
    try:
        response = requests.get("http://127.0.0.1:7860/sdapi/v1/options", timeout=15)
        if response.status_code == 200:
            checkpoint = str(response.json().get("sd_model_checkpoint", ""))
            return "xl" in checkpoint.lower()
    except requests.exceptions.RequestException:
        pass
    return None


def unload_ollama_model(model_name: str, max_wait_seconds: int = 20) -> None:
    """Tell Ollama to drop the model from VRAM immediately (keep_alive: 0) instead
    of waiting out its default ~5 minute idle timeout, then wait briefly for the
    VRAM to actually come back before Stage C starts fighting Ollama for it."""
    print(f"\n[VRAM] Unloading Ollama model '{model_name}' to free VRAM for cover generation...")
    before = get_free_vram_mb()
    try:
        requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model_name, "keep_alive": 0},
            timeout=15,
        )
    except requests.exceptions.RequestException as e:
        print(f"[WARN] Could not reach Ollama to unload the model ({e}). Continuing anyway.")
        return

    if before is None:
        # No nvidia-smi available to confirm — just give the unload a moment to land.
        time.sleep(3)
        return

    for _ in range(max_wait_seconds):
        time.sleep(1)
        now = get_free_vram_mb()
        if now is not None and now >= before + 500:
            print(f"[VRAM] Freed up to {now}MB available (was {before}MB).")
            return
    print(f"[WARN] VRAM didn't visibly increase after {max_wait_seconds}s "
          f"(still ~{before}MB free) — proceeding to Stage C anyway.")


def unload_a1111_checkpoint(a1111_url: str = "http://127.0.0.1:7860",
                            max_wait_seconds: int = 25) -> None:
    """The mirror image of unload_ollama_model(), added 2026-08-13.

    The original VRAM handoff was one-directional: Ollama got unloaded before
    AUTOMATIC1111 loaded SDXL (Stage B -> C), but nothing ever unloaded SDXL
    afterwards. A1111 holds its checkpoint resident indefinitely — it has no
    idle timeout the way Ollama's keep_alive does — so on an 8GB card SDXL
    (~6GB with its VAE) was still sitting in VRAM when editorial_agent.py
    started, leaving nowhere near enough room for llama3.1-16k (~4.7GB of
    weights plus ~2GB of 16k KV cache).

    Ollama's response to not fitting is not to fail loudly; it silently
    offloads most layers to system RAM and runs them on the CPU. That single
    fact explains all three symptoms reported 2026-08-13:

      * the "rogue" ollama.exe pinning ~50% CPU and ~50% RAM in Task Manager
        (that IS the model, running on the CPU because it couldn't fit),
      * the editorial stage taking 274s/chapter vs. the writer's 165s
        (job_timing_history.json) despite doing strictly less generation,
      * the hang on the 2nd editorial pass and the crash partway through the
        5-book batch — CPU-offloaded inference plus SDXL's own host memory
        eventually exhausts system RAM.

    Called at the end of Stage C so the GPU is clean for whatever runs next,
    and again defensively at the start of editorial_agent.py / scoring_agent.py
    in case those are run standalone. A1111 reloads the checkpoint by itself
    on the next txt2img request, so this costs nothing but a ~20s reload on
    the next book — which now happens while nothing else needs the VRAM."""
    print("\n[VRAM] Unloading AUTOMATIC1111's checkpoint to free VRAM for the next stage...")
    before = get_free_vram_mb()
    try:
        response = requests.post(f"{a1111_url.rstrip('/')}/sdapi/v1/unload-checkpoint", timeout=60)
        if response.status_code not in (200, 204):
            print(f"[WARN] A1111 unload-checkpoint returned HTTP {response.status_code}. "
                  f"Continuing anyway.")
    except requests.exceptions.RequestException as e:
        # Not running at all is the normal case when the editor/scorer are run
        # standalone — that's a no-op, not a problem worth alarming about.
        print(f"[VRAM] AUTOMATIC1111 not reachable ({e.__class__.__name__}) — nothing to unload.")
        return

    if before is None:
        time.sleep(3)
        return

    for _ in range(max_wait_seconds):
        time.sleep(1)
        now = get_free_vram_mb()
        if now is not None and now >= before + 500:
            print(f"[VRAM] Freed up to {now}MB available (was {before}MB).")
            return
    print(f"[WARN] VRAM didn't visibly increase after {max_wait_seconds}s "
          f"(still ~{before}MB free) — continuing anyway.")


def wait_for_free_vram(target_mb: int, max_wait_seconds: int = 60) -> bool:
    """Block until at least target_mb of VRAM is free, or the timeout expires.

    Used as the last gate before handing the GPU to the next stage or the next
    book. Driver-side frees are not instant after a process exits, which is why
    a batch run with no cooldown could start book N+1 while book N's memory was
    still being reclaimed. Returns True if the target was met, False on timeout
    (callers warn and proceed rather than aborting a run over a soft
    condition)."""
    free = get_free_vram_mb()
    if free is None:
        # No nvidia-smi to gate on — fall back to a fixed settle pause.
        time.sleep(5)
        return True
    if free >= target_mb:
        return True
    print(f"[VRAM] Waiting for {target_mb}MB free (currently {free}MB)...")
    for _ in range(max_wait_seconds):
        time.sleep(1)
        free = get_free_vram_mb()
        if free is not None and free >= target_mb:
            print(f"[VRAM] {free}MB free — proceeding.")
            return True
    print(f"[WARN] Only {free}MB free after {max_wait_seconds}s (wanted {target_mb}MB) — "
          f"proceeding anyway, but the next stage may be slow or fall back to CPU.")
    return False


class Auto1111CoverTool(BaseTool):
    name: str = "Local Book Cover Generator"
    description: str = (
        "Sends an image generation prompt to the local AUTOMATIC1111 WebUI API, "
        "hi-res upscales it, crops it to ebook cover proportions, overlays the book's "
        "title and author, and saves the finished cover locally."
    )
    args_schema: Type[BaseModel] = CoverToolInput

    # These are set at construction time in Python (Stage C), never by the LLM.
    output_dir: str = "."
    book_title: str = "Untitled"
    author_name: str = AUTHOR_NAME
    genre: str = RUN_GENRE
    subtitle: str = ""

    def _generate_single(self, prompt, negative_prompt, gen_width, gen_height, hr_scale, seed):
        """Run the hi-res -> no-hi-res -> reduced-resolution fallback ladder for one
        seed. Returns (image_bytes, label) on success, or (None, error_message)."""
        url = "http://127.0.0.1:7860/sdapi/v1/txt2img"
        base_payload = {
            "prompt": prompt,
            "negative_prompt": negative_prompt,
            "steps": 30,
            "cfg_scale": 7.0,
            "sampler_name": "DPM++ 2M Karras",
            "seed": seed,
        }
        hr_extra = {
            "enable_hr": True,
            "hr_scale": hr_scale,
            "hr_upscaler": "R-ESRGAN 4x+",
            "denoising_strength": 0.4,
        }

        # Fallback ladder: hi-res fix -> plain base resolution -> reduced resolution.
        # Each tier needs meaningfully less VRAM than the last, so an 8GB card that
        # can't fit the hi-res second pass (the most memory-hungry step, since it
        # runs on top of whatever the base pass + upscaler already allocated) still
        # gets *a* cover instead of that option failing outright.
        attempts = [
            ("hi-res fix", {**base_payload, "width": gen_width, "height": gen_height, **hr_extra}),
            ("no hi-res fix", {**base_payload, "width": gen_width, "height": gen_height}),
            ("reduced resolution, no hi-res fix", {**base_payload, "width": 640, "height": 960}),
        ]

        last_error = "unknown error"
        for label, attempt_payload in attempts:
            try:
                print(f"\n[Tool] Triggering AUTOMATIC1111 API (seed {seed}, {label})...")
                response = requests.post(url, json=attempt_payload, timeout=300)

                if response.status_code != 200:
                    last_error = f"HTTP {response.status_code} - {response.text[:300]}"
                    reason = "ran out of VRAM" if "memory" in response.text.lower() else "failed"
                    print(f"[WARN] {label} {reason} ({response.status_code}), trying a lighter setting...")
                    continue

                data = response.json()
                image_bytes = base64.b64decode(data["images"][0])
                return image_bytes, label

            except requests.exceptions.ConnectionError:
                return None, (
                    "Could not connect to AUTOMATIC1111 at http://127.0.0.1:7860. "
                    "Ensure WebUI is running with the '--api' flag enabled."
                )
            except Exception as e:
                last_error = str(e)
                print(f"[WARN] {label} raised an error ({last_error}), trying a lighter setting...")
                continue

        return None, last_error

    def _run(self, prompt: str, negative_prompt: str = "", num_options: int = 3) -> str:
        free_vram = get_free_vram_mb()
        if free_vram is not None:
            print(f"[Tool] Free VRAM before cover generation: {free_vram}MB")

        # Pick resolution/hi-res strategy based on whichever checkpoint is actually
        # loaded right now, rather than assuming SDXL. SDXL's native buckets (e.g.
        # 832x1216) are already high-detail, so it only needs a light hi-res assist;
        # SD1.5 was trained at ~512x768 and both wastes VRAM and loses quality when
        # forced to SDXL-sized buckets, so it gets its own native-resolution ladder
        # and a slightly stronger hi-res pass to compensate for the smaller base.
        is_sdxl = detect_sdxl_checkpoint()
        if is_sdxl:
            gen_width, gen_height, hr_scale = 832, 1216, 1.2
            print("[Tool] SDXL checkpoint detected — using SDXL-native resolution, light hi-res pass.")
        elif is_sdxl is False:
            gen_width, gen_height, hr_scale = 512, 768, 1.4
            print("[Tool] SD1.5-style checkpoint detected — using its native resolution to save VRAM/quality.")
        else:
            gen_width, gen_height, hr_scale = BASE_GEN_WIDTH, BASE_GEN_HEIGHT, 1.5
            print("[Tool] Could not detect checkpoint type (is A1111 running?) — using default resolution.")

        # Generate `num_options` independent candidates (different seeds each) so
        # there's a real choice to pick from instead of one auto-selected cover.
        saved_options = []
        errors = []
        for i in range(1, num_options + 1):
            seed = random.randint(0, 2**31 - 1)
            image_bytes, result = self._generate_single(
                prompt, negative_prompt, gen_width, gen_height, hr_scale, seed
            )
            if image_bytes is None:
                print(f"[WARN] Cover option {i}/{num_options} failed: {result}")
                errors.append(f"option {i}: {result}")
                continue

            raw_path = os.path.join(self.output_dir, f"_raw_art_option_{i}.png")
            with open(raw_path, "wb") as f:
                f.write(image_bytes)

            from io import BytesIO
            art_image = Image.open(BytesIO(image_bytes))
            final_cover = compose_cover(art_image, self.book_title, self.author_name, self.genre, self.subtitle)
            cover_path = os.path.join(self.output_dir, f"cover_option_{i}.png")
            final_cover.save(cover_path)
            # Storefront cover-upload steps commonly accept JPG or TIFF but not PNG
            # (checked 2026-08-12) — PNG stays the working format everywhere else
            # (compositing here, previews, edited/) since it's lossless, but a JPEG copy is
            # saved alongside every option specifically for that upload step. final_cover
            # is already RGB (compose_cover's return value), so no alpha-channel handling needed.
            cover_path_jpg = os.path.join(self.output_dir, f"cover_option_{i}.jpg")
            final_cover.save(cover_path_jpg, "JPEG", quality=95)
            print(f"[Tool] Cover option {i}/{num_options} saved to '{cover_path}' (+ '{cover_path_jpg}' for upload) (used: {result}, seed {seed}).")
            saved_options.append((i, cover_path, cover_path_jpg, raw_path))

        if not saved_options:
            return f"Error: all {num_options} cover option attempts failed. Details: {'; '.join(errors)}"

        # Stage D and the rest of the pipeline expect a single 'cover.png' / '_raw_art.png' —
        # default to the first option that actually succeeded until the dashboard picker
        # UI exists, without losing the other generated options.
        default_i, default_cover, default_cover_jpg, default_raw = saved_options[0]
        shutil.copyfile(default_cover, os.path.join(self.output_dir, "cover.png"))
        shutil.copyfile(default_cover_jpg, os.path.join(self.output_dir, "cover.jpg"))
        shutil.copyfile(default_raw, os.path.join(self.output_dir, "_raw_art.png"))

        option_list = ", ".join(f"option {i}" for i, _, _, _ in saved_options)
        return (
            f"Success! Generated {len(saved_options)}/{num_options} cover options ({option_list}) "
            f"in '{self.output_dir}'. cover.png/cover.jpg currently default to option {default_i} — open "
            f"cover_option_1.png / cover_option_2.png / cover_option_3.png to compare (each has a matching "
            f".jpg for upload) and manually replace cover.png + cover.jpg if you prefer a different "
            f"one (a dashboard picker is planned)."
        )


# =====================================================================
# 3. AGENTS
# =====================================================================

concept_architect = Agent(
    role="Book Concept & Outline Architect",
    goal=(
        "Design a complete, structured story outline — title, premise, and a "
        "chapter-by-chapter breakdown — that fits within a strict chapter and "
        "word budget."
    ),
    backstory=(
        "You are a developmental editor who specializes in tightly-plotted "
        "commercial genre fiction. You think in structure first: "
        "every chapter has a clear purpose, a hook, and moves the plot forward. "
        "You are precise about word counts and never overpromise scope you "
        "can't deliver within the budget you're given."
    ),
    llm=outline_llm,
    verbose=True,
)

writer = Agent(
    role="Creative Fiction Writer",
    goal="Draft vivid, atmospheric, continuity-consistent chapters based on an approved outline.",
    backstory=(
        "You are a professional novelist writing mature, sophisticated genre "
        "fiction for adult readers, in the tradition of classic detective, "
        "thriller, and speculative fiction. You write immersive prose with "
        "strong sensory detail, natural dialogue, and forward-moving tension. "
        "You strictly follow the outline beat you're given for each chapter, "
        "stay consistent with what happened in prior chapters, and respect "
        "the target word count."
    ),
    llm=writer_llm,
    verbose=True,
)

# Note: the "designer" agent is created later, in Stage C — its cover tool needs
# to know the book's output folder and title, which don't exist until Stage A/B finish.


# =====================================================================
# 4. STAGE A — CONCEPT & OUTLINE
# =====================================================================

outline_task = Task(
    description=(
        f"Create an original {RUN_GENRE} story concept and full outline for a "
        f"FULL-LENGTH book.\n\n"
        f"Hard constraints:\n"
        f"- Total chapters: between {MIN_CHAPTERS} and {MAX_CHAPTERS}.\n"
        f"- Total word count across all chapters MUST be between {MIN_TOTAL_WORDS} and "
        f"{MAX_TOTAL_WORDS}. This is a floor as well as a ceiling — a book that comes in under "
        f"{MIN_TOTAL_WORDS} words is too short to publish, so plan enough story to fill it.\n"
        f"- Every chapter's target_words must be between {MIN_CHAPTER_TARGET_WORDS} and "
        f"{MAX_CHAPTER_TARGET_WORDS}. Do NOT exceed {MAX_CHAPTER_TARGET_WORDS} for any single "
        f"chapter — if you need more total length, add more chapters rather than making any one "
        f"chapter longer.\n"
        f"- Add up your target_words before answering and check the total lands in range. With "
        f"{MIN_CHAPTERS}+ chapters averaging around 2,800 words each, that works out to roughly "
        f"{MIN_TOTAL_WORDS}-{MAX_TOTAL_WORDS}.\n\n"
        "EMOTIONAL SHAPE — plan this deliberately:\n"
        "Give every chapter an emotional_pitch from 0 to 100, where 0 is quiet//reflective and "
        "100 is the most intense moment in the book. The book needs real variation here: at least "
        "one chapter below 30, at least one above 85, and the pitches should rise and fall rather "
        "than climb steadily. Quiet chapters after loud ones are what make the loud ones land.\n\n"
        "Respond with ONLY a valid JSON object (no commentary, no markdown fences) in exactly "
        "this shape:\n"
        "{\n"
        '  "title": "string",\n'
        '  "subtitle": "a short punchy cover tagline, 4-9 words, tailored specifically to THIS '
        'story\'s premise and stakes - not a generic genre phrase",\n'
        '  "genre": "string",\n'
        '  "premise": "2-3 sentence premise",\n'
        '  "main_characters": ["Name: one-line description", "Name: one-line description"],\n'
        '  "chapters": [\n'
        '    {"chapter_number": 1, "title": "string", "summary": "2-4 sentence beat summary", '
        '"target_words": 2800, "emotional_pitch": 45}\n'
        "  ]\n"
        "}\n\n"
        "main_characters MUST be a list of plain strings in the exact format \"Name: description\" "
        "— NOT objects, NOT {\"name\": ...} dicts. Every value in the JSON must use correct "
        "key: value syntax with colons after every key. Double-check your JSON is strictly valid "
        "before answering."
    ),
    expected_output="A single valid JSON object matching the schema described, with nothing else in the response.",
    agent=concept_architect,
)

# =====================================================================
# 0c. TIMING INSTRUMENTATION (dashboard section 6.14 — ETA/countdown timers)
# =====================================================================


def _emit_timing(event: str, stage: str, **fields) -> None:
    """Machine-readable timing marker, printed alongside the normal
    human-readable progress lines — dashboard_server.py's Job._run() parses
    [TIMING] lines out of the streamed stdout to build a per-stage duration
    history (job_timing_history.json) and live countdown ETAs. Safe to
    ignore reading the terminal by eye; it's an extra line, not a
    replacement for the existing progress prints. 'event' is 'start' or
    'end'; 'stage' is a stable key used to bucket history across runs (NOT
    the human-readable label, which can change wording without breaking
    ETA history). Duplicated identically in editorial_agent.py and
    scoring_agent.py — same "each script runs standalone" convention as
    the rest of this codebase."""
    parts = [f"event={event}", f"stage={stage}"]
    for key, value in fields.items():
        # The dashboard parses this line by whitespace-splitting tokens, so
        # any value containing spaces (e.g. a chapter title passed as
        # label=...) would otherwise fragment into bogus extra tokens.
        # Sanitize here only — this doesn't touch the human-readable prints
        # elsewhere in the script.
        safe_value = re.sub(r"\s+", "_", str(value)) if value is not None else ""
        parts.append(f"{key}={safe_value}")
    print("[TIMING] " + " ".join(parts))


print(f"--- Stage A: Generating concept & outline ({RUN_GENRE}) ---")
_stage_a_t0 = time.time()
_emit_timing("start", "outline")
outline_crew = Crew(
    agents=[concept_architect],
    tasks=[outline_task],
    process=Process.sequential,
    verbose=True,
)


REFUSAL_PATTERN = re.compile(
    r"^\s*(I cannot|I can[\u2019']t|I can not|I[\u2019']m sorry|I am sorry|"
    r"I[\u2019']m unable|I am unable|As an AI|I won[\u2019']t|"
    r"I[\u2019']m not able|I don[\u2019']t (feel comfortable|think I))",
    re.IGNORECASE,
)

# A refusal doesn't always come back as a short "I'm sorry, I can't help with that."
# It can also be a *long* non-answer where the model declines to write the chapter
# and instead hands back an outline/instructions for how the reader could write it
# themselves (this is exactly what happened on Chapter 14 of "Infinity's Edge" — a
# ~1,900-character bulleted outline with roman-numeral section headers, ending in
# "Please let me know if I can help with anything else."). The old length cutoff
# (<400 chars) let that one through since it only caught short refusals. These
# phrases are chosen to be fairly specific multi-word "the model is talking to the
# user about the task" tells, so they shouldn't false-positive on genuine prose that
# happens to open with a first-person line like "I can't believe what I'm seeing."
META_INSTRUCTION_PATTERN = re.compile(
    r"(here's an outline|here is an outline|outline that can serve as a guide|"
    r"you can continue this story|add your own creative elements|"
    r"let me know if (i can help|you (need|would like)|there's anything else)|"
    r"i can (guide|help) you (on |through )?how to|"
    r"keeping to the word count target|feel free to (expand|continue)|"
    r"would you like me to (continue|write)|as a starting point for)",
    re.IGNORECASE,
)

# A second structural tell for the same "outline instead of prose" failure mode:
# roman-numeral section headers (I., II., III. ...), often bolded — real chapter
# prose never opens with a numbered outline structure like this.
OUTLINE_STRUCTURE_PATTERN = re.compile(r"^\s*(\*\*)?[IVX]{1,4}\.\s+", re.MULTILINE)


def looks_like_refusal(text: str) -> bool:
    """Heuristic check for local-model safety-refusal text landing in a chapter file
    instead of actual prose. Not foolproof, but catches the common phrasings —
    including both straight (') and smart/curly (\u2019) apostrophe variants that
    local models frequently use interchangeably. Short refusals ("I'm sorry, I
    can't help with that.") are caught by the opening phrase + length check alone;
    longer non-answers (the model declining and handing back an outline/instructions
    instead of prose) are caught by also matching meta-instruction phrasing or an
    outline-style structure anywhere in the response."""
    stripped = text.strip()
    if not REFUSAL_PATTERN.match(stripped):
        return False
    if len(stripped) < 400:
        return True
    return bool(META_INSTRUCTION_PATTERN.search(stripped) or OUTLINE_STRUCTURE_PATTERN.search(stripped))


def _repair_common_json_mistakes(text: str) -> str:
    """Local models occasionally write main_characters entries as malformed
    pseudo-dicts instead of the requested plain strings, e.g.:
        {"name: Jack Murphy", "a former cop..."}
    instead of:
        "Jack Murphy: a former cop..."
    This breaks the entire JSON parse even though the rest of the document is
    fine. Detect and rewrite that specific pattern before falling back."""
    pattern = re.compile(r'\{\s*"name:\s*([^"]+?)"\s*,\s*"([^"]+?)"\s*\}')
    return pattern.sub(lambda m: f'"{m.group(1).strip()}: {m.group(2).strip()}"', text)


def parse_outline(raw_text: str, allow_short: bool = False) -> dict:
    """Pull the JSON object out of the model's response, tolerating stray text/fences.

    allow_short=True accepts a plan with fewer than MIN_CHAPTERS chapters and pads
    it out to a publishable length instead of rejecting it. Callers pass True only
    on their final attempt — see the retry loop below."""
    match = re.search(r"\{.*\}", raw_text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in outline output.")
    json_text = match.group(0)
    try:
        data = json.loads(json_text)
    except json.JSONDecodeError:
        data = json.loads(_repair_common_json_mistakes(json_text))

    # Clamp chapter count and word budget defensively in case the model ignored instructions.
    chapters = data.get("chapters", [])
    chapters = chapters[:MAX_CHAPTERS]
    # Only a genuinely tiny chapter list means the model misunderstood the task
    # and the response is worth rejecting. A merely SHORT list is normal — the
    # outline model habitually under-plans — and rejecting those would be
    # actively harmful now that MIN_CHAPTERS is 10: three rejected attempts fall
    # through to the generic placeholder outline ("Story continues." for every
    # beat), which is a far worse book than a real 7-chapter outline padded out
    # by the length logic below. So short lists get grown, not thrown away.
    if len(chapters) < ABSOLUTE_MIN_CHAPTERS:
        raise ValueError(
            f"Model returned only {len(chapters)} chapters, which reads as a malformed response "
            f"rather than a short plan (need at least {ABSOLUTE_MIN_CHAPTERS})."
        )
    # A short-but-valid plan is worth one more roll of the dice before we pad it.
    # Padding works, but the added chapters carry a generic summary, and a book
    # where half the beats read "continue developing the plot" is a worse book
    # than one the model actually planned. So attempts 1..n-1 reject a short list
    # and ask again; the final attempt accepts whatever came back and pads it, so
    # a stubborn model still yields a full-length book rather than nothing.
    if len(chapters) < MIN_CHAPTERS and not allow_short:
        raise ValueError(
            f"Model planned only {len(chapters)} chapters; asked for at least {MIN_CHAPTERS} so the "
            f"book reaches a publishable length without padded filler beats."
        )

    # Clamp every chapter into the per-chapter band FIRST. The ceiling is the
    # one that matters for quality: collapse risk rises with chapter length (see
    # MAX_CHAPTER_TARGET_WORDS for the measurements), so an outline asking for a
    # 4,500-word chapter gets it cut whatever the totals say.
    for c in chapters:
        target = c.get("target_words") or 2500
        try:
            target = int(target)
        except (TypeError, ValueError):
            target = 2500
        c["target_words"] = max(MIN_CHAPTER_TARGET_WORDS,
                                min(MAX_CHAPTER_TARGET_WORDS, target))

    # Then bring the TOTAL into range. The scale-down half of this already
    # existed; the scale-up half is new as of 2026-08-14, because the outline
    # model reliably under-asks. The first clean run produced 11,604 words — 29%
    # of the budget — purely because nothing enforced a floor.
    total_words = sum(c["target_words"] for c in chapters)
    if total_words > MAX_TOTAL_WORDS:
        scale = MAX_TOTAL_WORDS / total_words
        for c in chapters:
            c["target_words"] = max(MIN_CHAPTER_TARGET_WORDS, int(c["target_words"] * scale))
        print(f"[Outline] Total target was {total_words} words (over the {MAX_TOTAL_WORDS} "
              f"ceiling) — scaled chapters down to {sum(c['target_words'] for c in chapters)}.")
    elif total_words < MIN_TOTAL_WORDS:
        # Two levers, in this order: grow the chapters we have up to the
        # per-chapter cap, then add chapters if that still isn't enough. Growing
        # first keeps the outline's own story structure intact; adding chapters
        # means inventing beats the model didn't plan, so it's the last resort.
        scale = MIN_TOTAL_WORDS / total_words
        for c in chapters:
            c["target_words"] = min(MAX_CHAPTER_TARGET_WORDS, int(c["target_words"] * scale))
        grown = sum(c["target_words"] for c in chapters)
        print(f"[Outline] Total target was only {total_words} words (under the "
              f"{MIN_TOTAL_WORDS} floor) — scaled chapters up to {grown}.")

        if grown < MIN_TOTAL_WORDS and len(chapters) < MAX_CHAPTERS:
            # Every existing chapter is already at the cap, so the only way to
            # reach a publishable length is more chapters. These are appended
            # with a deliberately honest placeholder summary rather than a
            # fabricated beat — the writer handles a thin summary far better
            # than a specific one that contradicts the real plot.
            needed = MIN_TOTAL_WORDS - grown
            to_add = min(MAX_CHAPTERS - len(chapters),
                         -(-needed // MAX_CHAPTER_TARGET_WORDS))  # ceil division
            next_num = max(c["chapter_number"] for c in chapters) + 1
            for i in range(to_add):
                chapters.append({
                    "chapter_number": next_num + i,
                    "title": f"Chapter {next_num + i}",
                    "summary": ("Continue developing the plot and characters toward the "
                                "resolution set up by the preceding chapters."),
                    "target_words": MAX_CHAPTER_TARGET_WORDS,
                    "emotional_pitch": 55,
                })
            print(f"[Outline] Still short after scaling — added {to_add} chapter(s) to reach "
                  f"{sum(c['target_words'] for c in chapters)} words.")

    # Emotional pitch drives the per-chapter tone instruction in Stage B. Missing
    # or nonsense values default to mid-range so a chapter never loses its
    # instruction entirely; the scorer separately rewards real variation here
    # (emotional_resonance measured 34/100 on a book with a spread of only 5.1).
    for c in chapters:
        pitch = c.get("emotional_pitch")
        try:
            pitch = int(pitch)
        except (TypeError, ValueError):
            pitch = 50
        c["emotional_pitch"] = max(0, min(100, pitch))

    data["chapters"] = chapters

    # Subtitle is a nice-to-have cover element, not a hard requirement — sanitize it
    # if present but never fail outline parsing over it. compose_cover() skips the
    # subtitle band entirely when this is empty, so a missing/blank value degrades
    # gracefully instead of putting placeholder text on the cover.
    subtitle = data.get("subtitle", "")
    data["subtitle"] = subtitle.strip().strip('"').strip("'") if isinstance(subtitle, str) else ""

    return data


# Local models are inconsistent about producing strictly valid JSON on the first try.
# A malformed-JSON fallback to the generic placeholder outline ("Story continues.",
# vague premise) tends to make Stage B chapter refusals *much* more likely, since the
# writer has nothing concrete to work with. Retrying the same task a couple of times
# is cheap and usually succeeds, so we do that before giving up and falling back.
MAX_OUTLINE_ATTEMPTS = 3
outline = None
last_error = None
for attempt in range(1, MAX_OUTLINE_ATTEMPTS + 1):
    outline_raw = str(outline_crew.kickoff())
    try:
        # Last attempt takes what it can get and pads; earlier ones hold out for
        # a plan the model actually thought through.
        outline = parse_outline(outline_raw, allow_short=(attempt == MAX_OUTLINE_ATTEMPTS))
        break
    except (ValueError, json.JSONDecodeError) as e:
        last_error = e
        print(f"\n[WARN] Outline attempt {attempt}/{MAX_OUTLINE_ATTEMPTS} failed to parse ({e}).")
        if attempt < MAX_OUTLINE_ATTEMPTS:
            print("[WARN] Retrying outline generation...")

if outline is None:
    print(
        f"\n[WARN] All {MAX_OUTLINE_ATTEMPTS} outline attempts failed to parse ({last_error}). "
        f"Using a fallback outline so the run can continue."
    )
    outline = {
        "title": f"Untitled {RUN_GENRE.title()} Novel",
        "subtitle": f"A Novel of {RUN_GENRE.title()}",
        "genre": RUN_GENRE,
        "premise": "A protagonist is pulled into a mystery that escalates chapter by chapter.",
        "main_characters": ["Protagonist: determined, resourceful"],
        "chapters": [
            {"chapter_number": i + 1, "title": f"Chapter {i + 1}", "summary": "Story continues.", "target_words": 2500}
            for i in range(MIN_CHAPTERS)
        ],
    }

_emit_timing("end", "outline", elapsed=f"{time.time() - _stage_a_t0:.1f}")
print(f"\n--- Outline ready: '{outline['title']}' ({len(outline['chapters'])} chapters) ---")

# Catch a title this local model has already used in a past run before it becomes
# the folder name / cover title / manuscript title — see "0b. TITLE HISTORY" above.
# Skipped when --writer-model is explicitly overriding the default: A/B testing
# writer models against the same premise/title would otherwise trip this guard
# against itself (title already "used" by the first model's run, second run gets
# silently renamed "... Part 2") — see dashboard section 6.12.
if WRITER_MODEL_OVERRIDE:
    print(f"[Config] --writer-model override active ({WRITER_MODEL_OVERRIDE}) — skipping title-history "
          f"dedupe/recording for this run so A/B testing doesn't trip the duplicate-title guard against itself.")
else:
    outline["title"] = dedupe_title(outline["title"], RUN_GENRE)

# =====================================================================
# 5. STAGE B — CHAPTER-BY-CHAPTER WRITING (dynamic task list, continuity via context)
# =====================================================================

slug = re.sub(r"[^a-z0-9]+", "-", outline["title"].lower()).strip("-") or "untitled-book"
run_stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
book_dir = os.path.join(OUTPUT_ROOT, f"{slug}-{run_stamp}")
os.makedirs(book_dir, exist_ok=True)

# Reclaim VRAM before writing starts, added 2026-08-14. The earlier fix only
# unloaded A1111 at the END of Stage C and at the start of the editor/scorer,
# which missed the longest stage of all. "Launch All Services" starts
# AUTOMATIC1111 up front and it loads its checkpoint immediately, so on the
# 2026-08-13 runs SDXL sat in VRAM through the entire writing stage — the
# dashboard read "1,445 / 8,188 MB free" while chapters were being written, and
# system RAM hit 91% with llama-server.exe holding 7.6GB. A1111 reloads its
# checkpoint by itself when Stage C asks for a cover, by which point Ollama has
# been unloaded and there's room for it.
unload_a1111_checkpoint()
# Drop the OUTLINE model too before measuring. Added 2026-08-14: without this
# the reading is taken while llama3.1 from Stage A is still resident, so it came
# back at ~2,670MB free on two consecutive runs and fired the low-VRAM warning
# every time — even though both runs then wrote chapters at a healthy 90-167s
# because Ollama evicts the outline model when the writer model loads. A warning
# that cries wolf on every run is worse than no warning, since the one time it
# matters it gets ignored. Unloading first makes the number mean what it says.
unload_ollama_model(OUTLINE_MODEL_NAME)
_free_before_write = get_free_vram_mb()
if _free_before_write is not None:
    print(f"[VRAM] {_free_before_write}MB free going into the writing stage.")
    if _free_before_write < 6500:
        print(f"[WARN] '{WRITER_MODEL_NAME}' needs roughly 6.5GB to stay fully on the GPU. Below "
              f"that Ollama silently offloads layers to system RAM and runs them on the CPU — "
              f"2-3x slower, and the cause of the stalls on 2026-08-13. Check for another process "
              f"holding VRAM before letting a long run continue.")

with open(os.path.join(book_dir, "outline.json"), "w", encoding="utf-8") as f:
    json.dump(outline, f, indent=2)

# Per-book model log (dashboard section 6.12) — records which writer model
# actually wrote this book, regardless of whether it was the default or an
# explicit --writer-model override, so a later A/B comparison in the dashboard
# doesn't have to guess from context.
model_info = {
    "writer_model": WRITER_MODEL_NAME,
    "writer_model_overridden": bool(WRITER_MODEL_OVERRIDE),
    "outline_model": OUTLINE_MODEL_NAME,
}
with open(os.path.join(book_dir, "model_info.json"), "w", encoding="utf-8") as f:
    json.dump(model_info, f, indent=2)

MAX_CHAPTER_ATTEMPTS = 5  # was 3 — stock llama3.1 refusals are flaky/probabilistic
# Separate, much smaller budget for repetition collapses — a refusal retry is
# cheap (the model declines in a sentence), a collapse reroll costs a whole
# generation. See the comment in the chapter loop for why 3 of them in a row
# was a 25-minute dead end rather than three chances at better luck.
MAX_DEGEN_REROLLS = 2
                          # (same prompt, same premise, refuses 3/3 times on one chapter
                          # then writes fine on attempt 2/3 for another), so a couple of
                          # extra tries meaningfully improves the odds without real cost
                          # for a local model. Chapter 1 of "Starfall Rebellion" refused
                          # all 3 of the old attempts on totally mainstream rebellion-vs-
                          # empire content — this is the known base-model behavior noted
                          # above, not a sign anything is wrong with the premise.
CONTINUITY_TAIL_CHARS = 1500  # how much of the previous chapter's ending to carry forward
# How much of the previous chapter's ending to ASSESS before deciding the tail
# above is safe to carry forward. Deliberately larger than the seed itself —
# see the comment at the check site.
CONTINUITY_CHECK_WORDS = 800

# Sentence-length ceiling for the continuity seed, added 2026-08-15.
#
# MEASURED, across the raw drafts of the 7 books with complete chapter sets:
#
#   Chapter 1 — the ONLY chapter that gets no continuity seed:
#       mean 26.5 words/sentence, stdev 1.7, range 25.2-30.6  (n=7)
#   Chapters 2+ — every one of them seeded:
#       mean 31.8 words/sentence, stdev 7.9, range 19.3-55.0  (n=37)
#
# Unseeded, this writer is consistent — seven books across seven genres land
# within a few words of each other, close to what the PROSE STYLE block below
# asks for. Seeded, the mean climbs 5 words and the spread widens 4.5x. Chapter
# 1 is not written differently; it is the only chapter whose prompt ends with an
# instruction rather than with an example.
#
# The propagation is specifically through the SEED, not through the previous
# chapter generally. Correlating each chapter's average sentence length against
# the chapter before it, centred within book so this isn't a between-book
# artifact, with a permutation test since n is small:
#
#   previous chapter's TAIL (the 1500 chars actually pasted in): r = +0.42, p = 0.012
#   previous chapter's BODY (its overall style):                 r = +0.18, p = 0.27
#
# The text we paste in predicts the next chapter; the text we don't paste in
# doesn't. That is the causal signature — the model is imitating the sample, and
# since the sample is the LAST thing in the prompt it outranks the abstract
# style instruction earlier. Once a chapter runs long the next one inherits it,
# and nothing pulls it back: a ratchet with no restoring force. Starbound
# Odyssey climbed 26.8 -> 55.0 over six chapters that way.
#
# 32.0 is set above every tail Rainy Night Requiem produced (its worst was 29.6)
# and below the ones that ran away, so the book this pipeline got most right is
# left completely untouched while the runaway books get the summary fallback.
CONTINUITY_MAX_TAIL_SENTENCE = 32.0


def average_sentence_words(text: str) -> float:
    """Mean words per sentence. Same splitter and word pattern the scorer uses,
    so this measures what the scorer will later grade."""
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return 0.0
    sentences = [p for p in re.split(r"(?<=[.!?])\s+", clean) if p.strip()]
    lengths = [len(re.findall(r"[A-Za-z']+", s)) for s in sentences]
    lengths = [n for n in lengths if n]
    if not lengths:
        return 0.0
    return sum(lengths) / len(lengths)

# A short reminder that this is fiction, placed right up front before the (sometimes
# refusal-triggering) beat/premise text. This does NOT reliably prevent llama3.1's
# keyword-triggered refusals on its own — see the WRITER_MODEL_NAME note above — but
# it's free, harmless, and stacks with the retry loop and fallback below.
FICTION_FRAMING = (
    "This is entirely fictional, original genre fiction being written for a published "
    "novel — not real-world advice, instructions, or commentary. Write the chapter "
    "directly as narrative prose.\n\n"
)

prev_chapter_text = None  # only ever set to genuine prose, never a refusal — see note above
chapters_needing_regeneration = []  # (ch_num, ch_title) pairs that never got past a refusal
# (ch_num, ch_title, outcome) for chapters where the model spiralled into
# repetition — tracked separately from refusals because the cause and the fix
# are different (see detect_degeneration()). Surfaced in the run summary so a
# degraded chapter can't slip by unnoticed the way Rashomon's did.
chapters_degenerated = []

# Lookup for the continuity-seed fallback below, which needs the PREVIOUS
# chapter's outline summary when that chapter's prose isn't safe to carry.
chapters_by_number = {c["chapter_number"]: c for c in outline["chapters"]}

for chapter in outline["chapters"]:
    ch_num = chapter["chapter_number"]
    ch_title = chapter["title"]
    ch_summary = chapter["summary"]
    ch_words = chapter["target_words"]

    # Continuity seed. Vetted before use as of 2026-08-14 — this was the reason
    # a single bad chapter poisoned every chapter after it.
    #
    # What happened on the 17:11 run: Chapter 4 drifted onto the word
    # "laughter" (5.5% of a 400-word stretch). The guard salvaged its coherent
    # opening, but the continuity seed is the LAST 1500 characters of whatever
    # is carried forward — which came from exactly the drifting end of that
    # chapter. Chapter 5 was then handed degraded prose and told "continue
    # naturally from this point", so it reproduced the same collapse — on
    # "laughter" again, at 6.2% — and did so on all three attempts, because
    # every reroll re-used the same poisoned seed. Rerolling cannot escape a
    # bad seed; the seed has to be checked.
    #
    # If the tail doesn't pass the collapse check, fall back to the previous
    # chapter's outline summary. That's weaker continuity than real prose, but
    # weaker continuity is far better than deliberately seeding the next
    # chapter with the failure mode we're trying to prevent.
    continuity_note = ""
    if prev_chapter_text:
        tail = prev_chapter_text[-CONTINUITY_TAIL_CHARS:]
        # Judge the health of the chapter's ENDING using a wider slice than the
        # 1500 characters actually pasted in. The first version of this check
        # ran the detector on the seed itself and was dead code: 1500 characters
        # is only ~280 words, below the detector's 400-word minimum window, so
        # it returned "too short to assess" every single time and the poisoned
        # tail sailed through anyway. Checking the last 800 words instead gives
        # the sliding window room to work and reuses the thresholds already
        # calibrated against 19 real chapters — verified to flag all 5 known-bad
        # endings and none of the 14 known-good ones.
        tail_verdict = detect_degeneration(
            " ".join(prev_chapter_text.split()[-CONTINUITY_CHECK_WORDS:])
        )
        # Two independent reasons to refuse the tail. The collapse check catches
        # prose that has broken down; the sentence-length check catches prose
        # that is intact but stylistically runaway. They fail differently and
        # are reported differently, but the remedy is the same — fall back to
        # the outline summary, which carries the events without carrying the
        # style. See CONTINUITY_MAX_TAIL_SENTENCE above for the measurements.
        tail_sentence_avg = average_sentence_words(tail)
        tail_too_long = tail_sentence_avg > CONTINUITY_MAX_TAIL_SENTENCE

        if tail_verdict["degenerate"] or tail_too_long:
            if tail_verdict["degenerate"]:
                print(f"[Writer] Previous chapter's ending doesn't pass the coherence check "
                      f"({tail_verdict['detail']}). Seeding continuity from the outline summary "
                      f"instead of that prose, so the drift isn't carried forward.")
            else:
                print(f"[Writer] Previous chapter's ending averages "
                      f"{tail_sentence_avg:.0f} words/sentence (ceiling "
                      f"{CONTINUITY_MAX_TAIL_SENTENCE:.0f}). Seeding continuity from the outline "
                      f"summary instead, so the long-sentence style isn't carried forward.")
            prev_meta = chapters_by_number.get(ch_num - 1, {})
            prev_summary = prev_meta.get("summary", "")
            if prev_summary:
                continuity_note = (
                    f"\n\nFor continuity, the previous chapter covered this beat: "
                    f"{prev_summary}\n"
                    "Continue the story naturally from there; do not repeat or restate it."
                )
        else:
            # The seed is labelled as reference material and followed by a
            # restatement of the length constraint. Ordering is deliberate: this
            # block is the last thing in the prompt, and until now that last
            # thing was a raw block of prose with no instruction after it, which
            # is exactly how you ask a model to imitate something.
            continuity_note = (
                f"\n\nFor continuity, here is how the previous chapter ended:\n"
                f"\"\"\"\n{tail}\n\"\"\"\n"
                "Continue the story naturally from this point; do not repeat or restate it.\n"
                "This excerpt is here for CONTINUITY OF EVENTS ONLY — for who is present, where "
                "they are, and what just happened. It is NOT a style sample. Do not match its "
                "sentence rhythm or paragraph length. The PROSE STYLE guidance above still "
                f"applies in full: keep your sentences short to medium, averaging well under "
                f"{CONTINUITY_MAX_TAIL_SENTENCE:.0f} words, and break long ones at the joint."
            )

    # Emotional pitch for this chapter, from the outline. Added 2026-08-14: the
    # scorer's emotional_resonance metric rewards a real spread of intensity
    # across the book and scored 34/100 on a book whose chapter-to-chapter spread
    # was only 5.1 — every chapter written at the same emotional volume. Telling
    # the writer where this chapter sits is what creates the variation.
    pitch = chapter.get("emotional_pitch", 50)
    if pitch >= 85:
        pitch_note = (
            f"\nEmotional pitch for this chapter: {pitch}/100 — this is one of the book's peaks. "
            "Write it at high intensity: short sentences, fast cuts, physical stakes, dialogue "
            "that lands like blows. Do not let the tension slacken."
        )
    elif pitch <= 30:
        pitch_note = (
            f"\nEmotional pitch for this chapter: {pitch}/100 — this is a deliberately quiet "
            "chapter. Let it breathe: longer beats, reflection, small human detail, characters "
            "processing what just happened. Its job is to make the next peak land harder, so "
            "resist manufacturing drama here."
        )
    else:
        pitch_note = (
            f"\nEmotional pitch for this chapter: {pitch}/100 — moderate intensity. Build and "
            "release tension within the chapter rather than holding one pitch throughout."
        )

    description = (
        f"{FICTION_FRAMING}"
        f"Write Chapter {ch_num}: '{ch_title}' of the book '{outline['title']}' ({outline['genre']}).\n\n"
        f"Chapter beat to cover: {ch_summary}\n\n"
        f"Book premise (for context): {outline['premise']}\n\n"
        f"Target length: approximately {ch_words} words. Stay within +/-15% of that target.\n"
        f"{pitch_note}\n"
        # Readability guidance added 2026-08-14. The scorer targets a
        # Flesch-Kincaid grade of 7-10 for commercial genre fiction and the last
        # book measured 14.4, costing it 53 points on that metric. Grade level is
        # driven almost entirely by sentence length and syllables per word, so
        # the instruction targets exactly those two things rather than asking
        # vaguely for "simpler prose" — which tends to flatten voice instead.
        "\nPROSE STYLE — this genre's readers expect accessible prose, roughly an 8th-to-10th "
        "grade reading level. That is about sentence construction, not about dumbing anything "
        "down:\n"
        "- Favour short and medium sentences. Break long ones at the natural joint instead of "
        "stacking three subordinate clauses into one.\n"
        "- Prefer the concrete, everyday word over the ornate one where both work — 'walked' over "
        "'perambulated', 'dark' over 'tenebrous'.\n"
        "- Vary rhythm deliberately: a run of short sentences, then a longer one. Uniform "
        "sentence length reads as monotonous however good the words are.\n"
        "- Keep vivid sensory detail and distinctive voice. Accessible does not mean plain.\n"
        "\nWrite only the chapter prose itself — no title header, no meta-commentary, no chapter "
        "summary."
        f"{continuity_note}"
    )

    print(f"\n--- Stage B: Writing Chapter {ch_num}/{len(outline['chapters'])}: '{ch_title}' ---")
    _chapter_t0 = time.time()
    _emit_timing("start", "chapter_write", ch=ch_num, total=len(outline["chapters"]), label=ch_title)

    # Per-chapter output ceiling — see CHAPTER_WORD_CAP_MULTIPLIER above.
    word_cap = int(ch_words * CHAPTER_WORD_CAP_MULTIPLIER)
    num_predict = int(word_cap * TOKENS_PER_WORD)
    print(f"[Writer] Chapter {ch_num} target {ch_words} words, hard cap {word_cap} "
          f"(num_predict {num_predict}), repeat_penalty {WRITER_REPEAT_PENALTY}.")

    chapter_text = None
    best_partial = None      # longest verified-clean prefix seen across attempts
    degen = None
    degen_rerolls = 0
    # Refusals and collapses need different retry budgets, split 2026-08-14.
    # A refusal costs almost nothing to retry — the model declines in a
    # sentence or two, so 5 attempts is cheap. A collapse costs a full-length
    # generation each time. On the 17:11 run Chapter 5 burned three of them
    # (~25 minutes) and failed identically every time, because the cause was
    # the poisoned continuity seed rather than bad luck — rerolling could never
    # have fixed it. Two collapse rerolls is enough to rule out bad luck;
    # beyond that, take the salvage and move on.
    for attempt in range(1, MAX_CHAPTER_ATTEMPTS + 1):
        try:
            result = generate_chapter_streaming(description, num_predict, ch_num)
        except requests.exceptions.RequestException as e:
            print(f"[WARN] Chapter {ch_num} attempt {attempt}/{MAX_CHAPTER_ATTEMPTS}: Ollama request "
                  f"failed ({e}).")
            if attempt < MAX_CHAPTER_ATTEMPTS:
                print("[WARN] Retrying...")
                continue
            chapter_text = ""
            break

        chapter_text = result["text"]
        if result["aborted_early"]:
            print(f"[WARN] Chapter {ch_num} attempt {attempt}: aborted mid-generation — "
                  f"{result['reason']}. (Stopped early rather than paying for the rest.)")

        if looks_like_refusal(chapter_text):
            print(
                f"[WARN] Chapter {ch_num} attempt {attempt}/{MAX_CHAPTER_ATTEMPTS} looks like a model "
                f"refusal rather than actual prose."
                + (" Retrying..." if attempt < MAX_CHAPTER_ATTEMPTS else "")
            )
            continue

        # Collapse gate. A chapter that spirals into repetition is worse than a
        # refusal: a refusal is obvious on sight and already gets a loud
        # placeholder, whereas 5,000 words of looping prose looks like a
        # finished chapter, passes the editor, and — as of the Rashomon run —
        # scored 80/100.
        degen = detect_degeneration(chapter_text)
        if not degen["degenerate"]:
            written = degen["total_words"]
            if word_cap and written > word_cap * 1.15:
                # Not degenerate, just long. Trim at a paragraph break rather
                # than rerolling — the prose itself is fine. With num_predict
                # now actually enforced this should be rare; it was routine
                # while the cap was being silently dropped by litellm.
                print(f"[Writer] Chapter {ch_num} came back at {written} words against a "
                      f"{ch_words}-word target; trimming to ~{word_cap} at a paragraph break.")
                chapter_text = truncate_at_word_index(chapter_text, word_cap)
            break

        print(f"[WARN] Chapter {ch_num} attempt {attempt}: {degen['detail']}.")

        # A salvage is only a salvage if the surviving prose actually passes the
        # check — never assume it. Cutting at the first bad window should make
        # the prefix clean by construction, but verifying costs nothing, and
        # this is exactly the assumption that produced a "salvaged" chapter
        # which was still broken while this guard was being built.
        salvaged_text = None
        if (degen["first_bad_word_index"] or 0) >= ch_words * 0.6:
            candidate = truncate_at_word_index(chapter_text, degen["first_bad_word_index"])
            if not detect_degeneration(candidate)["degenerate"]:
                salvaged_text = candidate
                if best_partial is None or len(_DEGEN_WORD_RE.findall(candidate)) > len(
                        _DEGEN_WORD_RE.findall(best_partial)):
                    best_partial = candidate
            else:
                print("[WARN] The opening doesn't pass the coherence check either — it degraded "
                      "earlier than the worst stretch. Not salvageable.")

        degen_rerolls += 1
        if degen_rerolls < MAX_DEGEN_REROLLS and attempt < MAX_CHAPTER_ATTEMPTS:
            print(f"[WARN] Rerolling (collapse reroll {degen_rerolls}/{MAX_DEGEN_REROLLS})." + (
                f" A coherent {len(_DEGEN_WORD_RE.findall(best_partial))}-word opening is held as a "
                f"fallback." if best_partial else ""))
            continue

        # Out of collapse rerolls. Prefer a shorter coherent chapter over a
        # full-length broken one: a 2,200-word chapter that reads properly is
        # publishable, 8,300 words ending in "hell hell hell" is not.
        if best_partial:
            print(f"[WARN] Giving up on a clean full-length Chapter {ch_num} after "
                  f"{degen_rerolls} collapse reroll(s). Keeping the coherent opening "
                  f"({len(_DEGEN_WORD_RE.findall(best_partial))} words) and discarding the "
                  f"repetition. This chapter will read short — worth a manual reroll.")
            chapter_text = best_partial
            chapters_degenerated.append((ch_num, ch_title, "truncated to coherent opening"))
        else:
            print(f"[WARN] Chapter {ch_num} collapsed on every attempt with no salvageable "
                  f"opening. Writing a placeholder.")
            chapters_degenerated.append((ch_num, ch_title, "no usable prose"))
            chapter_text = (
                f"[NEEDS MANUAL REGENERATION — the writer model collapsed into repetition on every "
                f"attempt at this chapter ({degen['detail']}). This is a generation failure, not a "
                f"refusal: try a shorter target_words for this chapter, or reroll it on its own.]"
            )
        break

    path = os.path.join(book_dir, f"chapter_{ch_num:02d}.txt")

    if looks_like_refusal(chapter_text):
        # Don't ship the model's raw refusal text into the manuscript disguised as a
        # chapter — this exact problem showed up when Chapter 14 of a previous book
        # was just the model's decline sitting where prose should be. Write an
        # unmistakable placeholder instead, so a skim of the manuscript makes it
        # obvious which chapters still need a manual regenerate, instead of reading
        # like a broken/unfinished sentence.
        print(
            f"[WARN] Chapter {ch_num} ('{ch_title}') still looks like a refusal after "
            f"{MAX_CHAPTER_ATTEMPTS} attempts. Writing a placeholder — you'll want to "
            f"regenerate this chapter manually. See the terminal output above for the "
            f"model's exact responses."
        )
        chapters_needing_regeneration.append((ch_num, ch_title))
        chapter_text = (
            f"[NEEDS MANUAL REGENERATION — the local writer model declined to write this "
            f"chapter {MAX_CHAPTER_ATTEMPTS} times in a row. WRITER_MODEL_NAME is already set "
            f"to an uncensored storytelling model, so this particular beat likely just needs a "
            f"reroll rather than a model change — re-run just this chapter. If a specific "
            f"chapter beat keeps failing repeatedly, that's worth rewording in the outline.]"
        )
        # Deliberately NOT updating prev_chapter_text here — feeding a refusal/placeholder
        # forward as "what happened last chapter" tends to prime the model to keep refusing.
    else:
        prev_chapter_text = chapter_text

    with open(path, "w", encoding="utf-8") as fh:
        fh.write(f"{ch_title}\n\n{chapter_text}")
    print(f"[Saved] {path}")
    _emit_timing("end", "chapter_write", ch=ch_num, total=len(outline["chapters"]),
                 elapsed=f"{time.time() - _chapter_t0:.1f}")

print(f"\n--- Stage B complete: {len(outline['chapters'])} chapters processed ---")

# Free the VRAM Ollama has been holding since Stage A/B before Stage C asks
# AUTOMATIC1111 to load SD/SDXL into it — see "2b. VRAM UTILITIES" above.
unload_ollama_model(WRITER_MODEL_NAME)

# =====================================================================
# 6. STAGE C — COVER GENERATION
# =====================================================================

# Use RUN_GENRE (the exact lowercase string Python picked from GENRE_POOL), not
# outline["genre"], for every genre-keyed dict lookup below. The model echoes the
# genre back into its own JSON output, but not reliably in the same casing it was
# given ("Epic Fantasy" vs "epic fantasy") — a dict.get() miss from that mismatch
# silently falls back to the sci-fi thriller style/font/glow for every other genre,
# which is exactly what happened on a real test run before this was caught.
style_hint = GENRE_STYLE_HINTS.get(RUN_GENRE, "atmospheric genre-fiction illustration")

cover_tool = Auto1111CoverTool(
    output_dir=book_dir,
    book_title=outline["title"],
    author_name=AUTHOR_NAME,
    genre=RUN_GENRE,
    subtitle=outline.get("subtitle", ""),
)

# NOTE: earlier versions had the "designer" agent call cover_tool directly via
# CrewAI's tool-calling. Local models served through Ollama are unreliable at
# actually triggering tool calls — they often just print text that *looks like*
# a tool-call JSON blob instead of invoking anything, silently skipping cover
# generation entirely. To make this deterministic, the agent's only job now is
# to write the art prompt as plain text (something local models do reliably),
# and we invoke the AUTOMATIC1111 tool ourselves in plain Python afterward.
designer = Agent(
    role="Book Cover Art Director",
    goal="Translate story themes into a strong, vivid visual art prompt for a book cover.",
    backstory="You specialize in digital illustration prompts for genre fiction covers.",
    llm=outline_llm,
    verbose=True,
)

detail_hint = GENRE_DETAIL_HINTS.get(RUN_GENRE, "")

cover_prompt_task = Task(
    description=(
        f"Write a single, rich visual art prompt (2-4 sentences) for this book's cover ART ONLY.\n\n"
        f"Title: {outline['title']}\n"
        f"Genre: {outline['genre']}\n"
        f"Premise: {outline['premise']}\n"
        f"Style direction for this genre: {style_hint}\n"
        f"Also work in this detail if it fits naturally: {detail_hint}\n\n"
        "IMPORTANT: Do not include any title text, lettering, or typography in the "
        "prompt itself. Focus entirely on composition, subject, color, lighting, and mood.\n\n"
        "Respond with ONLY the prompt text itself — no commentary, no JSON, no markdown "
        "fences, no surrounding quotation marks, nothing else."
    ),
    expected_output="A single plain-text paragraph describing the cover art, with nothing else in the response.",
    agent=designer,
)

print("\n--- Stage C: Generating cover ---")
_stage_c_t0 = time.time()
_emit_timing("start", "cover_gen")
cover_crew = Crew(
    agents=[designer],
    tasks=[cover_prompt_task],
    process=Process.sequential,
    verbose=True,
)
art_prompt_raw = str(cover_crew.kickoff())


def clean_art_prompt(raw_text: str) -> str:
    """Local models sometimes wrap the prompt in JSON/quotes/fences despite instructions
    not to. Strip common wrapping so we're left with just the descriptive prompt text."""
    text = raw_text.strip()
    if text.startswith("{"):
        try:
            data = json.loads(text)
            if isinstance(data, dict):
                if "prompt" in data:
                    text = data["prompt"]
                elif "parameters" in data and isinstance(data["parameters"], dict) and "prompt" in data["parameters"]:
                    text = data["parameters"]["prompt"]
        except (ValueError, KeyError, TypeError):
            pass
    text = re.sub(r"^```[a-zA-Z]*\n|```$", "", text).strip()
    if len(text) >= 2 and text[0] == text[-1] and text[0] in ('"', "'"):
        text = text[1:-1].strip()
    return text or f"{style_hint}, book cover illustration for a {outline['genre']} novel"


art_prompt = clean_art_prompt(art_prompt_raw)
print(f"[Stage C] Using art prompt: {art_prompt}")

cover_result = cover_tool._run(prompt=art_prompt)
print(f"[Stage C] {cover_result}")
_emit_timing("end", "cover_gen", elapsed=f"{time.time() - _stage_c_t0:.1f}")

# =====================================================================
# 7. STAGE D — MANUSCRIPT ASSEMBLY (raw draft + formatted .docx)
# =====================================================================


def _read_chapter_body(path: str) -> str:
    """Chapter files were saved as 'Title\\n\\nBody' by the Stage B callback — strip the title line."""
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    parts = content.split("\n\n", 1)
    return parts[1].strip() if len(parts) > 1 else content.strip()


def build_raw_manuscript(book_dir: str, outline: dict) -> str:
    """Compile all chapters into a single plain-text draft file."""
    lines = [outline["title"], f"by {AUTHOR_NAME}", ""]
    for chapter in outline["chapters"]:
        path = os.path.join(book_dir, f"chapter_{chapter['chapter_number']:02d}.txt")
        if not os.path.isfile(path):
            print(f"[WARN] Missing {path}, skipping in raw manuscript.")
            continue
        lines.append(f"\n\nChapter {chapter['chapter_number']}: {chapter['title']}\n")
        lines.append(_read_chapter_body(path))

    output_path = os.path.join(book_dir, "manuscript_raw.txt")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path


def _add_static_toc(document: Document, chapters: list) -> None:
    """Write a plain, pre-rendered chapter list — 'Chapter N — Title', one per
    line, no page numbers (ebooks reflow, so a printed page number would just be
    wrong on an e-reader anyway). Replaces an earlier real Word TOC *field*
    (dynamic, computed by Word itself on 'Update Field'): that approach reads
    fine in Word, but ebook converters generally don't run
    Word's field-update logic, so it imported the field as a permanently empty
    Contents page — confirmed by a trial-publish run, 2026-08-12. A
    static list has real text baked in at build time, so it always renders
    correctly no matter what app opens or converts the file."""
    for chapter in chapters:
        line = document.add_paragraph()
        line.add_run(f"Chapter {chapter['chapter_number']} — {chapter['title']}")


def build_docx_manuscript(book_dir: str, outline: dict) -> str:
    """Build a formatted .docx: title page, copyright page, TOC, then chapters."""
    document = Document()

    # Base body style
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # --- Title page ---
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(outline["title"])
    title_run.font.size = Pt(28)
    title_run.font.bold = True

    for _ in range(3):
        document.add_paragraph()

    by_p = document.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_run = by_p.add_run(f"by {AUTHOR_NAME}")
    by_run.font.size = Pt(16)

    document.add_page_break()

    # --- Copyright page ---
    year = datetime.now().year
    copyright_p = document.add_paragraph()
    copyright_p.add_run(f"Copyright \u00a9 {year} {AUTHOR_NAME}").bold = True
    document.add_paragraph("All rights reserved.")
    document.add_paragraph()
    document.add_paragraph(COPYRIGHT_BOILERPLATE)
    document.add_page_break()

    # --- Table of contents ---
    if INCLUDE_TOC:
        toc_heading = document.add_paragraph()
        toc_heading.add_run("Contents").bold = True
        toc_heading.style = document.styles["Heading 1"]
        _add_static_toc(document, outline["chapters"])
        document.add_page_break()

    # --- Chapters ---
    for chapter in outline["chapters"]:
        ch_num = chapter["chapter_number"]
        path = os.path.join(book_dir, f"chapter_{ch_num:02d}.txt")
        if not os.path.isfile(path):
            print(f"[WARN] Missing {path}, skipping in docx manuscript.")
            continue

        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.add_run(f"Chapter {ch_num}")

        subheading = document.add_paragraph()
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subheading.add_run(chapter["title"])
        sub_run.italic = True

        document.add_paragraph()

        body = _read_chapter_body(path)
        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            p = document.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Inches(0) if i == 0 else Inches(0.3)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)

        if ch_num != outline["chapters"][-1]["chapter_number"]:
            document.add_page_break()

    output_path = os.path.join(book_dir, "manuscript.docx")
    document.save(output_path)
    return output_path


# Stage C is the last thing in this script that touches the GPU. Hand it back
# clean: AUTOMATIC1111 has no idle timeout of its own, so without this the
# checkpoint stays resident and the editor stage that follows (via
# pipeline_chain.py, the dashboard, or a manual run) gets pushed onto the CPU.
# See unload_a1111_checkpoint()'s docstring for the full failure chain.
unload_a1111_checkpoint()
wait_for_free_vram(target_mb=6000, max_wait_seconds=45)

print("\n--- Stage D: Assembling manuscript (raw draft + formatted docx) ---")
_stage_d_t0 = time.time()
_emit_timing("start", "manuscript_build")
raw_manuscript_path = build_raw_manuscript(book_dir, outline)
docx_manuscript_path = build_docx_manuscript(book_dir, outline)
print(f"[Saved] {raw_manuscript_path}")
print(f"[Saved] {docx_manuscript_path}")
_emit_timing("end", "manuscript_build", elapsed=f"{time.time() - _stage_d_t0:.1f}")

# =====================================================================
# 8. SUMMARY
# =====================================================================

total_words = sum(c["target_words"] for c in outline["chapters"])
print("\n=======================================================")
print("RUN COMPLETE")
print("=======================================================")
print(f"Title:     {outline['title']}")
if outline.get("subtitle"):
    print(f"Subtitle:  {outline['subtitle']}")
print(f"Genre:     {outline['genre']}")
print(f"Pen name:  {AUTHOR_NAME}")
print(f"Chapters:  {len(outline['chapters'])}")
print(f"Target words (approx): {total_words}")
print(f"Output folder: {book_dir}")
print(f"Raw manuscript: {raw_manuscript_path}")
print(f"Formatted docx: {docx_manuscript_path}")
print(f"Cover step result: {cover_result}")
if chapters_needing_regeneration:
    ch_list = ", ".join(f"#{n} ('{t}')" for n, t in chapters_needing_regeneration)
    print(
        f"\n[!] {len(chapters_needing_regeneration)} chapter(s) need manual regeneration "
        f"(model refused {MAX_CHAPTER_ATTEMPTS}/{MAX_CHAPTER_ATTEMPTS} times): {ch_list}"
    )
    print("    These are marked with a placeholder in the manuscript — see the [WARN] lines above.")

if chapters_degenerated:
    print(f"\n[!] {len(chapters_degenerated)} chapter(s) hit the repetition-collapse guard:")
    for n, t, outcome in chapters_degenerated:
        print(f"    #{n} ('{t}') — {outcome}")
    print("    The looping text was removed, so these chapters are coherent but shorter than")
    print("    their outline target. scoring_agent.py will also flag this book. Worth rerolling")
    print("    those chapters, or lowering their target_words — an over-long target is what")
    print("    triggers the spiral in the first place.")
