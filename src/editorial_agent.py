import sys

# Force UTF-8 stdout/stderr before anything else runs — same reasoning as
# local-book-generator.py: on Windows the default console codepage (cp1252) can't
# encode unicode CrewAI's console output uses, and this must happen before crewai
# is imported since some of its output can fire on import.
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", line_buffering=True)
    sys.stderr.reconfigure(encoding="utf-8", line_buffering=True)

import argparse
import glob
import json
import os
import re
import shutil
import subprocess
import time
from collections import Counter
from datetime import datetime

import requests  # added 2026-08-13 for the A1111/Ollama VRAM unload calls below

from crewai import Agent, Crew, Process, Task, LLM
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    sys.stdout.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
except AttributeError:
    pass  # Python < 3.7 fallback, not expected here

# =====================================================================
# 0. CONFIG
# =====================================================================

OUTPUT_ROOT = "output_books"

# Same base model as outline generation in local-book-generator.py. Editing is
# closer to structured analysis/judgment than open-ended creative generation,
# so the stock instruct model llama3.1 is a better fit here than the
# uncensored storytelling model that writes chapters — swap this if you want
# to compare.
#
# Points at a custom Ollama tag (not the base pull) with num_ctx raised to
# 16384 instead of Ollama's 2048-token default — see Modelfile.llama31-16k
# and project plan section 5. The editor reads full chapter text as INPUT
# (up to ~8,000 words on a long chapter), which was silently getting
# truncated at 2048 tokens before this fix. Requires a one-time
# `ollama create llama3.1-16k -f Modelfile.llama31-16k` on this machine
# before a run will find it — the preflight check below will say so clearly
# if it's missing.
EDITOR_MODEL_NAME = "llama3.1-16k"

# timeout added 2026-08-13. Without it, a single stalled Ollama request blocks
# forever with no output — which is exactly what "hung on the second editorial
# pass" looked like from the dashboard: the job still reads as running, the ETA
# ticks down past zero, and nothing ever fails or returns. A hard ceiling turns
# that indefinite hang into an ordinary failed attempt, which run_json_pass()
# and revise_chapter() already know how to retry and then degrade past.
#
# 900s (15 min) is deliberately generous — a healthy GPU-resident pass on a
# long chapter finishes in 60-90s, so anything approaching this ceiling is
# already a pathological run, not a slow-but-working one.
EDITOR_REQUEST_TIMEOUT_SECONDS = 900

editor_llm = LLM(
    model=f"ollama/{EDITOR_MODEL_NAME}",
    base_url="http://localhost:11434",
    timeout=EDITOR_REQUEST_TIMEOUT_SECONDS,
)


def check_ollama_models_available(
    required_models: list, ollama_url: str = "http://localhost:11434", fix_hints: dict = None
) -> None:
    """Preflight check, run once at startup before any CrewAI work begins.
    Identical to local-book-generator.py's check_ollama_models_available() — see
    that copy for the full reasoning (a missing model otherwise surfaces as a
    wall of CrewAI retry-panel noise ending in a raw traceback instead of one
    clear message, confirmed 2026-08-12).

    fix_hints: optional {model_name: fix_command} overrides for models not
    fixable with a plain `ollama pull` — added 2026-08-12 for the custom
    -16k context-window tag (see Modelfile.llama31-16k), which is built
    locally with `ollama create`, not pulled from a registry."""
    import urllib.request as _urllib_request
    import json as _json

    fix_hints = fix_hints or {}

    try:
        with _urllib_request.urlopen(f"{ollama_url.rstrip('/')}/api/tags", timeout=5) as resp:
            data = _json.loads(resp.read())
    except Exception as e:
        print(f"[FATAL] Could not reach Ollama at {ollama_url} to verify required models: {e}")
        print("        Is Ollama running? Start it (e.g. 'ollama serve') and try again.")
        sys.exit(1)

    available = {m.get("name", "") for m in data.get("models", [])}
    available_bare = {name.split(":")[0] for name in available if name}

    missing = [
        m for m in required_models
        if m not in available and m.split(":")[0] not in available_bare
    ]
    if missing:
        print(f"[FATAL] Required Ollama model(s) not found on this machine: {', '.join(missing)}")
        for m in missing:
            print(f"        Fix: {fix_hints.get(m, f'ollama pull {m}')}")
        print(f"        Currently pulled models: {', '.join(sorted(available)) or '(none)'}")
        sys.exit(1)

    print(f"[Config] Ollama models verified present: {', '.join(required_models)}")


check_ollama_models_available(
    [EDITOR_MODEL_NAME],
    fix_hints={
        EDITOR_MODEL_NAME: (
            "one-time local setup, not a registry pull: "
            "ollama create llama3.1-16k -f Modelfile.llama31-16k "
            "(requires 'ollama pull llama3.1' first if you haven't already)"
        ),
    },
)

editor = Agent(
    role="Senior Fiction Editor",
    goal=(
        "Review genre fiction manuscripts the way a senior editor at a publishing "
        "house would — catching continuity errors, pacing problems, weak prose, "
        "and mechanical mistakes — and produce specific, actionable fixes."
    ),
    backstory=(
        "You have edited hundreds of commercial genre novels across thriller, "
        "fantasy, mystery, horror, and romance. You are precise and concrete: "
        "every issue you flag quotes the exact text it applies to and states "
        "exactly what should change. You never invent problems that aren't there "
        "and never pad your notes with vague generalities."
    ),
    llm=editor_llm,
    verbose=True,
)

# Manuscript-formatting constants, kept identical to local-book-generator.py's Stage D
# (BODY_FONT_NAME / BODY_FONT_SIZE_PT / INCLUDE_TOC / COPYRIGHT_BOILERPLATE) so the
# auto-apply .docx built below matches the raw one exactly. Duplicated rather than
# imported — the two scripts are meant to run standalone — so if you ever change
# these in local-book-generator.py, update them here too to keep the two in sync.
BODY_FONT_NAME = "Garamond"
BODY_FONT_SIZE_PT = 12
INCLUDE_TOC = True
COPYRIGHT_BOILERPLATE = (
    "This is a work of fiction. Names, characters, places, and incidents either are "
    "the product of the author's imagination or are used fictitiously. Any "
    "resemblance to actual events, locales, or persons, living or dead, is entirely "
    "coincidental."
)

MAX_REVIEW_ATTEMPTS = 3  # JSON parse retries per pass, same pattern as outline parsing

# A short reminder this is fiction, same rationale as local-book-generator.py's
# FICTION_FRAMING — free, harmless, stacks with retries if it ever matters here too.
FICTION_FRAMING = (
    "This is entirely fictional, original genre fiction being edited for a "
    "published novel — not real-world advice, instructions, or commentary.\n\n"
)

# =====================================================================
# 1. CLI
# =====================================================================

parser = argparse.ArgumentParser(description="Editorial pass over a local-book-generator.py output.")
parser.add_argument(
    "--book-dir",
    default=None,
    help="Path to a specific output_books/<slug-timestamp> folder. Defaults to the most "
         "recently modified book folder under output_books/.",
)
parser.add_argument(
    "--auto-apply",
    action="store_true",
    help="Apply the editor's fixes and write revised chapters to <book-dir>/edited/. "
         "Without this flag, the editor only writes a review report — the default, "
         "so you can build trust in its judgment before letting it rewrite chapters.",
)
parser.add_argument(
    "--separate-passes",
    action="store_true",
    help="Revert to the original three separate review calls per chapter (macro, then style, "
         "then micro) instead of the single combined review pass. Roughly 3x the review time "
         "per chapter, since each pass re-sends the whole chapter as input — kept as an escape "
         "hatch in case the combined pass ever misses something you care about.",
)
args = parser.parse_args()

AUTO_APPLY = args.auto_apply
SEPARATE_PASSES = args.separate_passes


def find_latest_book_dir(output_root: str) -> "str | None":
    """Most recently modified subfolder of output_root that actually looks like a
    book (has an outline.json) — filters out stray folders so a typo'd/incomplete
    directory doesn't get silently picked."""
    if not os.path.isdir(output_root):
        return None
    candidates = [
        os.path.join(output_root, name)
        for name in os.listdir(output_root)
        if os.path.isdir(os.path.join(output_root, name))
        and os.path.isfile(os.path.join(output_root, name, "outline.json"))
    ]
    if not candidates:
        return None
    candidates.sort(key=os.path.getmtime, reverse=True)
    return candidates[0]


BOOK_DIR = args.book_dir or find_latest_book_dir(OUTPUT_ROOT)
if not BOOK_DIR or not os.path.isdir(BOOK_DIR) or not os.path.isfile(os.path.join(BOOK_DIR, "outline.json")):
    print(
        f"[ERROR] Could not find a book folder to edit. Looked in '{OUTPUT_ROOT}' for the most "
        f"recently modified subfolder containing an outline.json, or pass one explicitly with "
        f"--book-dir \"path\\to\\output_books\\your-book-folder\"."
    )
    sys.exit(1)

print(f"[Config] Editing book at: {BOOK_DIR}")
print(f"[Config] Editor model: {EDITOR_MODEL_NAME}")
print(f"[Config] Mode: {'AUTO-APPLY (writes revised chapters to edited/)' if AUTO_APPLY else 'SUGGEST-ONLY (report only, no files changed)'}")
print(f"[Config] Review passes: {'3 separate (macro/style/micro)' if SEPARATE_PASSES else '1 combined'}")

# =====================================================================
# 2. VRAM UTILITIES (same pattern as local-book-generator.py)
# =====================================================================


def get_free_vram_mb() -> "int | None":
    """Query free VRAM in MB via nvidia-smi. Returns None (not an exception) if
    nvidia-smi isn't available, so callers can treat 'unknown' as 'proceed
    cautiously' instead of crashing on machines without it on PATH."""
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=memory.free", "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode == 0 and result.stdout.strip():
            return int(result.stdout.strip().splitlines()[0])
    except (subprocess.SubprocessError, FileNotFoundError, ValueError, OSError):
        pass
    return None


def unload_a1111_checkpoint(a1111_url: str = "http://127.0.0.1:7860",
                            max_wait_seconds: int = 25) -> None:
    """Free VRAM held by AUTOMATIC1111 before this script loads its own model.

    Corrects a wrong assumption in the note that used to live here, which said
    freeing A1111 was "a dashboard-level concern, not something this standalone
    script manages." That reasoning held that because this script only talks to
    Ollama, it has no handoff to manage — but the handoff that matters is the
    *inbound* one. Cover generation leaves SDXL (~6GB) resident with no idle
    timeout, so on an 8GB card llama3.1-16k (~4.7GB weights + ~2GB of 16k KV
    cache) could not fit, and Ollama silently offloaded it to the CPU instead
    of failing. Measured cost of that: 274s/chapter here vs. 165s/chapter for
    the writer stage doing strictly more generation (job_timing_history.json),
    plus the CPU-pinned ollama.exe and the mid-run hangs seen 2026-08-13.

    Safe to call unconditionally: if A1111 isn't running (the normal case for a
    genuinely standalone run) this is a no-op. A1111 reloads its checkpoint by
    itself on the next txt2img request."""
    print("\n[VRAM] Unloading AUTOMATIC1111's checkpoint (if loaded) before the editor model...")
    before = get_free_vram_mb()
    try:
        response = requests.post(f"{a1111_url.rstrip('/')}/sdapi/v1/unload-checkpoint", timeout=60)
        if response.status_code not in (200, 204):
            print(f"[WARN] A1111 unload-checkpoint returned HTTP {response.status_code}. "
                  f"Continuing anyway.")
    except requests.exceptions.RequestException as e:
        print(f"[VRAM] AUTOMATIC1111 not reachable ({e.__class__.__name__}) — nothing to unload.")
        return

    if before is None:
        time.sleep(3)
        return
    for _ in range(max_wait_seconds):
        time.sleep(1)
        now = get_free_vram_mb()
        if now is not None and now >= before + 500:
            print(f"[VRAM] Freed up to {now}MB available (was {before}MB).")
            return
    print(f"[WARN] VRAM didn't visibly increase after {max_wait_seconds}s "
          f"(still ~{before}MB free) — continuing anyway.")


def unload_ollama_model(model_name: str, max_wait_seconds: int = 20) -> None:
    """Drop a model from VRAM immediately (keep_alive: 0). Same helper as
    local-book-generator.py's — duplicated per this codebase's "each script
    runs standalone" convention. Called on the way OUT of this script so the
    scorer stage that follows doesn't have to wait out Ollama's ~5 minute idle
    timeout (or fight the editor model for room) before loading its own."""
    print(f"\n[VRAM] Unloading Ollama model '{model_name}'...")
    try:
        requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model_name, "keep_alive": 0},
            timeout=15,
        )
    except requests.exceptions.RequestException as e:
        print(f"[WARN] Could not reach Ollama to unload the model ({e}). Continuing anyway.")
        return
    for _ in range(max_wait_seconds):
        time.sleep(1)
        if (get_free_vram_mb() or 0) >= 4000:
            return


_free_vram = get_free_vram_mb()
if _free_vram is not None:
    print(f"[VRAM] Free before starting: {_free_vram}MB")

# Inbound handoff: reclaim VRAM from cover generation before loading the editor
# model, so it runs on the GPU instead of silently falling back to the CPU.
unload_a1111_checkpoint()

_free_vram = get_free_vram_mb()
if _free_vram is not None:
    print(f"[VRAM] Free after A1111 unload: {_free_vram}MB")
    if _free_vram < 6000:
        print(f"[WARN] Only {_free_vram}MB VRAM free. '{EDITOR_MODEL_NAME}' needs roughly "
              f"6.5GB (weights + 16k KV cache) to stay fully on the GPU. Ollama will not "
              f"error if it doesn't fit — it will quietly offload layers to system RAM and "
              f"run them on the CPU, which is 2-3x slower and is what caused the stalls on "
              f"2026-08-13. Check for a leftover A1111 or ollama process holding VRAM.")

# =====================================================================
# 3. JSON PARSING HELPERS (same retry/repair pattern as parse_outline)
# =====================================================================


def _repair_common_json_mistakes(text: str) -> str:
    """Local models occasionally leave a trailing comma before a closing brace/
    bracket, which breaks json.loads even though the rest of the document is
    fine. Strip those before giving up."""
    return re.sub(r",(\s*[}\]])", r"\1", text)


def _parse_json_response(raw_text: str) -> dict:
    """Pull the JSON object out of the model's response, tolerating stray text/
    fences, same approach as local-book-generator.py's parse_outline()."""
    match = re.search(r"\{.*\}", raw_text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in response.")
    json_text = match.group(0)
    try:
        return json.loads(json_text)
    except json.JSONDecodeError:
        return json.loads(_repair_common_json_mistakes(json_text))


def run_json_pass(description: str, expected_output: str, required_keys=("issues",)) -> dict:
    """Run one editor task expecting a JSON response, retrying up to
    MAX_REVIEW_ATTEMPTS times on parse failure. Degrades gracefully on total
    failure — returns an empty-issues result with _parse_failed=True instead of
    crashing the whole chapter's review over one bad response, since a missed
    pass on one chapter shouldn't take down the rest of the run."""
    last_error = None
    for attempt in range(1, MAX_REVIEW_ATTEMPTS + 1):
        task = Task(description=description, expected_output=expected_output, agent=editor)
        crew = Crew(agents=[editor], tasks=[task], process=Process.sequential, verbose=True)
        raw = str(crew.kickoff())
        try:
            data = _parse_json_response(raw)
            if not all(k in data for k in required_keys):
                raise ValueError(f"Response missing required key(s) {required_keys}.")
            data.setdefault("issues", [])
            if not isinstance(data["issues"], list):
                raise ValueError("'issues' was not a list.")
            return data
        except (ValueError, json.JSONDecodeError) as e:
            last_error = e
            print(f"[WARN] Editor pass attempt {attempt}/{MAX_REVIEW_ATTEMPTS} failed to parse ({e}).")
            if attempt < MAX_REVIEW_ATTEMPTS:
                print("[WARN] Retrying...")
    print(f"[WARN] All {MAX_REVIEW_ATTEMPTS} attempts failed to parse a valid response ({last_error}). "
          f"Skipping this pass — no issues recorded from it.")
    return {"issues": [], "_parse_failed": True}


# =====================================================================
# 4. REFUSAL DETECTION (same as local-book-generator.py — see that file's
#    comments for the full rationale; duplicated here since these are
#    intentionally separate standalone scripts)
# =====================================================================

REFUSAL_PATTERN = re.compile(
    r"^\s*(I cannot|I can[’']t|I can not|I[’']m sorry|I am sorry|"
    r"I[’']m unable|I am unable|As an AI|I won[’']t|"
    r"I[’']m not able|I don[’']t (feel comfortable|think I))",
    re.IGNORECASE,
)

META_INSTRUCTION_PATTERN = re.compile(
    r"(here's an outline|here is an outline|outline that can serve as a guide|"
    r"you can continue this story|add your own creative elements|"
    r"let me know if (i can help|you (need|would like)|there's anything else)|"
    r"i can (guide|help) you (on |through )?how to|"
    r"keeping to the word count target|feel free to (expand|continue)|"
    r"would you like me to (continue|write)|as a starting point for)",
    re.IGNORECASE,
)

OUTLINE_STRUCTURE_PATTERN = re.compile(r"^\s*(\*\*)?[IVX]{1,4}\.\s+", re.MULTILINE)


def looks_like_refusal(text: str) -> bool:
    stripped = text.strip()
    if not REFUSAL_PATTERN.match(stripped):
        return False
    if len(stripped) < 400:
        return True
    return bool(META_INSTRUCTION_PATTERN.search(stripped) or OUTLINE_STRUCTURE_PATTERN.search(stripped))


# =====================================================================
# 5. LOADING THE BOOK: outline, chapters, story bible / style sheet
# =====================================================================


def load_outline(book_dir: str) -> dict:
    with open(os.path.join(book_dir, "outline.json"), "r", encoding="utf-8") as f:
        return json.load(f)


def parse_main_characters(main_characters: list) -> list:
    """outline.json stores main_characters as ["Name: description", ...] strings —
    split each into a {"name", "description"} dict, tolerating entries with no
    colon (kept whole as the name, empty description) rather than raising."""
    parsed = []
    for entry in main_characters or []:
        if ":" in entry:
            name, desc = entry.split(":", 1)
            parsed.append({"name": name.strip(), "description": desc.strip()})
        else:
            parsed.append({"name": entry.strip(), "description": ""})
    return parsed


CHAPTER_FILE_PATTERN = re.compile(r"chapter_(\d+)\.txt$")


def discover_chapters(book_dir: str) -> list:
    """Find chapter_NN.txt files, sorted numerically by chapter number (not
    lexicographically — safe either way at 2-digit padding, but numeric sorting
    doesn't silently break if that ever changes)."""
    found = []
    for path in glob.glob(os.path.join(book_dir, "chapter_*.txt")):
        match = CHAPTER_FILE_PATTERN.search(os.path.basename(path))
        if match:
            found.append((int(match.group(1)), path))
    found.sort(key=lambda pair: pair[0])
    return found


def split_chapter_file(path: str) -> tuple:
    """Chapter files are written as '{title}\\n\\n{body}' by local-book-generator.py
    — split back into (title, body) so the editor reviews prose only, not the
    title line repeated as if it were the first sentence."""
    with open(path, "r", encoding="utf-8") as f:
        raw = f.read()
    parts = raw.split("\n\n", 1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1]
    return "", raw


STYLE_SHEET_FILENAME = "style_sheet.json"


def load_or_seed_style_sheet(book_dir: str, outline: dict) -> dict:
    """The 'story bible' the editor carries forward chapter to chapter. Seeded
    from the outline on first run; on later runs (e.g. re-editing after a manual
    fix) the existing file is reused so accumulated facts/terms aren't lost."""
    path = os.path.join(book_dir, STYLE_SHEET_FILENAME)
    if os.path.isfile(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "title": outline.get("title", ""),
        "genre": outline.get("genre", ""),
        "premise": outline.get("premise", ""),
        "characters": parse_main_characters(outline.get("main_characters", [])),
        "established_facts": [],
        "style_terms": [],
    }


def save_style_sheet(book_dir: str, style_sheet: dict) -> None:
    path = os.path.join(book_dir, STYLE_SHEET_FILENAME)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(style_sheet, f, indent=2)


# =====================================================================
# 6. EBOOK TYPOGRAPHY FIXES (deterministic — not left to the local model's
#    character-level precision, which isn't reliable at this scale)
# =====================================================================


def apply_ebook_typography(text: str) -> tuple:
    """Mechanical ebook-formatting cleanup: smart/curly quotes, no more than one
    blank line between paragraphs, no trailing whitespace. Returns (fixed_text,
    notes) — notes is a list of human-readable strings describing what changed,
    for the report, even when this function is called in suggest-only mode
    (where fixed_text is computed but never written to disk)."""
    notes = []
    fixed = text

    straight_double = fixed.count('"')
    if straight_double:
        # Alternate opening/closing curly quotes for straight double-quotes.
        out, is_open = [], True
        for ch in fixed:
            if ch == '"':
                out.append("“" if is_open else "”")
                is_open = not is_open
            else:
                out.append(ch)
        fixed = "".join(out)
        notes.append(f"Converted {straight_double} straight double-quote(s) to curly quotes.")

    # Straight apostrophes -> right single quote, except when used as a leading
    # apostrophe for elision (e.g. 'em) — good enough for prose; not attempting
    # to distinguish opening single-quotes since genre fiction rarely needs them.
    straight_single = fixed.count("'")
    if straight_single:
        fixed = fixed.replace("'", "’")
        notes.append(f"Converted {straight_single} straight apostrophe(s)/single-quote(s) to curly.")

    collapsed = re.sub(r"\n{3,}", "\n\n", fixed)
    if collapsed != fixed:
        notes.append("Collapsed extra blank lines down to a single blank line between paragraphs.")
    fixed = collapsed

    stripped_lines = "\n".join(line.rstrip() for line in fixed.split("\n"))
    if stripped_lines != fixed:
        notes.append("Removed trailing whitespace from one or more lines.")
    fixed = stripped_lines

    return fixed, notes


# =====================================================================
# 7. EDITOR PASSES (macro / style / micro — matching the intended pass
#    structure) + the auto-apply revise pass
# =====================================================================


def macro_pass(chapter_text: str, ch_num: int, ch_title: str, ch_summary: str, style_sheet: dict) -> dict:
    facts = "\n".join(f"- {f}" for f in style_sheet["established_facts"]) or "(none yet)"
    characters = "\n".join(f"- {c['name']}: {c['description']}" for c in style_sheet["characters"]) or "(none listed)"
    description = (
        f"{FICTION_FRAMING}"
        f"You are reviewing Chapter {ch_num} ('{ch_title}') of the novel '{style_sheet['title']}' "
        f"({style_sheet['genre']}) for STRUCTURE AND CONTINUITY ONLY. Do not comment on grammar or "
        f"line-level wording — those are handled in separate passes.\n\n"
        f"Book premise: {style_sheet['premise']}\n\n"
        f"Main characters:\n{characters}\n\n"
        f"Established facts from earlier chapters (treat these as canon — flag anything in this "
        f"chapter that contradicts them):\n{facts}\n\n"
        f"This chapter's intended beat, from the outline: {ch_summary}\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        "Check for:\n"
        "- Continuity errors against the established facts/character list above (changed physical "
        "descriptions, a dead character reappearing, timeline slips).\n"
        "- Whether this chapter's events actually deliver on its intended beat.\n"
        "- Whether every scene serves a clear purpose — advances the plot and/or develops a "
        "character — flag scenes that do neither.\n"
        "- Pacing within this chapter (dragging, or rushing through something that needed more room).\n\n"
        "Respond with ONLY a valid JSON object (no commentary, no markdown fences) in exactly this "
        "shape:\n"
        "{\n"
        '  "issues": [\n'
        '    {"category": "continuity|pacing|scene_purpose", "severity": "low|medium|high", '
        '"excerpt": "short quoted snippet from the chapter", "issue": "what is wrong", '
        '"fix": "specific suggested fix"}\n'
        "  ],\n"
        '  "new_facts": ["short factual statements established in this chapter worth remembering, '
        'e.g. \'Kael has a scar on his left hand\'"]\n'
        "}\n"
        "If there are no issues, return an empty issues list. Keep new_facts to concrete, checkable "
        "facts only (names, physical descriptions, deaths, locations, dates/times) — not summary."
    )
    result = run_json_pass(
        description,
        expected_output="A JSON object with 'issues' and 'new_facts', matching the schema described.",
        required_keys=("issues",),
    )
    result.setdefault("new_facts", [])
    if not isinstance(result["new_facts"], list):
        result["new_facts"] = []
    return result


def style_pass(chapter_text: str, ch_num: int, ch_title: str, style_sheet: dict) -> dict:
    description = (
        f"{FICTION_FRAMING}"
        f"You are reviewing Chapter {ch_num} ('{ch_title}') of the novel '{style_sheet['title']}' "
        f"for PROSE STYLE ONLY. Do not comment on plot/continuity or grammar/punctuation — those "
        f"are handled in separate passes.\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        "Check for:\n"
        "- Point-of-view or narrative voice inconsistency.\n"
        "- Filter words that distance the reader (\"she noticed,\" \"he felt,\" \"she saw,\" "
        "\"he realized\") where direct description would be stronger.\n"
        "- Overused passive voice.\n"
        "- Words or phrases repeated too close together.\n"
        "- \"Telling\" instead of \"showing\" emotional states, or exposition dumped rather than "
        "dramatized.\n\n"
        "Respond with ONLY a valid JSON object (no commentary, no markdown fences) in exactly this "
        "shape:\n"
        "{\n"
        '  "issues": [\n'
        '    {"category": "pov_voice|filter_words|passive_voice|repetition|show_dont_tell", '
        '"severity": "low|medium|high", "excerpt": "short quoted snippet from the chapter", '
        '"issue": "what is wrong", "fix": "specific suggested fix"}\n'
        "  ]\n"
        "}\n"
        "If there are no issues, return an empty issues list. Only flag real, specific instances you "
        "can quote — not general impressions."
    )
    return run_json_pass(
        description,
        expected_output="A JSON object with 'issues', matching the schema described.",
        required_keys=("issues",),
    )


def micro_pass(chapter_text: str, ch_num: int, ch_title: str, style_sheet: dict) -> dict:
    terms = "\n".join(f"- {t}" for t in style_sheet["style_terms"]) or "(none yet)"
    description = (
        f"{FICTION_FRAMING}"
        f"You are reviewing Chapter {ch_num} ('{ch_title}') of the novel '{style_sheet['title']}' "
        f"for GRAMMAR, MECHANICS, AND STYLE-SHEET CONSISTENCY ONLY. Do not comment on plot or "
        f"prose style — those are handled in separate passes.\n\n"
        f"Project style sheet (recurring terms that must stay consistent — flag any deviation from "
        f"these):\n{terms}\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        "Check for:\n"
        "- Grammar and syntax errors: dangling modifiers, tense slips, faulty parallelism.\n"
        "- Dialogue mechanics: tag placement, action beats, comma/quote punctuation.\n"
        "- Deviations from the project style sheet above.\n"
        "- New recurring terms in this chapter (character/place names, invented jargon, hyphenated "
        "compounds) worth adding to the style sheet so future chapters stay consistent with them.\n\n"
        "Respond with ONLY a valid JSON object (no commentary, no markdown fences) in exactly this "
        "shape:\n"
        "{\n"
        '  "issues": [\n'
        '    {"category": "grammar|dialogue_mechanics|style_sheet", "severity": "low|medium|high", '
        '"excerpt": "short quoted snippet from the chapter", "issue": "what is wrong", '
        '"fix": "specific suggested fix"}\n'
        "  ],\n"
        '  "new_style_terms": ["Term: how it should always be styled/capitalized/hyphenated"]\n'
        "}\n"
        "If there are no issues, return an empty issues list."
    )
    result = run_json_pass(
        description,
        expected_output="A JSON object with 'issues' and 'new_style_terms', matching the schema described.",
        required_keys=("issues",),
    )
    result.setdefault("new_style_terms", [])
    if not isinstance(result["new_style_terms"], list):
        result["new_style_terms"] = []
    return result


def combined_review_pass(chapter_text: str, ch_num: int, ch_title: str, ch_summary: str,
                         style_sheet: dict) -> dict:
    """One review call covering all three of macro_pass/style_pass/micro_pass's
    territory, added 2026-08-13 as the main runtime optimisation.

    Why this is worth doing: the three separate passes each resend the ENTIRE
    chapter (4-6k words, ~6-8k tokens) as input. On a local 8B model the
    prompt-evaluation of that text dominates the cost of the call — the actual
    JSON that comes back is a few hundred tokens. So three passes cost roughly
    three full prompt evaluations to produce three short outputs, and merging
    them into one call removes two of those evaluations outright rather than
    trading quality for speed. Combined with the VRAM fix above, this is what
    takes the editorial stage from ~274s/chapter to a projected ~70-90s.

    The tradeoff being accepted: the separate passes' "review for X ONLY, the
    others are handled elsewhere" framing does keep the model narrowly focused,
    and a single prompt asking for all three at once will find somewhat fewer
    total issues. That is an acceptable trade here for two reasons — the issues
    most likely to be dropped are the low-severity style nits that contribute
    least to a revision, and the review output feeds revise_chapter(), which
    works better from a shorter list of real problems than a long list padded
    with marginal ones.

    Returns the same merged shape the three separate passes produced between
    them ({"issues": [...], "new_facts": [...], "new_style_terms": [...]}), so
    every downstream consumer — revise_chapter(), the report writer, and
    scoring_agent.py's parsing of editorial_review.txt — is unaffected.

    Pass --separate-passes to fall back to the original three-call behaviour."""
    facts = "\n".join(f"- {f}" for f in style_sheet["established_facts"]) or "(none yet)"
    characters = "\n".join(f"- {c['name']}: {c['description']}" for c in style_sheet["characters"]) or "(none listed)"
    terms = "\n".join(f"- {t}" for t in style_sheet["style_terms"]) or "(none yet)"
    description = (
        f"{FICTION_FRAMING}"
        f"You are the editor on Chapter {ch_num} ('{ch_title}') of the novel "
        f"'{style_sheet['title']}' ({style_sheet['genre']}). Review it on three levels in a "
        f"single pass: STRUCTURE/CONTINUITY, PROSE STYLE, and GRAMMAR/MECHANICS.\n\n"
        f"Book premise: {style_sheet['premise']}\n\n"
        f"Main characters:\n{characters}\n\n"
        f"Established facts from earlier chapters (treat these as canon — flag anything in this "
        f"chapter that contradicts them):\n{facts}\n\n"
        f"Project style sheet (recurring terms that must stay consistent — flag any "
        f"deviation):\n{terms}\n\n"
        f"This chapter's intended beat, from the outline: {ch_summary}\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        "Check for, in one combined list:\n"
        "STRUCTURE & CONTINUITY\n"
        "- Continuity errors against the established facts/character list above (changed physical "
        "descriptions, a dead character reappearing, timeline slips).\n"
        "- Whether this chapter's events actually deliver on its intended beat.\n"
        "- Scenes that neither advance the plot nor develop a character.\n"
        "- Pacing within this chapter (dragging, or rushing something that needed more room).\n"
        "PROSE STYLE\n"
        "- Point-of-view or narrative voice inconsistency.\n"
        "- Filter words that distance the reader (\"she noticed,\" \"he felt,\" \"she saw,\" "
        "\"he realized\") where direct description would be stronger.\n"
        "- Overused passive voice.\n"
        "- Words or phrases repeated too close together.\n"
        "- \"Telling\" instead of \"showing\" emotional states, or exposition dumped rather than "
        "dramatized.\n"
        "GRAMMAR & MECHANICS\n"
        "- Grammar and syntax errors: dangling modifiers, tense slips, faulty parallelism.\n"
        "- Dialogue mechanics: tag placement, action beats, comma/quote punctuation.\n"
        "- Deviations from the project style sheet above.\n\n"
        "Respond with ONLY a valid JSON object (no commentary, no markdown fences) in exactly "
        "this shape:\n"
        "{\n"
        '  "issues": [\n'
        '    {"category": "continuity|pacing|scene_purpose|pov_voice|filter_words|passive_voice|'
        'repetition|show_dont_tell|grammar|dialogue_mechanics|style_sheet", '
        '"severity": "low|medium|high", "excerpt": "short quoted snippet from the chapter", '
        '"issue": "what is wrong", "fix": "specific suggested fix"}\n'
        "  ],\n"
        '  "new_facts": ["short factual statements established in this chapter worth remembering, '
        'e.g. \'Kael has a scar on his left hand\'"],\n'
        '  "new_style_terms": ["Term: how it should always be styled/capitalized/hyphenated"]\n'
        "}\n"
        "The \"category\" value must be exactly one of the eleven strings listed above — do not "
        "invent new category names. Only flag real, specific instances you can quote — not "
        "general impressions. If there are no issues, return an empty issues list. Keep "
        "new_facts to concrete, checkable facts only (names, physical descriptions, deaths, "
        "locations, dates/times) — not summary."
    )
    result = run_json_pass(
        description,
        expected_output="A JSON object with 'issues', 'new_facts' and 'new_style_terms', "
                        "matching the schema described.",
        required_keys=("issues",),
    )
    for key in ("new_facts", "new_style_terms"):
        result.setdefault(key, [])
        if not isinstance(result[key], list):
            result[key] = []
    return result


# =====================================================================
# DEGENERATION DETECTION — added 2026-08-13
# =====================================================================
# Catches the repetition spiral by measuring the text statistically. The LLM
# review pass cannot see it — on "Rainy Night, Lonely Soul" chapter 9 it
# reported zero issues on a passage repeating "murder" twenty times — because
# locally every sentence still reads like prose. Duplicated from
# local-book-generator.py (same "each script runs standalone" convention as the rest
# of this codebase) so the scorer can refuse to publish-rate a broken book
# even if it was written before this fix existed.
#
# Two signals, measured over a sliding 400-word window:
#
#   type-token ratio  — unique words / total words in the window. Collapses
#                       when the model starts cycling the same vocabulary.
#   runaway content word — the most frequent NON-function word's share of the
#                       window. Function words ("the", "a", "and") are
#                       legitimately 4-6% of any healthy English text, so
#                       they're excluded; a content word above 5% means the
#                       model is stuck on it.
#
# Thresholds were calibrated against 19 real chapters from this pipeline's own
# output_books/ — 5 known-bad (Rashomon) and 14 known-good (Murder at
# Willowbrook, Golden Heirloom, Aurora Initiative, Raining Ashes, Stardust
# Rebellion, Echoes of Eternity). Worst-window TTR measured 0.245-0.34 on the
# bad chapters and 0.365-0.505 on the good ones, so the 0.36 floor sits in a
# clean gap: it caught all 5 bad chapters and flagged not one window in any of
# the 14 good ones. Requiring two CONSECUTIVE bad windows keeps a single
# dialogue-heavy passage (naturally low TTR — short lines, repeated names)
# from tripping it.
DEGEN_WINDOW_WORDS = 400
DEGEN_STEP_WORDS = 200
DEGEN_TTR_FLOOR = 0.36
DEGEN_WORD_CEILING = 0.05
DEGEN_MIN_CONSECUTIVE = 2

# Streaming early-abort tuning (see generate_chapter_streaming). The tail slice
# is 1000 words so it spans enough windows for the 2-consecutive rule to be
# meaningful; checks start at 1200 words because a collapse before then is
# vanishingly rare and the first few hundred words of any chapter are naturally
# repetitive (establishing names, place, weather).
DEGEN_STREAM_TAIL_WORDS = 1000
DEGEN_STREAM_FIRST_CHECK_WORDS = 1200
DEGEN_STREAM_CHECK_EVERY_WORDS = 300

_DEGEN_WORD_RE = re.compile(r"[A-Za-z']+")
_DEGEN_STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "if", "of", "to", "in", "on", "at", "for", "with",
    "from", "by", "as", "is", "was", "were", "be", "been", "are", "am", "he", "she", "it",
    "they", "them", "his", "her", "him", "their", "its", "i", "you", "we", "me", "my", "that",
    "this", "these", "those", "not", "no", "had", "has", "have", "did", "do", "does", "would",
    "could", "should", "will", "can", "there", "then", "so", "up", "out", "down", "into",
    "over", "about", "just", "like", "what", "when", "who", "which", "all", "one", "more",
    "than", "too", "very", "said", "being", "own", "only", "other", "new",
}


def detect_degeneration(text: str) -> dict:
    """Return a verdict on whether `text` collapses into repetition.

    Keys: degenerate (bool), first_bad_word_index (int|None — where the
    collapse starts, for salvaging the healthy prefix), worst_ttr,
    runaway_word, total_words, longest_run, detail (human-readable)."""
    from collections import Counter as _Counter

    words = _DEGEN_WORD_RE.findall(text.lower())
    total = len(words)
    if total < DEGEN_WINDOW_WORDS:
        # Too short to measure meaningfully — don't guess. A chapter this short
        # has other problems the length check will catch anyway.
        return {"degenerate": False, "first_bad_word_index": None, "worst_ttr": None,
                "runaway_word": None, "total_words": total, "longest_run": 0,
                "detail": f"only {total} words — too short to assess"}

    # Scan on SEVERAL window phases and keep the worst verdict, rather than one
    # grid starting at word 0.
    #
    # Why (found 2026-08-14, "Rainy Night, Lonely Soul"): a single grid makes the
    # result phase-dependent, and on a chapter sitting right at the boundary that
    # decides the verdict. Chapter 9 of that book tested CLEAN as the writer saw
    # it (the chapter body, longest_run 1) and DEGENERATE as the scorer saw it
    # (the same body with a 6-word title line prepended, longest_run 2). Shifting
    # the start by 4 words flipped it. The chapter was in fact badly collapsed —
    # 20 repetitions of "murder" in one stretch — so the writer's miss let it
    # through to the editor, and only the scorer caught it.
    #
    # Four phases at quarter-step intervals means a bad stretch is measured
    # wherever it happens to begin, so the same text gets the same answer no
    # matter what precedes it.
    phases = [0, DEGEN_STEP_WORDS // 4, DEGEN_STEP_WORDS // 2, 3 * DEGEN_STEP_WORDS // 4]
    worst_ttr = 1.0
    runaway = None
    longest_run = 0
    first_bad = None
    for phase in phases:
        bad_offsets = []
        for i in range(phase, total - DEGEN_WINDOW_WORDS + 1, DEGEN_STEP_WORDS):
            window = words[i:i + DEGEN_WINDOW_WORDS]
            ttr = len(set(window)) / len(window)
            worst_ttr = min(worst_ttr, ttr)
            content = [w for w in window if w not in _DEGEN_STOPWORDS]
            share, top = 0.0, ""
            if content:
                top, count = _Counter(content).most_common(1)[0]
                share = count / len(window)
            if ttr < DEGEN_TTR_FLOOR or share > DEGEN_WORD_CEILING:
                bad_offsets.append(i)
                if share > DEGEN_WORD_CEILING and (runaway is None or share > runaway[1]):
                    runaway = (top, round(share, 4))
        run, prev = 0, None
        for i in bad_offsets:
            if prev is not None and i - prev == DEGEN_STEP_WORDS:
                run += 1
            else:
                run = 1
            longest_run = max(longest_run, run)
            prev = i
        if bad_offsets and (first_bad is None or bad_offsets[0] < first_bad):
            first_bad = bad_offsets[0]

    degenerate = longest_run >= DEGEN_MIN_CONSECUTIVE

    # first_bad_word_index is the FIRST bad window anywhere in the text, not the
    # start of the longest bad run — and that distinction matters for salvage.
    # Truncating at the longest run's start looked right but left earlier bad
    # patches in place: on Rashomon chapter 2 it produced a 2,596-word "salvaged"
    # prefix that was still degenerate, because that chapter was already
    # collapsing well before its worst stretch. Cutting at the first bad window
    # instead makes the surviving prefix clean by construction — every window
    # entirely inside it was measured and passed.
    detail = ""
    if degenerate:
        detail = (f"repetition collapse from around word {first_bad} of {total} "
                  f"(lowest vocabulary richness {round(worst_ttr, 3)}, floor {DEGEN_TTR_FLOOR})")
        if runaway:
            detail += f"; the word '{runaway[0]}' reaches {runaway[1] * 100:.1f}% of a 400-word stretch"
    return {"degenerate": degenerate, "first_bad_word_index": first_bad,
            "worst_ttr": round(worst_ttr, 3), "runaway_word": runaway, "total_words": total,
            "longest_run": longest_run, "detail": detail}


# =====================================================================
# REVISION ACCEPTANCE + DIRECT OLLAMA CALL (2026-08-14)
# =====================================================================
_WORD_RE = re.compile(r"[A-Za-z']+")
# Both quote styles — the writer model emits curly quotes, but typography
# fixes and hand edits can leave straight ones, and missing half the dialogue
# would make the "current dialogue %" figure quoted to the editor wrong.
_DIALOGUE_RE = re.compile(r'[“"]([^”"]*)[”"]')
TOKENS_PER_WORD = 1.4          # same rule of thumb as local-book-generator.py
OLLAMA_URL = "http://localhost:11434"
EDITOR_NUM_CTX = 16384         # matches Modelfile.llama31-16k

# How short a revision may be before it stops being "tighter" and starts being
# "incomplete". The old code used 0.50, which sat exactly on top of where this
# editor naturally lands (measured 51-65% across a real book), so acceptance was
# close to a coin flip. 0.35 leaves room for genuine condensing while still
# catching a revision that quit halfway.
# Raised from 0.35 to 0.75 on 2026-08-14, after the "Stars Without End" run.
#
# The history here is worth keeping, because both previous values were wrong in
# opposite directions and the fix is not a third guess at the number:
#
#   0.50 (original) — sat exactly on top of where this editor naturally lands
#     (51-65%), so acceptance was near a coin flip. Chapters that fell just under
#     were rejected three times and shipped UNEDITED. Two whole chapters of
#     "Ashes & Ember" went out that way, and they were the worst prose in it.
#
#   0.35 (my fix for that) — stopped the silent reverts, and the quality metrics
#     jumped: readability 46 -> 81, emotional resonance 37 -> 100. But it also
#     licensed the editor to cut far harder than "tightening". On the next run
#     the writer delivered 14,554 words, correctly in range, and the editor
#     shipped 8,331 — a 43% cut that dropped the book below the 12,000-word floor
#     its own length preset had asked for.
#
# A threshold alone cannot separate "condensed well" from "cut too much", because
# both look identical to a length check. So the threshold is now set where real
# tightening actually lives (10-25% reduction), and the work of hitting it is
# moved into the prompt: every attempt states the required word count outright,
# and a short return is retried WITH the shortfall quoted back. Falling back to
# the unedited original is the last resort it was always meant to be, not the
# routine outcome.
REVISION_MIN_RATIO = 0.75
# Ceiling too — a "revision" much longer than the source has started writing new
# story rather than editing the existing one.
REVISION_MAX_RATIO = 1.30
# Below MIN_RATIO but complete and not gutted, a revision still beats shipping
# the chapter unedited. Only used after every attempt has failed.
#
# Set to 0.35 deliberately, BELOW what this editor produces when it's cutting
# hardest. On the "Stars Without End" run its most aggressive chapters came back
# at 41% and 42%; a salvage floor of 0.45 would have rejected exactly those and
# shipped them unedited — reinstating the original bug through a side door.
# Caught in testing, not in production.
#
# The ordering this creates is what matters: a full-length real edit is best, a
# short-but-complete edit is second, and the unedited original is last. Those
# heavily-cut revisions are the ones that took readability from 46 to 81 and
# emotional resonance from 37 to 100 — worse than a full-length edit, clearly
# better than no edit at all.
REVISION_SALVAGE_RATIO = 0.35

# A revision that passes the length check is not necessarily the BEST revision
# this editor can produce, and until 2026-08-15 the loop below returned the first
# one that fit and threw the rest away. "Starbound Odyssey" showed what that
# costs. Nine chapters, three attempts allowed each:
#
#   fully revised (6 chapters)   FK 13.1   3.8% dialogue   28.1-word sentences
#   salvaged      (3 chapters)   FK 10.9  14.1% dialogue   23.1-word sentences
#
# The chapters the length check REJECTED read better on every measure than the
# ones it accepted. That is the length floor fighting the style directives: told
# it came back short, the model complies the safe way — it restores the original
# dense narration instead of rewriting it into scenes. So a first attempt passing
# the length test is weak evidence that it is any good.
#
# The fix is to score each attempt on the same signals the scorer will later
# grade the book on, and keep the best length-valid one instead of the first.
# PROSE_GOOD_ENOUGH bounds the cost: an attempt at or above it ends the loop
# immediately, so a genuinely good first revision still costs exactly one
# generation, as before. Only a mediocre-but-passing attempt spends the retries
# that were already budgeted. Worst case is unchanged at MAX_REVIEW_ATTEMPTS.
PROSE_GOOD_ENOUGH = 80.0
# A composite averages, and an average can be bought off. Measured on the real
# "Starbound Odyssey" salvaged profile — grade 10.9, 14% dialogue, 23-word
# sentences — the composite comes out at 79: two strong axes carrying one weak
# one. But under-dialogue is not an incidental weakness of this pipeline, it is
# THE weakness, present in every book it has produced (3.8-14.1% against a 25-40%
# convention). So the early exit needs it to be genuinely addressed and not
# merely outvoted. A chapter must clear the composite AND carry real dialogue to
# end the loop early; anything else gets the second look.
EARLY_EXIT_MIN_DIALOGUE = 20.0
# ...and this bounds it properly. PROSE_GOOD_ENOUGH alone would not: scored on
# the composite below, "Starbound Odyssey" chapters land around 30 (the fully
# revised ones) to 69 (the salvaged ones), so nothing would clear 75 and every
# chapter would burn all three attempts — roughly tripling the editorial pass on
# a book where two thirds of chapters currently stop at one or two.
#
# So a chapter gets at most ONE extra generation motivated purely by prose
# quality. Retries caused by an attempt actually failing (too short, truncated,
# refusal) are unchanged and still bounded by MAX_REVIEW_ATTEMPTS. Worst case
# per chapter is therefore identical to today; the realistic case is one extra
# generation on chapters that used to stop at the first passing attempt.
MAX_QUALITY_RETRIES = 1

# Flesch-Kincaid grade bands, copied from scoring_agent.GENRE_READABILITY_TARGETS
# so the editor optimises for the same target the scorer measures against. Kept
# as a copy rather than an import for the same reason detect_degeneration() is
# copied into all three files: these run as independent scripts and a shared
# module would make the writer stage depend on the scorer.
EDITOR_READABILITY_TARGETS = {
    "cozy mystery": (5, 8),
    "romantic suspense": (6, 9),
    "sci-fi thriller": (7, 9),
    "dystopian survival": (6, 9),
    "noir detective": (7, 9),
    "space opera": (7, 10),
    "psychological horror": (7, 10),
    "epic fantasy": (8, 11),
}
EDITOR_DEFAULT_READABILITY = (7, 10)
DIALOGUE_TARGET_LO = 0.25       # commercial genre fiction runs 25-40% dialogue
DIALOGUE_TARGET_HI = 0.40
LONG_SENTENCE_WORDS = 35        # the ceiling the style directives ask for
_VOWEL_GROUP_RE = re.compile(r"[aeiouy]+")


def _syllables(word: str) -> int:
    """Vowel-group heuristic, identical to scoring_agent._count_syllables."""
    word = word.lower()
    count = len(_VOWEL_GROUP_RE.findall(word))
    if word.endswith("e") and count > 1:
        count -= 1
    return max(count, 1)


def _split_sentences(text: str) -> list:
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return []
    return [p.strip() for p in re.split(r"(?<=[.!?])\s+", clean) if p.strip()]


def prose_quality(text: str, genre: str = "") -> dict:
    """0-100 composite of the three prose signals this pipeline actually grades:
    reading level against the genre band, dialogue share, and sentence
    discipline. Deterministic and free — no model call — so it can be run on
    every revision attempt without adding runtime.

    Deliberately NOT a clone of the scorer's full rubric. Plot, coherence and
    continuity are properties of the chapter's content, which revision is not
    supposed to change; grading attempts on them would mostly measure noise.
    These three are the things a revision moves, and the three the "Starbound
    Odyssey" salvaged chapters beat the accepted ones on.
    """
    # NORMALISE QUOTES FIRST. This line is the whole reason the first run of
    # this function was measuring noise on its most important axis.
    #
    # The writer emits MIXED quote characters — "Echoes of Elyria" chapter 3 came
    # out of the model with 89 straight quotes, 12 curly-open and 23 curly-close,
    # and chapter 6 with 138 / 3 / 37. _DIALOGUE_RE pairs an opening mark with
    # the next closing mark, so on straight quotes (where one character is both)
    # a single unbalanced quote inverts the alternation and every subsequent
    # match captures the NARRATION BETWEEN the dialogue instead of the dialogue.
    #
    # Measured on the real shipped chapter 3, same text, same regex:
    #     as shipped (curly-normalised)   39.5% dialogue
    #     the model's raw straight quotes  4.8% dialogue
    #
    # The editor was measuring the second and the scorer the first, which is why
    # the first run logged 2-10% dialogue on every chapter while the scorer gave
    # the book 100/100 for a 24.6% dialogue ratio. Nothing could clear the
    # dialogue gate, so every chapter burned its extra generation for nothing.
    #
    # apply_ebook_typography() is what the pipeline itself runs before the text is
    # written to edited/, so normalising through it here means this function
    # measures exactly the text the scorer will later see.
    text, _ = apply_ebook_typography(text)

    words = _WORD_RE.findall(text)
    n = len(words)
    if n < 50:
        return {"score": 0.0, "grade": 0.0, "dialogue_pct": 0.0,
                "avg_sentence": 0.0, "long_pct": 0.0, "detail": "too short to measure"}

    sentences = _split_sentences(text)
    lengths = [len(_WORD_RE.findall(s)) for s in sentences]
    lengths = [l for l in lengths if l]
    avg_sentence = n / max(1, len(lengths))
    long_share = sum(1 for l in lengths if l > LONG_SENTENCE_WORDS) / max(1, len(lengths))

    grade = 0.39 * avg_sentence + 11.8 * (sum(_syllables(w) for w in words) / n) - 15.59
    lo, hi = EDITOR_READABILITY_TARGETS.get(genre, EDITOR_DEFAULT_READABILITY)
    if lo <= grade <= hi:
        readability = 100.0
    else:
        deviation = (lo - grade) if grade < lo else (grade - hi)
        readability = max(0.0, 100.0 - deviation * 12)

    dialogue_share = sum(len(_WORD_RE.findall(d)) for d in _DIALOGUE_RE.findall(text)) / n
    if DIALOGUE_TARGET_LO <= dialogue_share <= DIALOGUE_TARGET_HI:
        dialogue = 100.0
    elif dialogue_share < DIALOGUE_TARGET_LO:
        # Linear from 0% dialogue = 0 up to the band. Under-dialogue is the
        # failure mode this pipeline actually has (measured 3.8-14.1%), so it is
        # scored on a full-range slope rather than a token penalty.
        dialogue = max(0.0, dialogue_share / DIALOGUE_TARGET_LO * 100)
    else:
        dialogue = max(0.0, 100.0 - (dialogue_share - DIALOGUE_TARGET_HI) * 250)

    # Sentence discipline: the average should sit near 18, and long sentences
    # should be rare. Both halves matter — a chapter can average 20 words and
    # still be 30% unreadable 45-word sentences.
    avg_penalty = max(0.0, abs(avg_sentence - 18.0) - 4.0) * 6.0
    long_penalty = long_share * 220.0
    discipline = max(0.0, 100.0 - avg_penalty - long_penalty)

    score = 0.35 * readability + 0.35 * dialogue + 0.30 * discipline
    return {
        "score": round(score, 1),
        "grade": round(grade, 1),
        "dialogue_pct": round(dialogue_share * 100, 1),
        "avg_sentence": round(avg_sentence, 1),
        "long_pct": round(long_share * 100, 1),
        "detail": (f"prose {score:.0f}/100 (FK {grade:.1f} vs {lo}-{hi}, "
                   f"{dialogue_share * 100:.0f}% dialogue, {avg_sentence:.0f}-word "
                   f"sentences, {long_share * 100:.0f}% over {LONG_SENTENCE_WORDS})"),
    }


def _weakest_prose_signal(q: dict, genre: str = "") -> str:
    """Which of the three signals is furthest from target — used to aim the
    retry at the actual weakness instead of repeating the whole style sheet."""
    lo, hi = EDITOR_READABILITY_TARGETS.get(genre, EDITOR_DEFAULT_READABILITY)
    if q["dialogue_pct"] < DIALOGUE_TARGET_LO * 100:
        return (
            f"\n\nIMPORTANT — that attempt was usable but thin on dialogue: only "
            f"{q['dialogue_pct']:.0f}% of it is spoken lines, against the 25-40% commercial "
            f"genre fiction runs at. Do it again and dramatise more of the narration. Find the "
            f"passages where you REPORT what characters said, thought at each other, argued "
            f"about or decided together, and write those as scenes — actual lines of dialogue, "
            f"back and forth, with action beats between them. Keep every event. This swaps "
            f"narration for dialogue at about the same length, so the word count is unaffected."
        )
    if q["long_pct"] > 12 or q["avg_sentence"] > 26:
        return (
            f"\n\nIMPORTANT — that attempt was usable but the sentences run long: an average of "
            f"{q['avg_sentence']:.0f} words, with {q['long_pct']:.0f}% of them over "
            f"{LONG_SENTENCE_WORDS}. Do it again and break the long ones at their natural joints. "
            f"Where a sentence stacks two or three subordinate clauses, make it two or three "
            f"sentences. Aim for an average nearer 18. Do not cut material to do it."
        )
    if q["grade"] > hi:
        return (
            f"\n\nIMPORTANT — that attempt was usable but reads heavier than this genre wants "
            f"(grade level {q['grade']:.1f}; the target is {lo}-{hi}). Do it again with plainer "
            f"word choice and more concrete verbs. Keep every scene and event."
        )
    return (
        "\n\nThat attempt was usable. Try once more for a sharper version: more dialogue, "
        "shorter sentences, concrete detail over summary. Keep every scene and event."
    )


def ollama_chat(prompt: str, num_predict: int, timeout: int = 1800) -> str:
    """One prompt in, text out, straight to Ollama with an explicit options dict.

    Used for revision instead of CrewAI because CrewAI/litellm silently drops
    max_tokens for Ollama models — proven on the writer side, where a chapter
    capped at 2,730 words came back with 8,759. For revision the same bug bites
    in the opposite direction: without a stated budget, a long chapter's rewrite
    can stop early and then get rejected for being 'too short'."""
    response = requests.post(
        f"{OLLAMA_URL}/api/chat",
        json={
            "model": EDITOR_MODEL_NAME,
            "messages": [
                {"role": "system", "content":
                    "You are a senior fiction editor. You return complete revised manuscript "
                    "prose and nothing else — no commentary, no notes, no summaries."},
                {"role": "user", "content": prompt},
            ],
            "stream": False,
            "options": {
                "num_predict": num_predict,
                "num_ctx": EDITOR_NUM_CTX,
                "temperature": 0.6,   # lower than the writer's — this is revision, not invention
            },
        },
        timeout=timeout,
    )
    response.raise_for_status()
    return (response.json().get("message") or {}).get("content", "")


def revision_is_usable(revised: str, original: str) -> "tuple[bool, str]":
    """Decide whether a revision can replace the original. Returns (ok, reason).

    Judging on length alone was the old mistake — it conflated two very different
    things. A revision at 55% of the original because it cut padding is exactly
    what this editor is for. A revision at 55% because it stopped mid-scene is
    useless. Length can't tell those apart, so this checks completeness directly:
    does it end on a finished sentence, and is it coherent prose throughout."""
    text = (revised or "").strip()
    if not text:
        return False, "empty response"
    if looks_like_refusal(text):
        return False, "refusal"

    orig_words = len(_WORD_RE.findall(original))
    new_words = len(_WORD_RE.findall(text))
    if orig_words and new_words < orig_words * REVISION_MIN_RATIO:
        return False, f"too short ({new_words:,} words vs {orig_words:,} — under " \
                      f"{int(REVISION_MIN_RATIO * 100)}%)"
    if orig_words and new_words > orig_words * REVISION_MAX_RATIO:
        return False, f"too long ({new_words:,} words vs {orig_words:,} — over " \
                      f"{int(REVISION_MAX_RATIO * 100)}%, likely writing new story)"

    # Truncation check. A rewrite that ran out of budget stops mid-sentence; a
    # finished one lands on terminal punctuation (allowing a trailing quote mark
    # for a chapter that ends on dialogue).
    if text.rstrip().rstrip('"”’\'') and text.rstrip().rstrip('"”’\'')[-1] not in ".!?…":
        return False, "ends mid-sentence (looks truncated)"

    # Meta-commentary leaking into the manuscript — the model occasionally
    # prefaces the prose with "Here is the revised chapter:".
    head = text[:200].lower()
    if any(p in head for p in ("here is the revised", "here's the revised", "revised chapter:",
                               "i have revised", "below is the revised")):
        return False, "starts with meta-commentary instead of prose"

    return True, "ok"


def revise_chapter(chapter_text: str, ch_num: int, ch_title: str, style_sheet: dict,
                   issues: list, emotional_pitch: "int | None" = None) -> "str | None":
    """Auto-apply only: produce a clean revised chapter incorporating the review's
    fixes. Returns None (not an exception) when no attempt yields usable prose,
    so the caller falls back to the original (typography-cleaned) text rather
    than losing the chapter.

    Rewritten 2026-08-14 for two reasons, both from the "Ashes & Ember" run:

    1. IT WAS SILENTLY REVERTING CHAPTERS. The old acceptance test was
       `len(revised) > 0.5 * len(original)`. But a real revision from this
       editor condenses — measured across that book's 15 chapters, accepted
       revisions came back at 51-65% of the original, i.e. sitting right on top
       of the 50% line. Chapters 4 and 14 fell just under it, were rejected
       three times each, and shipped UNEDITED. They were then the worst prose in
       the book: chapter 4 had 47% of its sentences over 40 words and contained
       a single 112-word sentence. So the guard meant to protect quality was the
       reason the two chapters that most needed editing didn't get any.

    2. IT WENT THROUGH CREWAI, which drops max_tokens for Ollama (proven on the
       writer side — a chapter capped at 2,730 words returned 8,759). For the
       two longest chapters that most likely meant the revision ran out of room
       and came back truncated, which the length test then read as "too short"
       — the same failure by a second route. This now calls Ollama directly with
       an explicit num_predict sized to the chapter, so there is enough budget
       to rewrite it in full and the setting is verifiable.

    The analysis passes still use CrewAI: they return JSON and genuinely benefit
    from its retry handling. Revision is one prompt in, prose out, which CrewAI
    adds nothing to.
    """
    if not issues:
        return chapter_text

    issues_text = "\n".join(
        f"- [{i.get('category', '?')}] {i.get('issue', '')} -> {i.get('fix', '')}" for i in issues
    ) or "(none)"

    original_words = len(_WORD_RE.findall(chapter_text))
    min_words = int(original_words * REVISION_MIN_RATIO)
    max_words = int(original_words * REVISION_MAX_RATIO)
    dialogue_words = sum(len(_WORD_RE.findall(d)) for d in _DIALOGUE_RE.findall(chapter_text))
    dialogue_pct = 100 * dialogue_words / max(1, original_words)

    # Style requirements live HERE rather than in the writer prompt, on evidence:
    # the writer was told to write short sentences and produced a 26-word average
    # with a 112-word outlier, while moving the same instruction to the editor
    # took readability from 46/100 to 81/100 in a single run, and the emotional
    # pitch below from 37/100 to 100/100. Instructions the writer ignores, the
    # editor acts on.
    style_directives = (
        "\nALSO APPLY THESE, whether or not the issue list mentions them:\n"
        f"- LENGTH. The revised chapter must be between {min_words:,} and {max_words:,} words "
        f"(the original is {original_words:,}). This is a hard requirement. Tightening wordy "
        f"sentences is wanted; CUTTING SCENES, EVENTS OR PARAGRAPHS IS NOT. Every scene in the "
        f"original must still be present, with the same events in the same order. If your "
        f"revision is coming in short, you are deleting content instead of editing it — keep the "
        f"material and improve the sentences in place.\n"
        f"- SENTENCE LENGTH. Break up long sentences. No sentence should run past about 35 words, "
        f"and the chapter should average nearer 15 than 25. Split at the natural joint — where a "
        f"sentence stacks two or three subordinate clauses, make it two or three sentences.\n"
        f"- DIALOGUE. This chapter is currently {dialogue_pct:.0f}% dialogue; commercial genre "
        f"fiction runs 25-40%. Where two or more characters share a scene, have them SPEAK. Turn "
        f"reported speech into real exchanges — replace 'she told him what she had found in the "
        f"hold' with the actual lines, back and forth, with action beats between them. This "
        f"converts narration into dialogue at roughly equal length, so it costs nothing against "
        f"the word count above. Do not invent new plot; dramatise what the narration already "
        f"describes.\n"
        f"- Keep the author's voice and every story beat.\n"
    )
    if emotional_pitch is not None:
        # Confirmed effective 2026-08-14: the same instruction given to the
        # WRITER was a complete no-op (planned pitches 10-100, delivered prose
        # flat, intensity spread 5.6). Given to the editor it took the spread to
        # 20.7 and the metric from 37/100 to 100/100.
        if emotional_pitch >= 85:
            style_directives += (
                f"- EMOTIONAL PITCH {emotional_pitch}/100 — a peak chapter. Sharpen it: shorter "
                f"sentences, harder verbs, less reflection, more immediacy.\n"
            )
        elif emotional_pitch <= 30:
            style_directives += (
                f"- EMOTIONAL PITCH {emotional_pitch}/100 — a deliberately quiet chapter. Let it "
                f"breathe. Interiority and small detail belong here; resist manufactured "
                f"tension.\n"
            )

    prompt = (
        f"{FICTION_FRAMING}"
        f"You are a senior editor producing a clean revised draft of Chapter {ch_num} "
        f"('{ch_title}') of the novel '{style_sheet['title']}', incorporating the specific fixes "
        f"an editorial review already identified.\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        f"Issues to fix:\n{issues_text}\n"
        f"{style_directives}\n"
        "Write the complete revised chapter, from its first line to its last. Write only the "
        "chapter prose itself — no title header, no meta-commentary, no summary of what you "
        "changed, and do not stop early or abbreviate the ending."
    )

    # Generous headroom: the revision may run slightly LONGER than the original
    # (dialogue conversion adds lines), so size the budget above the source.
    num_predict = int(original_words * 1.6 * TOKENS_PER_WORD)

    genre = style_sheet.get("genre", "")
    best = None            # best salvageable (too-short but complete) attempt
    best_quality = -1.0    # its prose score
    winner = None          # best length-VALID attempt
    winner_quality = -1.0
    winner_attempt = 0
    quality_retries = 0
    correction = ""   # feedback appended to the next attempt after a near miss
    for attempt in range(1, MAX_REVIEW_ATTEMPTS + 1):
        try:
            revised = ollama_chat(prompt + correction, num_predict=num_predict)
        except requests.exceptions.RequestException as e:
            print(f"[WARN] Revision attempt {attempt}/{MAX_REVIEW_ATTEMPTS} for Chapter {ch_num}: "
                  f"Ollama request failed ({e}).")
            continue

        ok, why = revision_is_usable(revised, chapter_text)
        quality = prose_quality(revised, genre)

        if ok:
            # Keep the best, not the first. See PROSE_GOOD_ENOUGH above for why:
            # passing the length check turned out to be nearly uncorrelated with
            # reading well, and on "Starbound Odyssey" the attempts this loop
            # discarded measured better than the ones it kept.
            if quality["score"] > winner_quality:
                winner, winner_quality, winner_attempt = revised, quality["score"], attempt
            print(f"[Editor] Chapter {ch_num} attempt {attempt}/{MAX_REVIEW_ATTEMPTS} valid "
                  f"({len(_WORD_RE.findall(revised)):,} words) — {quality['detail']}")
            good_enough = (quality["score"] >= PROSE_GOOD_ENOUGH
                           and quality["dialogue_pct"] >= EARLY_EXIT_MIN_DIALOGUE)
            if (good_enough
                    or attempt == MAX_REVIEW_ATTEMPTS
                    or quality_retries >= MAX_QUALITY_RETRIES):
                break
            # Usable but unremarkable: spend one retry aimed at its weakest
            # signal rather than banking it. If that retry comes back valid,
            # whichever of the two reads better is kept and the loop stops.
            quality_retries += 1
            correction = _weakest_prose_signal(quality, genre)
            continue

        # Retry with the miss quoted back, rather than re-sending the identical
        # prompt and hoping for a different sample. The previous code retried
        # blind, which is why a chapter that came back short came back short
        # three times in a row — nothing in the second attempt told the model
        # what was wrong with the first.
        got = len(_WORD_RE.findall(revised))
        if why.startswith("too short"):
            # Reworded 2026-08-15. The old text — "you cut roughly N words of
            # material that needs to stay" — reliably produced length, and
            # reliably produced it the worst way: the model pasted the original
            # narration back in. That is the mechanism behind the table in the
            # PROSE_GOOD_ENOUGH comment. Asking for the SAME words as dialogue
            # gets the count back without undoing the edit.
            correction = (
                f"\n\nIMPORTANT — your previous attempt was rejected. It came back at {got:,} "
                f"words, but this chapter must be at least {min_words:,} — about "
                f"{min_words - got:,} words more.\n"
                f"Do NOT get there by pasting the original narration back in, and do NOT pad "
                f"with description. Get there by DRAMATISING. Go through the chapter and find "
                f"every passage that reports something happening between characters — an "
                f"argument, a briefing, a confession, a negotiation, someone explaining what "
                f"they found — and write it as a scene instead: real lines of dialogue, back "
                f"and forth, with action beats between them. A paragraph of summarised exchange "
                f"becomes a page of played-out exchange, which is where the words come from.\n"
                f"Keep every scene and event, in the same order. Do not summarise or skip."
            )
        elif why.startswith("too long"):
            correction = (
                f"\n\nIMPORTANT — your previous attempt was rejected. It came back at {got:,} "
                f"words against a {max_words:,} maximum. Do not add new scenes or events; edit "
                f"what is already there."
            )
        elif why.startswith("ends mid-sentence"):
            correction = (
                "\n\nIMPORTANT — your previous attempt stopped mid-sentence before the end of "
                "the chapter. Write the chapter through to its final line."
            )
        # Hold the least-bad attempt: a revision that is merely SHORTER than we'd
        # like still beats shipping the chapter unedited, which is what the old
        # code did.
        #
        # But only that failure mode qualifies. An earlier version of this
        # salvaged anything that wasn't a refusal, which meant a revision
        # rejected for ending mid-sentence could still be salvaged — putting a
        # truncated chapter into the manuscript while the log claimed it "ends
        # cleanly". Caught in testing before it ever ran. A short-but-complete
        # revision is a worse edit; a truncated one is a broken chapter, and the
        # unedited original beats it every time.
        #
        # Ranked on prose quality rather than raw length as of 2026-08-15, with
        # the salvage floor still applied first. Length is the eligibility test
        # ("is enough of the chapter here?"); among the attempts that clear it,
        # the better-reading one wins. Ranking by length alone was picking the
        # attempt that had restored the most original narration — precisely the
        # attempt the "Starbound Odyssey" numbers say to avoid.
        if (why.startswith("too short")
                and len(_WORD_RE.findall(revised)) >= original_words * REVISION_SALVAGE_RATIO
                and quality["score"] > best_quality):
            best, best_quality = revised, quality["score"]
        print(f"[WARN] Revision attempt {attempt}/{MAX_REVIEW_ATTEMPTS} for Chapter {ch_num} "
              f"rejected ({why})." + (" Retrying..." if attempt < MAX_REVIEW_ATTEMPTS else ""))

    if winner is not None:
        print(f"[Editor] Chapter {ch_num}: kept attempt {winner_attempt} of "
              f"{MAX_REVIEW_ATTEMPTS} — best prose score {winner_quality:.0f}/100 "
              f"({len(_WORD_RE.findall(winner)):,} words).")
        return winner

    if best is not None:
        print(f"[WARN] No attempt fully passed for Chapter {ch_num}, but the best one covers "
              f"{len(_WORD_RE.findall(best)):,} of {original_words:,} words, ends cleanly and "
              f"scores {best_quality:.0f}/100 on prose — using it rather than shipping the "
              f"chapter unedited.")
        return best

    print(f"[WARN] Could not get a usable revision for Chapter {ch_num} after {MAX_REVIEW_ATTEMPTS} "
          f"attempts. Keeping the original text (with typography fixes only) in edited/.")
    return None


# =====================================================================
# 8. WHOLE-BOOK CHECKS (deterministic, computed from data already gathered
#    or from cheap file scans — no extra LLM calls needed)
# =====================================================================

STOPWORDS = set((
    "the a an and or but if of to in on at for with as by from is are was were be been being "
    "this that these those it its his her their they he she him them we you your our i me my "
    "not no so than then there here when where what who which how why all any both each few "
    "more most other some such only own same just also into out up down over under again "
    "further once had have has do does did doing would could should will shall can may might "
    "must about after before between during through above below because while"
).split())


def extract_candidate_keywords(full_text: str, character_names: list, top_n: int = 15) -> list:
    """Simple frequency-based keyword extraction for the metadata pass — no NLP
    library dependency, just stopword-filtered word frequency. Character names
    are excluded since they're already known, not useful as search keywords
    on their own."""
    exclude = {n.lower() for full in character_names for n in full.split()}
    words = re.findall(r"[A-Za-z']{4,}", full_text.lower())
    counts = Counter(w for w in words if w not in STOPWORDS and w not in exclude)
    return [word for word, _ in counts.most_common(top_n)]


def audit_front_back_matter(book_dir: str) -> list:
    """Lightweight check that Stage D's manuscript assembly actually produced
    what it's supposed to — a sanity/regression check, not a re-implementation
    of Stage D's own logic."""
    notes = []
    raw_path = os.path.join(book_dir, "manuscript_raw.txt")
    docx_path = os.path.join(book_dir, "manuscript.docx")
    if not os.path.isfile(docx_path):
        notes.append("manuscript.docx not found — Stage D may not have completed.")
    if not os.path.isfile(raw_path):
        notes.append("manuscript_raw.txt not found — Stage D may not have completed.")
        return notes
    with open(raw_path, "r", encoding="utf-8") as f:
        raw_text = f.read()
    if "Copyright" not in raw_text:
        notes.append("No 'Copyright' text found in the manuscript — check the copyright page.")
    if not notes:
        notes.append("Front/back matter looks present (copyright text found, both output files exist).")
    return notes


def extract_author_name(book_dir: str) -> str:
    """Same approach as scoring_agent.py's extract_author_name() — parse the pen name
    straight out of manuscript_raw.txt's copyright/byline rather than duplicating
    local-book-generator.py's GENRE_PEN_NAMES mapping into a third file where it could
    drift out of sync."""
    path = os.path.join(book_dir, "manuscript_raw.txt")
    if not os.path.isfile(path):
        return "(author name not found — manuscript_raw.txt missing; check Stage D output)"
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    m = re.search(r"^Copyright\s*©?\s*\d{4}\s+(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    m = re.search(r"^by\s+(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    return "(author name not found in manuscript_raw.txt — check the file manually)"


def _add_static_toc(document: "Document", chapters: list) -> None:
    """Write a plain, pre-rendered chapter list — 'Chapter N — Title', one per
    line, no page numbers (ebooks reflow, so a printed page number would just be
    wrong on an e-reader anyway). Identical to local-book-generator.py's
    _add_static_toc(). Replaces an earlier real Word TOC *field* (dynamic,
    computed by Word itself on 'Update Field'): that reads fine in Word, but
    ebook converters generally don't run Word's field-update logic, so the
    field imported as a permanently empty Contents page — confirmed by a
    trial conversion, 2026-08-12. A static list has
    real text baked in at build time, so it always renders correctly no matter
    what app opens or converts the file."""
    for chapter in chapters:
        line = document.add_paragraph()
        line.add_run(f"Chapter {chapter['chapter_number']} — {chapter['title']}")


def build_edited_docx_manuscript(edited_dir: str, outline: dict, author_name: str,
                                  chapter_bodies: dict) -> str:
    """Auto-apply convenience: build a reader-ready .docx from the EDITED chapter
    text, using the exact same formatting as local-book-generator.py's
    build_docx_manuscript() (title page, copyright page, static Contents list,
    centered 'Chapter N' headings). The raw manuscript.docx that Stage D writes to the book's
    root folder is generated once, before the editor ever runs — it never reflects
    auto-apply's revisions, so it's the wrong file to read from. This is the one
    that should: it lives in edited/ next to the revised chapters, the covers,
    the score files, and Finished Product Notes.txt, so that folder is the single
    self-contained place holding the finished book.
    Chapters with no entry in chapter_bodies (shouldn't normally happen) are skipped
    with a [WARN], same as the raw builder does for a missing chapter file."""
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(outline["title"])
    title_run.font.size = Pt(28)
    title_run.font.bold = True

    for _ in range(3):
        document.add_paragraph()

    by_p = document.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_run = by_p.add_run(f"by {author_name}")
    by_run.font.size = Pt(16)

    document.add_page_break()

    year = datetime.now().year
    copyright_p = document.add_paragraph()
    copyright_p.add_run(f"Copyright © {year} {author_name}").bold = True
    document.add_paragraph("All rights reserved.")
    document.add_paragraph()
    document.add_paragraph(COPYRIGHT_BOILERPLATE)
    document.add_page_break()

    if INCLUDE_TOC:
        toc_heading = document.add_paragraph()
        toc_heading.add_run("Contents").bold = True
        toc_heading.style = document.styles["Heading 1"]
        _add_static_toc(document, outline["chapters"])
        document.add_page_break()

    chapters = outline["chapters"]
    for chapter in chapters:
        ch_num = chapter["chapter_number"]
        body = chapter_bodies.get(ch_num)
        if body is None:
            print(f"[WARN] No edited text in memory for chapter {ch_num}, skipping in edited/manuscript.docx.")
            continue

        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.add_run(f"Chapter {ch_num}")

        subheading = document.add_paragraph()
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subheading.add_run(chapter["title"])
        sub_run.italic = True

        document.add_paragraph()

        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            p = document.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Inches(0) if i == 0 else Inches(0.3)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)

        if ch_num != chapters[-1]["chapter_number"]:
            document.add_page_break()

    output_path = os.path.join(edited_dir, "manuscript.docx")
    document.save(output_path)
    return output_path


COVER_FILENAMES = (
    "cover_option_1.png", "cover_option_1.jpg",
    "cover_option_2.png", "cover_option_2.jpg",
    "cover_option_3.png", "cover_option_3.jpg",
)


def copy_cover_images(book_dir: str, edited_dir: str) -> "tuple[list, list]":
    """Auto-apply convenience: copy the 3 finished cover options into edited/
    alongside the revised chapters, so the edited folder is self-contained
    (text + art) rather than requiring you to hop back to the raw book folder
    for cover files. Best-effort — a missing cover (older book, or a run where
    cover generation failed) is noted, not a fatal error."""
    copied, missing = [], []
    for filename in COVER_FILENAMES:
        src = os.path.join(book_dir, filename)
        if os.path.isfile(src):
            shutil.copyfile(src, os.path.join(edited_dir, filename))
            copied.append(filename)
        else:
            missing.append(filename)
    return copied, missing


# =====================================================================
# 8b. TIMING INSTRUMENTATION (dashboard section 6.14 — ETA/countdown timers)
# =====================================================================


def _emit_timing(event: str, stage: str, **fields) -> None:
    """Machine-readable timing marker — see local-book-generator.py's copy
    of this same function for the full rationale. Duplicated here per this
    codebase's "each script runs standalone" convention."""
    parts = [f"event={event}", f"stage={stage}"]
    for key, value in fields.items():
        # Sanitize spaces (e.g. a chapter title in label=...) so the
        # dashboard's whitespace-tokenized parser doesn't fragment it.
        safe_value = re.sub(r"\s+", "_", str(value)) if value is not None else ""
        parts.append(f"{key}={safe_value}")
    print("[TIMING] " + " ".join(parts))


# =====================================================================
# 9. MAIN
# =====================================================================

outline = load_outline(BOOK_DIR)
style_sheet = load_or_seed_style_sheet(BOOK_DIR, outline)
chapters_meta = {c["chapter_number"]: c for c in outline.get("chapters", [])}

chapter_files = discover_chapters(BOOK_DIR)
if not chapter_files:
    print(f"[ERROR] No chapter_NN.txt files found in '{BOOK_DIR}'.")
    sys.exit(1)

edited_dir = os.path.join(BOOK_DIR, "edited")
if AUTO_APPLY:
    os.makedirs(edited_dir, exist_ok=True)

chapter_reports = []  # one dict per chapter, for the consolidated report
full_manuscript_parts = []
edited_chapter_bodies = {}  # ch_num -> final body text, captured in-memory for build_edited_docx_manuscript()

print(f"\n--- Editorial pass starting: {len(chapter_files)} chapter(s) ---")

for ch_num, path in chapter_files:
    ch_title, chapter_text = split_chapter_file(path)
    ch_meta = chapters_meta.get(ch_num, {})
    ch_summary = ch_meta.get("summary", "(no outline summary on file)")
    full_manuscript_parts.append(chapter_text)

    print(f"\n--- Editorial pass: Chapter {ch_num}/{len(chapter_files)}: '{ch_title}' ---")
    _chapter_t0 = time.time()
    _emit_timing("start", "chapter_edit", ch=ch_num, total=len(chapter_files), label=ch_title)

    if "NEEDS MANUAL REGENERATION" in chapter_text:
        print(f"[SKIP] Chapter {ch_num} is a placeholder from a failed writer pass — skipping review "
              f"until it's regenerated.")
        chapter_reports.append({
            "chapter_number": ch_num, "title": ch_title, "skipped": True,
            "issues": [], "typography_notes": [],
        })
        if AUTO_APPLY:
            with open(os.path.join(edited_dir, os.path.basename(path)), "w", encoding="utf-8") as fh:
                fh.write(f"{ch_title}\n\n{chapter_text}")
            edited_chapter_bodies[ch_num] = chapter_text
        _emit_timing("end", "chapter_edit", ch=ch_num, total=len(chapter_files),
                     elapsed=f"{time.time() - _chapter_t0:.1f}")
        continue

    if SEPARATE_PASSES:
        macro_result = macro_pass(chapter_text, ch_num, ch_title, ch_summary, style_sheet)
        style_result = style_pass(chapter_text, ch_num, ch_title, style_sheet)
        micro_result = micro_pass(chapter_text, ch_num, ch_title, style_sheet)
        all_issues = macro_result["issues"] + style_result["issues"] + micro_result["issues"]
        new_facts = macro_result.get("new_facts", [])
        new_style_terms = micro_result.get("new_style_terms", [])
        breakdown = (f"{len(macro_result['issues'])} structural, "
                     f"{len(style_result['issues'])} style, "
                     f"{len(micro_result['issues'])} mechanics")
    else:
        combined = combined_review_pass(chapter_text, ch_num, ch_title, ch_summary, style_sheet)
        all_issues = combined["issues"]
        new_facts = combined.get("new_facts", [])
        new_style_terms = combined.get("new_style_terms", [])
        breakdown = "single combined pass"

    style_sheet["established_facts"].extend(new_facts)
    style_sheet["style_terms"].extend(new_style_terms)
    save_style_sheet(BOOK_DIR, style_sheet)  # incremental save — a crash mid-run doesn't lose progress

    # Deterministic collapse check, added 2026-08-14. The review pass is blind to
    # repetition collapse: on "Rainy Night, Lonely Soul" chapter 9 it reported
    # ZERO issues on a stretch containing 20 repetitions of the word "murder" and
    # "Lena bore it all" four times. With no issues, revise_chapter() returns
    # early and the chapter passes through untouched — so the one stage capable of
    # repairing the damage never even tried, and the scorer blocked the whole book
    # at the end instead.
    #
    # An LLM asked "find continuity and style problems" doesn't flag this, because
    # locally each sentence looks like prose. It's only visible statistically. So
    # the check is done in Python and injected as a high-severity issue, which
    # gives the revision pass an explicit mandate to rewrite that stretch.
    collapse = detect_degeneration(chapter_text)
    if collapse["degenerate"]:
        runaway = collapse.get("runaway_word")
        fix = (f"Rewrite the passage starting around word "
               f"{collapse['first_bad_word_index']} of {collapse['total_words']}. It has "
               f"collapsed into repetition — the same phrases and images cycle without "
               f"advancing the scene.")
        if runaway:
            fix += (f" In particular the word '{runaway[0]}' accounts for "
                    f"{runaway[1] * 100:.1f}% of a 400-word stretch; almost every use of it "
                    f"there should go.")
        fix += (" Replace the repetition with prose that actually moves the scene forward to "
                "the chapter's ending. Keep the events, cut the echoing.")
        all_issues.insert(0, {
            "category": "repetition_collapse",
            "severity": "high",
            "excerpt": " ".join(chapter_text.split()[
                (collapse["first_bad_word_index"] or 0):(collapse["first_bad_word_index"] or 0) + 40]),
            "issue": collapse["detail"],
            "fix": fix,
        })
        print(f"[Editor] Chapter {ch_num}: repetition collapse detected by the deterministic "
              f"check — {collapse['detail']}. Added as a high-severity issue so the revision "
              f"pass has to address it.")

    typography_text, typography_notes = apply_ebook_typography(chapter_text)

    print(f"[Editor] Chapter {ch_num}: {len(all_issues)} issue(s) found "
          f"({breakdown}), {len(typography_notes)} typography fix(es).")

    if AUTO_APPLY:
        revised = None
        if all_issues:
            # Emotional pitch comes from the outline (local-book-generator.py asks
            # for it per chapter). Passed to revision as a second attempt at
            # varying tone across the book: the writer received the same
            # instruction on the "Ashes & Ember" run and ignored it completely —
            # planned pitches ran 10 to 100, delivered prose was flat, with quiet
            # and peak chapters statistically identical. The editor at least
            # demonstrably acts on its instructions, so it's the better place to
            # try. Lower confidence than the sentence-length and dialogue
            # directives, which have direct evidence behind them.
            revised = revise_chapter(chapter_text, ch_num, ch_title, style_sheet, all_issues,
                                     emotional_pitch=ch_meta.get("emotional_pitch"))
        final_text = revised if revised else chapter_text
        final_text, _ = apply_ebook_typography(final_text)  # always land the mechanical fixes
        out_path = os.path.join(edited_dir, os.path.basename(path))
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(f"{ch_title}\n\n{final_text}")
        print(f"[Saved] {out_path}")
        edited_chapter_bodies[ch_num] = final_text

    chapter_reports.append({
        "chapter_number": ch_num, "title": ch_title, "skipped": False,
        "issues": all_issues, "typography_notes": typography_notes,
        "word_count": len(chapter_text.split()),
        "target_words": ch_meta.get("target_words"),
    })
    _emit_timing("end", "chapter_edit", ch=ch_num, total=len(chapter_files),
                 elapsed=f"{time.time() - _chapter_t0:.1f}")

print(f"\n--- Editorial pass complete: {len(chapter_files)} chapter(s) processed ---")

covers_copied, covers_missing = ([], [])
edited_docx_path = None
edited_docx_error = None
if AUTO_APPLY:
    covers_copied, covers_missing = copy_cover_images(BOOK_DIR, edited_dir)
    if covers_copied:
        print(f"[Covers] Copied {len(covers_copied)}/3 cover option(s) into {edited_dir}: "
              f"{', '.join(covers_copied)}")
    if covers_missing:
        print(f"[Covers] Not found in {BOOK_DIR}, skipped: {', '.join(covers_missing)}")

    # reader-ready .docx built from the EDITED text — the raw manuscript.docx Stage D
    # wrote to the book's root predates the editor and never reflects these revisions,
    # so this is the one file that represents the finished book. Best-effort:
    # a failure here (e.g. a corrupt/unreadable font on this machine) is reported, not
    # fatal — the edited chapter .txt files themselves are already safely on disk.
    author_name = extract_author_name(BOOK_DIR)
    try:
        edited_docx_path = build_edited_docx_manuscript(edited_dir, outline, author_name, edited_chapter_bodies)
        print(f"[Manuscript] reader-ready .docx written to {edited_docx_path}")
    except Exception as e:
        edited_docx_error = str(e)
        print(f"[WARN] Failed to build edited/manuscript.docx: {edited_docx_error}. "
              f"The edited chapter .txt files are still saved — you can build the .docx "
              f"by re-running with --auto-apply, or by hand from those files.")

# =====================================================================
# 10. WHOLE-BOOK SYNTHESIS + REPORT
# =====================================================================

full_manuscript_text = "\n\n".join(full_manuscript_parts)
character_names = [c["name"] for c in style_sheet["characters"]]
candidate_keywords = extract_candidate_keywords(full_manuscript_text, character_names)
front_back_notes = audit_front_back_matter(BOOK_DIR)

pacing_notes = []
for report in chapter_reports:
    if report["skipped"] or not report.get("target_words"):
        continue
    target = report["target_words"]
    actual = report["word_count"]
    if target and abs(actual - target) / target > 0.25:
        pacing_notes.append(
            f"Chapter {report['chapter_number']} ('{report['title']}'): {actual} words vs. a "
            f"{target}-word target ({'over' if actual > target else 'under'} by "
            f"{abs(actual - target) * 100 // target}%) — worth a pacing look."
        )

total_issues = sum(len(r["issues"]) for r in chapter_reports)
severity_counts = Counter(
    issue.get("severity", "unknown") for r in chapter_reports for issue in r["issues"]
)
category_counts = Counter(
    issue.get("category", "unknown") for r in chapter_reports for issue in r["issues"]
)

report_lines = []
report_lines.append(f"EDITORIAL REVIEW — {style_sheet['title']}")
report_lines.append(f"Genre: {style_sheet['genre']}")
report_lines.append(f"Mode: {'auto-applied to ' + edited_dir if AUTO_APPLY else 'suggestions only, no files changed'}")
if AUTO_APPLY and covers_copied:
    report_lines.append(f"Cover option(s) copied to edited/: {', '.join(covers_copied)}")
if AUTO_APPLY and covers_missing:
    report_lines.append(f"Cover option(s) not found, skipped: {', '.join(covers_missing)}")
if AUTO_APPLY and edited_docx_path:
    report_lines.append("Reader-ready manuscript.docx written to edited/ (built from the edited chapters)")
if AUTO_APPLY and edited_docx_error:
    report_lines.append(f"Reader-ready manuscript.docx FAILED to build: {edited_docx_error}")
report_lines.append(f"Total issues found: {total_issues}")
if severity_counts:
    report_lines.append("By severity: " + ", ".join(f"{k}={v}" for k, v in severity_counts.most_common()))
if category_counts:
    report_lines.append("By category: " + ", ".join(f"{k}={v}" for k, v in category_counts.most_common()))
report_lines.append("")

for report in chapter_reports:
    report_lines.append("=" * 70)
    if report["skipped"]:
        report_lines.append(f"Chapter {report['chapter_number']}: '{report['title']}' — SKIPPED "
                             f"(needs manual regeneration first)")
        report_lines.append("")
        continue
    report_lines.append(f"Chapter {report['chapter_number']}: '{report['title']}' "
                         f"({report['word_count']} words) — {len(report['issues'])} issue(s)")
    report_lines.append("")
    if not report["issues"]:
        report_lines.append("  No issues found.")
    for issue in report["issues"]:
        report_lines.append(f"  [{issue.get('severity', '?')}/{issue.get('category', '?')}] "
                             f"{issue.get('issue', '')}")
        if issue.get("excerpt"):
            report_lines.append(f"    Excerpt: \"{issue['excerpt']}\"")
        if issue.get("fix"):
            report_lines.append(f"    Fix: {issue['fix']}")
        report_lines.append("")
    if report["typography_notes"]:
        report_lines.append("  Typography (ebook formatting):")
        for note in report["typography_notes"]:
            report_lines.append(f"    - {note}")
        report_lines.append("")

report_lines.append("=" * 70)
report_lines.append("WHOLE-BOOK NOTES")
report_lines.append("")
report_lines.append("Pacing (chapters whose length deviated >25% from their outline target):")
if pacing_notes:
    for note in pacing_notes:
        report_lines.append(f"  - {note}")
else:
    report_lines.append("  - None flagged.")
report_lines.append("")
report_lines.append("Front/back matter audit:")
for note in front_back_notes:
    report_lines.append(f"  - {note}")
report_lines.append("")
report_lines.append("Candidate keywords (frequency-based, for listing metadata — review before using):")
report_lines.append("  " + ", ".join(candidate_keywords) if candidate_keywords else "  (none extracted)")
report_lines.append("")
report_lines.append(f"Established facts tracked in {STYLE_SHEET_FILENAME}: {len(style_sheet['established_facts'])}")
report_lines.append(f"Style-sheet terms tracked in {STYLE_SHEET_FILENAME}: {len(style_sheet['style_terms'])}")

report_path = os.path.join(BOOK_DIR, "editorial_review.txt")
with open(report_path, "w", encoding="utf-8") as f:
    f.write("\n".join(report_lines))

print(f"\n=======================================================")
print("EDITORIAL PASS COMPLETE")
print(f"=======================================================")
print(f"Book: {style_sheet['title']}")
print(f"Chapters reviewed: {len([r for r in chapter_reports if not r['skipped']])}/{len(chapter_reports)}")
print(f"Total issues found: {total_issues}")
print(f"Report saved to: {report_path}")
print(f"Style sheet saved to: {os.path.join(BOOK_DIR, STYLE_SHEET_FILENAME)}")
if AUTO_APPLY:
    print(f"Revised chapters saved to: {edited_dir}")
    if covers_copied:
        print(f"Cover option(s) copied to: {edited_dir} ({', '.join(covers_copied)})")
    if covers_missing:
        print(f"Cover option(s) not found (skipped): {', '.join(covers_missing)}")
    if edited_docx_path:
        print(f"Reader-ready manuscript.docx: {edited_docx_path}")
    elif edited_docx_error:
        print(f"Reader-ready manuscript.docx: FAILED ({edited_docx_error})")
else:
    print("Suggest-only mode — no chapter files were changed. Re-run with --auto-apply to apply fixes.")

# Outbound handoff: drop the editor model now rather than letting it sit in VRAM
# for Ollama's default ~5 minute idle timeout. scoring_agent.py runs immediately
# after this in the chain and loads its own model — without this, the two overlap
# for those first few minutes and the scorer gets pushed onto the CPU exactly the
# way the editor was. Added 2026-08-13 alongside the A1111 unload above.
unload_ollama_model(EDITOR_MODEL_NAME)
