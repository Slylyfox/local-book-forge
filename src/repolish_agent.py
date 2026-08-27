"""
Local Book Forge — Repolish Agent
------------------------------------
Design note (2026-08-12): a "fixed reroll" that takes the ratings/notes the
editor and scorer already produced and uses them to push a book to an even
better version of itself, rather than starting from scratch.

What it actually does, scoped down to what the data supports:

    scoring_agent.py's book_score.json only carries qualitative "why" text
    (a *_note string) for 3 of its 18 sub-metrics: arc_tracking,
    subplot_resolution, and thematic_cohesion — these are the three
    whole-book synthesis metrics (see scoring_agent.py section 9,
    whole_book_synthesis()). Every other sub-metric is a bare number with
    no explanation a revision pass could act on.

    So: this script reads the EDITED version's score, finds which of those
    three metrics fell below LOW_SCORE_THRESHOLD, turns each weak metric's
    note into an "issue" (reusing editorial_agent.py's revise_chapter()
    issue shape), and asks the editor model to revise every chapter with
    that same feedback in mind — uniformly, not just the "worst" chapters,
    since whole-book metrics like arc/theme are properties of the book as a
    whole, not of any one chapter. Output goes to a brand-new repolish/
    folder; edited/ is never touched. This is a single pass — it does not
    loop trying to hit a target score (see phase-2-roadmap.md item 6.5 for
    the discussion; looping is a possible future follow-up, not built here).

    Finally it re-runs scoring_agent.py (via --score-dir, see that script's
    own section 1) against repolish/ and prints a before/after comparison.

Precondition: the book must already have an edited/ folder AND a score for
the edited version (edited/book_score.json, or a root book_score.json with
scored_version == "edited"). If either is missing, this script exits with a
clear message telling you which stage to run first — it does not try to
run the editor or scorer for you first, to keep this script's job narrow
and its failure modes obvious.

Usage:
    python repolish_agent.py
        (defaults to the most recently modified book folder under
        output_books/, same heuristic as the other three scripts)

    python repolish_agent.py --book-dir "output_books\\your-book-folder"

Must be run from the `files` folder, same as the other three scripts.
Following this codebase's established convention, this script duplicates
rather than imports the handful of helper functions it shares with
editorial_agent.py and scoring_agent.py (apply_ebook_typography,
revise_chapter, extract_author_name, copy_cover_images, the docx builder,
looks_like_refusal, the outline/style-sheet loaders) — each pipeline script
is meant to run fully standalone. If you change one of these in its home
file, update the copy here too.
"""

import sys

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", line_buffering=True)
    sys.stderr.reconfigure(encoding="utf-8", line_buffering=True)

import argparse
import glob
import json
import os
import re
import shutil
import subprocess
from datetime import datetime

from crewai import Agent, Crew, Process, Task, LLM
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    sys.stdout.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(line_buffering=True, encoding="utf-8", errors="replace")
except AttributeError:
    pass  # Python < 3.7 fallback, not expected here

# =====================================================================
# 0. CONFIG
# =====================================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_ROOT = os.path.join(SCRIPT_DIR, "output_books")
SCORING_AGENT = os.path.join(SCRIPT_DIR, "scoring_agent.py")

# Same model as editorial_agent.py's editor — repolishing is the same kind of
# work (revise chapter prose against specific fixes), just fed a different
# source of fixes (scorer notes instead of the editor's own review passes).
REPOLISH_MODEL_NAME = "llama3.1-16k"

repolish_llm = LLM(
    model=f"ollama/{REPOLISH_MODEL_NAME}",
    base_url="http://localhost:11434",
)


def check_ollama_models_available(
    required_models: list, ollama_url: str = "http://localhost:11434", fix_hints: dict = None
) -> None:
    """Preflight check, identical to the copy in local-book-generator.py /
    editorial_agent.py / scoring_agent.py — see those for the full reasoning."""
    import urllib.request as _urllib_request
    import json as _json

    fix_hints = fix_hints or {}

    try:
        with _urllib_request.urlopen(f"{ollama_url.rstrip('/')}/api/tags", timeout=5) as resp:
            data = _json.loads(resp.read())
    except Exception as e:
        print(f"[FATAL] Could not reach Ollama at {ollama_url} to verify required models: {e}")
        print("        Is Ollama running? Start it (e.g. 'ollama serve') and try again.")
        sys.exit(1)

    available = {m.get("name", "") for m in data.get("models", [])}
    available_bare = {name.split(":")[0] for name in available if name}

    missing = [
        m for m in required_models
        if m not in available and m.split(":")[0] not in available_bare
    ]
    if missing:
        print(f"[FATAL] Required Ollama model(s) not found on this machine: {', '.join(missing)}")
        for m in missing:
            print(f"        Fix: {fix_hints.get(m, f'ollama pull {m}')}")
        print(f"        Currently pulled models: {', '.join(sorted(available)) or '(none)'}")
        sys.exit(1)

    print(f"[Config] Ollama models verified present: {', '.join(required_models)}")


check_ollama_models_available(
    [REPOLISH_MODEL_NAME],
    fix_hints={
        REPOLISH_MODEL_NAME: (
            "one-time local setup, not a registry pull: "
            "ollama create llama3.1-16k -f Modelfile.llama31-16k "
            "(requires 'ollama pull llama3.1' first if you haven't already)"
        ),
    },
)

repolish_editor = Agent(
    role="Senior Fiction Editor (Repolish Pass)",
    goal=(
        "Revise genre-fiction chapters to address specific whole-book weaknesses "
        "a quality scorer already identified — arc tracking, subplot resolution, "
        "and/or thematic cohesion — without losing what already works."
    ),
    backstory=(
        "You have edited hundreds of commercial genre novels. You've been handed a "
        "specific, narrow critique from a manuscript assessment and asked to make "
        "the book measurably stronger against exactly that critique — not to "
        "rewrite the book wholesale, invent new problems, or change voice, plot "
        "events, or length any more than necessary."
    ),
    llm=repolish_llm,
    verbose=True,
)

# Manuscript-formatting constants, kept identical to local-book-generator.py's
# Stage D / editorial_agent.py's auto-apply docx builder.
BODY_FONT_NAME = "Garamond"
BODY_FONT_SIZE_PT = 12
INCLUDE_TOC = True
COPYRIGHT_BOILERPLATE = (
    "This is a work of fiction. Names, characters, places, and incidents either are "
    "the product of the author's imagination or are used fictitiously. Any "
    "resemblance to actual events, locales, or persons, living or dead, is entirely "
    "coincidental."
)

MAX_REPOLISH_ATTEMPTS = 3  # same retry count as editorial_agent.py's revise_chapter()

FICTION_FRAMING = (
    "This is entirely fictional, original genre fiction being revised for a "
    "published novel — not real-world advice, instructions, or commentary.\n\n"
)

# The only 3 of scoring_agent.py's 18 sub-metrics that carry a qualitative
# *_note explaining the score (see whole_book_synthesis() in scoring_agent.py) —
# these are the only metrics this script has anything actionable to feed the
# editor. Score below this threshold -> treated as "weak enough to repolish".
# Same boundary scoring_agent.py itself uses for its own low-score warning.
LOW_SCORE_THRESHOLD = 70

NOTED_METRICS = {
    "arc_tracking": {
        "label": "Arc tracking",
        "fix": "Strengthen how this chapter's events visibly connect to and advance the "
               "overall story arc, so a reader could point to what it moves forward.",
    },
    "subplot_resolution": {
        "label": "Subplot resolution",
        "fix": "Make sure this chapter meaningfully advances or pays off at least one "
               "subplot thread rather than leaving it idle.",
    },
    "thematic_cohesion": {
        "label": "Thematic cohesion",
        "fix": "Reinforce the book's central theme(s) through this chapter's choices, "
               "imagery, or character decisions, without stating the theme outright.",
    },
}

# =====================================================================
# 1. CLI
# =====================================================================

parser = argparse.ArgumentParser(
    description="Repolish pass: feed the editor's/scorer's own feedback back into a "
                "second revision of an already-edited-and-scored book, output to repolish/."
)
parser.add_argument(
    "--book-dir",
    default=None,
    help="Path to a specific output_books/<slug-timestamp> folder. Defaults to the most "
         "recently modified book folder under output_books/.",
)
parser.add_argument(
    "--python-exe",
    default=sys.executable,
    help="Python interpreter to launch scoring_agent.py's re-score subprocess with. "
         "Defaults to the interpreter running this script.",
)
args = parser.parse_args()


def find_latest_book_dir(output_root: str) -> "str | None":
    if not os.path.isdir(output_root):
        return None
    candidates = [
        os.path.join(output_root, name)
        for name in os.listdir(output_root)
        if os.path.isdir(os.path.join(output_root, name))
        and os.path.isfile(os.path.join(output_root, name, "outline.json"))
    ]
    if not candidates:
        return None
    candidates.sort(key=os.path.getmtime, reverse=True)
    return candidates[0]


BOOK_DIR = args.book_dir or find_latest_book_dir(OUTPUT_ROOT)
if not BOOK_DIR or not os.path.isdir(BOOK_DIR) or not os.path.isfile(os.path.join(BOOK_DIR, "outline.json")):
    print(
        f"[ERROR] Could not find a book folder to repolish. Looked in '{OUTPUT_ROOT}' for the most "
        f"recently modified subfolder containing an outline.json, or pass one explicitly with "
        f"--book-dir \"path\\to\\output_books\\your-book-folder\"."
    )
    sys.exit(1)

print(f"[Config] Repolishing book at: {BOOK_DIR}")
print(f"[Config] Repolish model: {REPOLISH_MODEL_NAME}")

# =====================================================================
# 2. PRECONDITION: edited/ must exist and already have a score for the
#    EDITED version. Repolishing the raw draft directly isn't supported —
#    the whole point is to act on the editor's + scorer's own feedback,
#    which only exists once both of those have actually run.
# =====================================================================

EDITED_DIR = os.path.join(BOOK_DIR, "edited")
if not os.path.isdir(EDITED_DIR):
    print(
        f"[ERROR] No edited/ folder found at '{EDITED_DIR}'. Repolish acts on the editor's "
        f"revised chapters plus the scorer's feedback on them — run editorial_agent.py "
        f"--auto-apply on this book first, then scoring_agent.py, then try repolish_agent.py again."
    )
    sys.exit(1)


def _load_score_data() -> "dict | None":
    """Prefer edited/book_score.json (always strictly the edited-version score,
    per scoring_agent.py's own mirroring rule). Fall back to the book-root
    book_score.json only if it explicitly says it scored the edited version —
    never fabricate or assume."""
    edited_score_path = os.path.join(EDITED_DIR, "book_score.json")
    if os.path.isfile(edited_score_path):
        with open(edited_score_path, "r", encoding="utf-8") as f:
            return json.load(f)
    root_score_path = os.path.join(BOOK_DIR, "book_score.json")
    if os.path.isfile(root_score_path):
        with open(root_score_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if data.get("scored_version") == "edited":
            return data
    return None


score_data = _load_score_data()
if score_data is None:
    print(
        f"[ERROR] No score found for the edited version of this book (looked for "
        f"'{os.path.join(EDITED_DIR, 'book_score.json')}'). Run scoring_agent.py on this book "
        f"(after editorial_agent.py --auto-apply) so there's real feedback to repolish against, "
        f"then try repolish_agent.py again."
    )
    sys.exit(1)

# =====================================================================
# 3. FIND WEAK, ACTIONABLE METRICS
# =====================================================================

sub_metrics = score_data.get("sub_metrics", {})
weak_issues = []       # [{"category", "issue", "fix"}, ...] — fed to revise_chapter()
weak_summary = []      # [(label, score), ...] — for the printed report

for key, meta in NOTED_METRICS.items():
    entry = sub_metrics.get(key) or {}
    score = entry.get("score")
    note = entry.get("note")
    if score is None:
        continue
    if score < LOW_SCORE_THRESHOLD:
        weak_summary.append((meta["label"], score))
        weak_issues.append({
            "category": key,
            "issue": note or f"{meta['label']} scored {score}/100 with no further detail recorded.",
            "fix": meta["fix"],
        })

print(f"\n[Repolish] Overall score being repolished from: {score_data.get('overall_score')}/100")
for key, meta in NOTED_METRICS.items():
    entry = sub_metrics.get(key) or {}
    print(f"[Repolish]   {meta['label']}: {entry.get('score')}/100"
          + (" <- WEAK, will repolish against this" if entry.get('score') is not None
             and entry.get('score') < LOW_SCORE_THRESHOLD else ""))

if not weak_issues:
    print(
        f"\n[Repolish] All three metrics repolish can act on (arc tracking, subplot resolution, "
        f"thematic cohesion) are already at/above {LOW_SCORE_THRESHOLD}/100 — nothing to repolish. "
        f"(These are the only 3 of the scorer's 18 sub-metrics with qualitative feedback text; the "
        f"other 15 are just numbers with no 'why' to act on.) Not creating a repolish/ folder."
    )
    sys.exit(0)

print(f"\n[Repolish] {len(weak_issues)} weak metric(s) found — every chapter will be revised "
      f"against the same feedback (uniform pass, not just the 'worst' chapters, since these are "
      f"whole-book properties, not single-chapter ones).")

# =====================================================================
# 4. JSON / REFUSAL HELPERS (duplicated, same as the other three scripts)
# =====================================================================

REFUSAL_PATTERN = re.compile(
    r"^\s*(I cannot|I can[’']t|I can not|I[’']m sorry|I am sorry|"
    r"I[’']m unable|I am unable|As an AI|I won[’']t|"
    r"I[’']m not able|I don[’']t (feel comfortable|think I))",
    re.IGNORECASE,
)

META_INSTRUCTION_PATTERN = re.compile(
    r"(here's an outline|here is an outline|outline that can serve as a guide|"
    r"you can continue this story|add your own creative elements|"
    r"let me know if (i can help|you (need|would like)|there's anything else)|"
    r"i can (guide|help) you (on |through )?how to|"
    r"keeping to the word count target|feel free to (expand|continue)|"
    r"would you like me to (continue|write)|as a starting point for)",
    re.IGNORECASE,
)

OUTLINE_STRUCTURE_PATTERN = re.compile(r"^\s*(\*\*)?[IVX]{1,4}\.\s+", re.MULTILINE)


def looks_like_refusal(text: str) -> bool:
    stripped = text.strip()
    if not REFUSAL_PATTERN.match(stripped):
        return False
    if len(stripped) < 400:
        return True
    return bool(META_INSTRUCTION_PATTERN.search(stripped) or OUTLINE_STRUCTURE_PATTERN.search(stripped))


# =====================================================================
# 5. LOADING THE BOOK: outline, chapters, style sheet (read-only)
# =====================================================================


def load_outline(book_dir: str) -> dict:
    with open(os.path.join(book_dir, "outline.json"), "r", encoding="utf-8") as f:
        return json.load(f)


def parse_main_characters(main_characters: list) -> list:
    parsed = []
    for entry in main_characters or []:
        if ":" in entry:
            name, desc = entry.split(":", 1)
            parsed.append({"name": name.strip(), "description": desc.strip()})
        else:
            parsed.append({"name": entry.strip(), "description": ""})
    return parsed


CHAPTER_FILE_PATTERN = re.compile(r"chapter_(\d+)\.txt$")


def discover_chapters(chapters_dir: str) -> list:
    found = []
    for path in glob.glob(os.path.join(chapters_dir, "chapter_*.txt")):
        match = CHAPTER_FILE_PATTERN.search(os.path.basename(path))
        if match:
            found.append((int(match.group(1)), path))
    found.sort(key=lambda pair: pair[0])
    return found


def split_chapter_file(path: str) -> tuple:
    with open(path, "r", encoding="utf-8") as f:
        raw = f.read()
    parts = raw.split("\n\n", 1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1]
    return "", raw


STYLE_SHEET_FILENAME = "style_sheet.json"


def load_style_sheet_readonly(book_dir: str, outline: dict) -> dict:
    """Read-only, same as scoring_agent.py's copy — this script never writes to
    style_sheet.json, that stays editorial_agent.py's file to own."""
    path = os.path.join(book_dir, STYLE_SHEET_FILENAME)
    if os.path.isfile(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "title": outline.get("title", ""),
        "genre": outline.get("genre", ""),
        "premise": outline.get("premise", ""),
        "characters": parse_main_characters(outline.get("main_characters", [])),
        "established_facts": [],
        "style_terms": [],
    }


# =====================================================================
# 6. EBOOK TYPOGRAPHY FIXES (deterministic — duplicated from editorial_agent.py)
# =====================================================================


def apply_ebook_typography(text: str) -> tuple:
    notes = []
    fixed = text

    straight_double = fixed.count('"')
    if straight_double:
        out, is_open = [], True
        for ch in fixed:
            if ch == '"':
                out.append("“" if is_open else "”")
                is_open = not is_open
            else:
                out.append(ch)
        fixed = "".join(out)
        notes.append(f"Converted {straight_double} straight double-quote(s) to curly quotes.")

    straight_single = fixed.count("'")
    if straight_single:
        fixed = fixed.replace("'", "’")
        notes.append(f"Converted {straight_single} straight apostrophe(s)/single-quote(s) to curly.")

    collapsed = re.sub(r"\n{3,}", "\n\n", fixed)
    if collapsed != fixed:
        notes.append("Collapsed extra blank lines down to a single blank line between paragraphs.")
    fixed = collapsed

    stripped_lines = "\n".join(line.rstrip() for line in fixed.split("\n"))
    if stripped_lines != fixed:
        notes.append("Removed trailing whitespace from one or more lines.")
    fixed = stripped_lines

    return fixed, notes


# =====================================================================
# 7. THE REPOLISH REVISION PASS (same shape as editorial_agent.py's
#    revise_chapter(), but fed uniform whole-book feedback instead of a
#    per-chapter editorial review)
# =====================================================================


def revise_chapter(chapter_text: str, ch_num: int, ch_title: str, style_sheet: dict, issues: list) -> "str | None":
    """Returns None (not an exception) if the model refuses or every attempt
    fails to produce usable prose, so the caller can fall back to the
    original edited text instead of losing the chapter."""
    if not issues:
        return chapter_text

    issues_text = "\n".join(
        f"- [{i.get('category', '?')}] {i.get('issue', '')} -> {i.get('fix', '')}" for i in issues
    ) or "(none)"

    description = (
        f"{FICTION_FRAMING}"
        f"You are a senior editor producing a second, more polished revision of Chapter {ch_num} "
        f"('{ch_title}') of the novel '{style_sheet['title']}'. This chapter has already been through "
        f"one editorial pass — you are now addressing specific whole-book feedback a quality "
        f"assessment identified, which applies to every chapter in the book (not just this one). "
        f"Preserve the author's voice, plot events, dialogue content, and approximate length — only "
        f"change what the feedback below calls for.\n\n"
        f"Chapter text:\n\"\"\"\n{chapter_text}\n\"\"\"\n\n"
        f"Whole-book feedback to address in this chapter:\n{issues_text}\n\n"
        "Write the complete revised chapter text, incorporating this feedback. Write only the chapter "
        "prose itself — no title header, no meta-commentary, no summary of what you changed."
    )

    for attempt in range(1, MAX_REPOLISH_ATTEMPTS + 1):
        task = Task(
            description=description,
            expected_output=f"The complete repolished prose text of Chapter {ch_num}.",
            agent=repolish_editor,
        )
        crew = Crew(agents=[repolish_editor], tasks=[task], process=Process.sequential, verbose=True)
        revised = str(crew.kickoff())
        if not looks_like_refusal(revised) and len(revised.strip()) > 0.5 * len(chapter_text.strip()):
            return revised
        print(f"[WARN] Repolish attempt {attempt}/{MAX_REPOLISH_ATTEMPTS} for Chapter {ch_num} looks like "
              f"a refusal or came back too short." + (" Retrying..." if attempt < MAX_REPOLISH_ATTEMPTS else ""))
    print(f"[WARN] Could not get a usable repolish revision for Chapter {ch_num} after "
          f"{MAX_REPOLISH_ATTEMPTS} attempts. Keeping the existing edited/ text (unchanged) in repolish/.")
    return None


# =====================================================================
# 8. COVERS + MANUSCRIPT.DOCX (duplicated from editorial_agent.py, pointed
#    at repolish/ instead of edited/)
# =====================================================================

COVER_FILENAMES = (
    "cover_option_1.png", "cover_option_1.jpg",
    "cover_option_2.png", "cover_option_2.jpg",
    "cover_option_3.png", "cover_option_3.jpg",
)


def copy_cover_images(src_dir: str, dest_dir: str) -> "tuple[list, list]":
    copied, missing = [], []
    for filename in COVER_FILENAMES:
        src = os.path.join(src_dir, filename)
        if os.path.isfile(src):
            shutil.copyfile(src, os.path.join(dest_dir, filename))
            copied.append(filename)
        else:
            missing.append(filename)
    return copied, missing


def extract_author_name(book_dir: str) -> str:
    path = os.path.join(book_dir, "manuscript_raw.txt")
    if not os.path.isfile(path):
        return "(author name not found — manuscript_raw.txt missing; check Stage D output)"
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    m = re.search(r"^Copyright\s*©?\s*\d{4}\s+(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    m = re.search(r"^by\s+(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    return "(author name not found in manuscript_raw.txt — check the file manually)"


def _add_static_toc(document: "Document", chapters: list) -> None:
    for chapter in chapters:
        line = document.add_paragraph()
        line.add_run(f"Chapter {chapter['chapter_number']} — {chapter['title']}")


def build_repolish_docx_manuscript(repolish_dir: str, outline: dict, author_name: str,
                                    chapter_bodies: dict) -> str:
    """Reader-ready .docx built from the REPOLISHED chapter text, same
    formatting as editorial_agent.py's build_edited_docx_manuscript() /
    local-book-generator.py's build_docx_manuscript()."""
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(outline["title"])
    title_run.font.size = Pt(28)
    title_run.font.bold = True

    for _ in range(3):
        document.add_paragraph()

    by_p = document.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_run = by_p.add_run(f"by {author_name}")
    by_run.font.size = Pt(16)

    document.add_page_break()

    year = datetime.now().year
    copyright_p = document.add_paragraph()
    copyright_p.add_run(f"Copyright © {year} {author_name}").bold = True
    document.add_paragraph("All rights reserved.")
    document.add_paragraph()
    document.add_paragraph(COPYRIGHT_BOILERPLATE)
    document.add_page_break()

    if INCLUDE_TOC:
        toc_heading = document.add_paragraph()
        toc_heading.add_run("Contents").bold = True
        toc_heading.style = document.styles["Heading 1"]
        _add_static_toc(document, outline["chapters"])
        document.add_page_break()

    chapters = outline["chapters"]
    for chapter in chapters:
        ch_num = chapter["chapter_number"]
        body = chapter_bodies.get(ch_num)
        if body is None:
            print(f"[WARN] No repolished text in memory for chapter {ch_num}, skipping in repolish/manuscript.docx.")
            continue

        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.add_run(f"Chapter {ch_num}")

        subheading = document.add_paragraph()
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subheading.add_run(chapter["title"])
        sub_run.italic = True

        document.add_paragraph()

        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            p = document.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Inches(0) if i == 0 else Inches(0.3)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)

        if ch_num != chapters[-1]["chapter_number"]:
            document.add_page_break()

    output_path = os.path.join(repolish_dir, "manuscript.docx")
    document.save(output_path)
    return output_path


# =====================================================================
# 9. MAIN
# =====================================================================

outline = load_outline(BOOK_DIR)
style_sheet = load_style_sheet_readonly(BOOK_DIR, outline)

chapter_files = discover_chapters(EDITED_DIR)
if not chapter_files:
    print(f"[ERROR] No chapter_NN.txt files found in '{EDITED_DIR}'.")
    sys.exit(1)

REPOLISH_DIR = os.path.join(BOOK_DIR, "repolish")
os.makedirs(REPOLISH_DIR, exist_ok=True)

print(f"\n--- Repolish pass starting: {len(chapter_files)} chapter(s) ---")

repolish_chapter_bodies = {}  # ch_num -> final body text, for the docx builder

for ch_num, path in chapter_files:
    ch_title, chapter_text = split_chapter_file(path)
    print(f"\n--- Repolish pass: Chapter {ch_num}/{len(chapter_files)}: '{ch_title}' ---")

    if "NEEDS MANUAL REGENERATION" in chapter_text:
        print(f"[SKIP] Chapter {ch_num} is a placeholder from a failed writer pass — carrying it "
              f"through unchanged rather than repolishing a placeholder.")
        final_text = chapter_text
    else:
        revised = revise_chapter(chapter_text, ch_num, ch_title, style_sheet, weak_issues)
        final_text = revised if revised else chapter_text
        final_text, _ = apply_ebook_typography(final_text)  # always land the mechanical fixes

    out_path = os.path.join(REPOLISH_DIR, os.path.basename(path))
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write(f"{ch_title}\n\n{final_text}")
    print(f"[Saved] {out_path}")
    repolish_chapter_bodies[ch_num] = final_text

print(f"\n--- Repolish pass complete: {len(chapter_files)} chapter(s) processed ---")

# ---- Covers: prefer edited/ (self-contained) then fall back to the book root ----

covers_copied, covers_missing = copy_cover_images(EDITED_DIR, REPOLISH_DIR)
still_missing = []
for filename in covers_missing:
    root_src = os.path.join(BOOK_DIR, filename)
    if os.path.isfile(root_src):
        shutil.copyfile(root_src, os.path.join(REPOLISH_DIR, filename))
        covers_copied.append(filename)
    else:
        still_missing.append(filename)
if covers_copied:
    print(f"\n[Covers] Copied {len(covers_copied)}/3 cover option(s) into {REPOLISH_DIR}: {', '.join(covers_copied)}")
if still_missing:
    print(f"[Covers] Not found in edited/ or the book root, skipped: {', '.join(still_missing)}")

# ---- Reader-ready .docx built from the repolished text ----

author_name = extract_author_name(BOOK_DIR)
try:
    docx_path = build_repolish_docx_manuscript(REPOLISH_DIR, outline, author_name, repolish_chapter_bodies)
    print(f"[Saved] {docx_path}")
except Exception as e:
    print(f"[WARN] Could not build repolish/manuscript.docx: {e}. Chapter .txt files were still written "
          f"successfully — you can build the .docx by hand later, or re-run this script.")

# =====================================================================
# 10. RE-SCORE repolish/ AND PRINT A BEFORE/AFTER COMPARISON
# =====================================================================

print(f"\n{'=' * 70}")
print("[Repolish] Re-scoring repolish/ ...")
print(f"{'=' * 70}\n")

scorer_cmd = [args.python_exe, "-u", SCORING_AGENT, "--book-dir", BOOK_DIR, "--score-dir", REPOLISH_DIR]
print(f"[Repolish] $ {' '.join(scorer_cmd)}")
process = subprocess.Popen(
    scorer_cmd,
    cwd=SCRIPT_DIR,
    stdout=subprocess.PIPE,
    stderr=subprocess.STDOUT,
    text=True,
    encoding="utf-8",
    errors="replace",
    bufsize=1,
)
for line in process.stdout:
    print(f"[scorer] {line.rstrip(chr(10))}")
process.wait()

print(f"\n{'=' * 70}")
print("REPOLISH COMPLETE")
print(f"{'=' * 70}")
print(f"Book folder: {BOOK_DIR}")
print(f"Repolished chapters + manuscript: {REPOLISH_DIR}")

if process.returncode != 0:
    print(f"[WARN] Re-scoring repolish/ failed (exit {process.returncode}) — the repolished chapters "
          f"and manuscript.docx were still written successfully, just no fresh score to compare. See "
          f"the [scorer] output above for details.")
    sys.exit(process.returncode)

repolish_score_path = os.path.join(REPOLISH_DIR, "book_score.json")
if not os.path.isfile(repolish_score_path):
    print(f"[WARN] Scorer reported success but '{repolish_score_path}' wasn't found — can't print a "
          f"before/after comparison.")
    sys.exit(0)

with open(repolish_score_path, "r", encoding="utf-8") as f:
    after_data = json.load(f)

before_overall = score_data.get("overall_score")
after_overall = after_data.get("overall_score")


def _delta(before, after) -> str:
    if before is None or after is None:
        return ""
    diff = after - before
    if diff > 0:
        return f" (+{diff})"
    if diff < 0:
        return f" ({diff})"
    return " (unchanged)"


print(f"\nBEFORE/AFTER — {score_data.get('title', outline.get('title', ''))}")
print(f"  Overall score:  {before_overall}/100  ->  {after_overall}/100{_delta(before_overall, after_overall)}")
after_sub_metrics = after_data.get("sub_metrics", {})
for key, meta in NOTED_METRICS.items():
    before_entry = sub_metrics.get(key) or {}
    after_entry = after_sub_metrics.get(key) or {}
    before_score = before_entry.get("score")
    after_score = after_entry.get("score")
    marker = " <- targeted this pass" if key in {i["category"] for i in weak_issues} else ""
    print(f"  {meta['label']}: {before_score}/100  ->  {after_score}/100{_delta(before_score, after_score)}{marker}")

print(f"\nRepolish score report: {os.path.join(REPOLISH_DIR, 'book_score_report.txt')}")
print("Note: repolish/ is a new, separate folder — edited/ and its score were not modified.")
